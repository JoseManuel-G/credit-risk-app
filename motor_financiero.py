import os
import re
import math
import tempfile
import json
import pandas as pd

from datetime import datetime
from collections import Counter
from unidecode import unidecode
from openpyxl import load_workbook
from openpyxl.cell.cell import MergedCell
from openai import OpenAI


# Cliente OpenAI lazy: no falla al importar el módulo si aún no existe la API key.
_openai_client = None

def get_openai_client():
    global _openai_client
    if _openai_client is None:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError(
                "Falta OPENAI_API_KEY. Configúrala en Streamlit Cloud > App settings > Secrets."
            )
        _openai_client = OpenAI(api_key=api_key)
    return _openai_client


def load_financial_sheets(file_path, pyg_sheet_name="PyG", balance_sheet_name="Balance"):
    pyg_df = pd.read_excel(file_path, sheet_name=pyg_sheet_name)
    balance_df = pd.read_excel(file_path, sheet_name=balance_sheet_name)

    return pyg_df, balance_df

########## ANALISIS DE ESTRUCTURA DE ARCHIVOS #######3


# =========================================================
# HELPERS GENERALES
# =========================================================

def _is_missing(x):
    try:
        return x is None or pd.isna(x)
    except Exception:
        return x is None


def _clean_text(x):
    if _is_missing(x):
        return ""
    return str(x).replace("\xa0", " ").strip()


def _norm_text(x):
    return re.sub(r"\s+", " ", _clean_text(x).lower()).strip()


def _letters_only(s):
    return re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ]", "", str(s))


def _uppercase_ratio(s):
    letters = _letters_only(s)
    if not letters:
        return 0.0
    return sum(1 for c in letters if c.isupper()) / len(letters)


def _parse_number(x):
    if _is_missing(x):
        return None

    if isinstance(x, (int, float)) and not isinstance(x, bool):
        return float(x)

    s = _clean_text(x)
    if s == "":
        return None

    s = (
        s.replace("€", "")
         .replace("$", "")
         .replace("£", "")
         .replace("%", "")
         .replace(" ", "")
    )

    is_negative = False

    if re.fullmatch(r"\(.*\)", s):
        is_negative = True
        s = s[1:-1]

    if s.endswith("-"):
        is_negative = True
        s = s[:-1]

    value = None

    if re.fullmatch(r"-?\d{1,3}(\.\d{3})*(,\d+)?", s):
        try:
            value = float(s.replace(".", "").replace(",", "."))
        except Exception:
            value = None

    elif re.fullmatch(r"-?\d{1,3}(,\d{3})*(\.\d+)?", s):
        try:
            value = float(s.replace(",", ""))
        except Exception:
            value = None

    else:
        try:
            value = float(s.replace(",", "."))
        except Exception:
            value = None

    if value is None:
        return None

    return -abs(value) if is_negative else value


# =========================================================
# HELPERS PARA DETECTAR AÑOS Y COLUMNAS DE IMPORTES
# =========================================================

def _extract_year(x):
    """
    Extrae año desde formatos habituales:

    - 2024
    - "2024"
    - "31/12/2024"
    - "31-12-2024"
    - "2024-12-31"
    - "2024/12"
    - "12/2024"
    - "diciembre 2024"
    - "dic-24"
    - "Dec-24"
    - "Dec-21"
    - "FY2024"
    - "Ejercicio 2024"

    Devuelve año como int.
    """

    if _is_missing(x):
        return None

    # Fechas reales
    if isinstance(x, (pd.Timestamp, datetime)):
        y = int(x.year)
        return y if 1990 <= y <= 2100 else None

    if hasattr(x, "year") and not isinstance(x, (int, float, str)):
        try:
            y = int(x.year)
            return y if 1990 <= y <= 2100 else None
        except Exception:
            pass

    # Números tipo 2024, 202412, 20241231
    if isinstance(x, (int, float)) and not isinstance(x, bool):
        if float(x).is_integer():
            xi = int(x)

            if 1990 <= xi <= 2100:
                return xi

            s_num = str(xi)

            # 20241231
            if re.fullmatch(r"(19|20)\d{6}", s_num):
                y = int(s_num[:4])
                return y if 1990 <= y <= 2100 else None

            # 202412
            if re.fullmatch(r"(19|20)\d{4}", s_num):
                y = int(s_num[:4])
                return y if 1990 <= y <= 2100 else None

    s = _clean_text(x)
    if not s:
        return None

    s_norm = (
        s.lower()
        .replace("\xa0", " ")
        .replace("_", " ")
        .strip()
    )

    # Normalizar separadores
    s_norm = re.sub(r"[.\s]+", " ", s_norm)

    month_words = (
        "enero|ene|febrero|feb|marzo|mar|abril|abr|mayo|may|junio|jun|"
        "julio|jul|agosto|ago|septiembre|setiembre|sep|sept|octubre|oct|"
        "noviembre|nov|diciembre|dic|"
        "january|jan|february|feb|march|mar|april|apr|may|june|jun|july|jul|"
        "august|aug|september|sep|sept|october|oct|november|nov|december|dec"
    )

    # 1) Año completo en cualquier parte: 2024, FY2024, Dec-2023, 31/12/2024
    years = re.findall(r"\b(19\d{2}|20\d{2})\b", s_norm)

    valid_years = []
    for y in years:
        yi = int(y)
        if 1990 <= yi <= 2100:
            valid_years.append(yi)

    if valid_years:
        return max(valid_years)

    # 2) Formatos compactos: 20241231 / 202412
    compact_match = re.search(r"\b((?:19|20)\d{2})(?:0[1-9]|1[0-2])(?:[0-3]\d)?\b", s_norm)
    if compact_match:
        y = int(compact_match.group(1))
        return y if 1990 <= y <= 2100 else None

    # 3) Mes palabra + año corto: Dec-21, dic-24, December 23
    m = re.search(rf"\b(?:{month_words})[\s\-/]+(\d{{2}})\b", s_norm, re.I)
    if m:
        yy = int(m.group(1))
        yyyy = 2000 + yy if yy <= 79 else 1900 + yy
        return yyyy if 1990 <= yyyy <= 2100 else None

    # 4) Año corto + mes palabra: 21-Dec, 24-dic
    m = re.search(rf"\b(\d{{2}})[\s\-/]+(?:{month_words})\b", s_norm, re.I)
    if m:
        yy = int(m.group(1))
        yyyy = 2000 + yy if yy <= 79 else 1900 + yy
        return yyyy if 1990 <= yyyy <= 2100 else None

    # 5) Mes número + año corto: 12/21, 12-24
    m = re.search(r"\b(?:0?[1-9]|1[0-2])[\s\-/]+(\d{2})\b", s_norm)
    if m:
        yy = int(m.group(1))
        yyyy = 2000 + yy if yy <= 79 else 1900 + yy
        return yyyy if 1990 <= yyyy <= 2100 else None

    # 6) Año corto + mes número: 21/12, 24-12
    m = re.search(r"\b(\d{2})[\s\-/]+(?:0?[1-9]|1[0-2])\b", s_norm)
    if m:
        yy = int(m.group(1))
        yyyy = 2000 + yy if yy <= 79 else 1900 + yy
        return yyyy if 1990 <= yyyy <= 2100 else None

    return None

def _is_year_like_cell(x):
    """
    Detecta si una celda parece cabecera temporal:
    - año
    - fecha
    - mes/año
    - FY2024
    - ejercicio 2024
    """
    if _is_missing(x):
        return False

    s = _clean_text(x)
    if not s:
        return False

    s_norm = _norm_text(s)

    year = _extract_year(s_norm)
    if year is None:
        return False

    if re.fullmatch(r"(19|20)\d{2}", s_norm):
        return True

    temporal_words = [
        "ejercicio", "año", "ano", "periodo", "período", "fecha",
        "cierre", "cerrado", "finalizado", "terminado",
        "fy", "year", "period", "ended", "ending", "as of",
        "month", "mes", "quarter", "trimestre"
    ]

    if any(w in s_norm for w in temporal_words):
        return True

    if re.search(r"\d{1,2}[\-/]\d{1,2}[\-/](19|20)?\d{2}", s_norm):
        return True

    if re.search(r"(19|20)\d{2}[\-/]\d{1,2}([\-/]\d{1,2})?", s_norm):
        return True

    if re.search(r"\d{1,2}[\-/](19|20)?\d{2}", s_norm):
        return True

    month_words = [
        "enero", "ene", "febrero", "feb", "marzo", "mar", "abril", "abr",
        "mayo", "may", "junio", "jun", "julio", "jul", "agosto", "ago",
        "septiembre", "setiembre", "sep", "sept", "octubre", "oct",
        "noviembre", "nov", "diciembre", "dic",
        "january", "jan", "february", "march", "april", "june", "july",
        "august", "aug", "september", "october", "november", "december", "dec"
    ]

    if any(m in s_norm for m in month_words):
        return True

    return False


def _detect_amount_columns_from_matrix(
    matrix,
    text_col_idx=None,
    min_numeric_ratio=0.35,
    min_numeric_count=3,
    max_header_rows=15
):
    """
    Detecta TODAS las columnas que parecen importes.
    Ignora cabeceras tipo 2024 / 31-12-2024 al calcular ratios.
    """
    col_values = {}

    for row in matrix:
        for cell in row:
            c = cell["col_index"]
            col_values.setdefault(c, []).append(cell["value"])

    numeric_candidates = []

    for col_idx, values in col_values.items():
        if text_col_idx is not None and col_idx == text_col_idx:
            continue

        non_empty = []
        numeric_count = 0

        for row_pos, v in enumerate(values):
            if _clean_text(v) == "":
                continue

            if row_pos < max_header_rows and _is_year_like_cell(v):
                continue

            non_empty.append(v)

            if _parse_number(v) is not None:
                numeric_count += 1

        if not non_empty:
            continue

        numeric_ratio = numeric_count / max(len(non_empty), 1)

        if numeric_count >= min_numeric_count and numeric_ratio >= min_numeric_ratio:
            numeric_candidates.append((col_idx, numeric_ratio, numeric_count))

    numeric_candidates.sort(key=lambda x: (x[1], x[2]), reverse=True)

    amount_col_idxs = [x[0] for x in numeric_candidates]

    return amount_col_idxs, numeric_candidates


def _detect_years_for_amount_columns_from_matrix(
    matrix,
    amount_col_idxs,
    header_search_rows=25
):
    """
    Detecta el año asociado a cada columna de importes.

    Regla importante:
    - Primero mira SOLO en la propia columna.
    - Esto evita que Dec-23 contamine a Dec-22 o Dec-21.
    """

    col_years = {c: None for c in amount_col_idxs}

    if not matrix:
        return col_years

    max_rows = min(header_search_rows, len(matrix))

    for amount_col in amount_col_idxs:
        candidates = []

        for r in range(max_rows):
            row = matrix[r]

            for cell in row:
                if cell["col_index"] != amount_col:
                    continue

                y = _extract_year(cell["value"])

                if y is not None:
                    candidates.append({
                        "year": y,
                        "row": cell["row_index"],
                        "value": cell["value"]
                    })

        if candidates:
            # Cogemos el último año encontrado en esa columna.
            # Si hay varios textos, normalmente el año real de cabecera será el más reciente.
            col_years[amount_col] = max(c["year"] for c in candidates)

    return col_years

def _assign_amount_columns_to_fixed_periods(
    amount_col_idxs,
    detected_col_years,
    latest_year=None,
    latest_position="left"
):
    """
    Convierte columnas reales de Excel a campos fijos:

    amount   = año más reciente
    amount_1 = año anterior
    amount_2 = dos años antes

    Regla principal:
    - Si hay años detectados y son únicos, manda el año, no la posición.
      Ejemplo:
        Dec-21 | Dec-22 | Dec-23
      Resultado:
        Dec-23 -> amount
        Dec-22 -> amount_1
        Dec-21 -> amount_2

    Fallback:
    - Si no se detectan años suficientes, usa latest_position.
    """

    amount_col_idxs = sorted(list(amount_col_idxs))

    if not amount_col_idxs:
        return {}, {}

    # Pares columna-año detectados
    col_year_pairs = [
        (col, detected_col_years.get(col))
        for col in amount_col_idxs
        if detected_col_years.get(col) is not None
    ]

    detected_years = [year for _, year in col_year_pairs]

    # -----------------------------------------------------
    # CASO BUENO: hay años detectados únicos
    # -----------------------------------------------------
    if len(col_year_pairs) >= 2 and len(detected_years) == len(set(detected_years)):
        # Ordenar por año descendente: 2023, 2022, 2021...
        ordered_cols = [
            col for col, year in sorted(
                col_year_pairs,
                key=lambda x: x[1],
                reverse=True
            )
        ]

        latest_year_detected = max(detected_years)

        # Añadir posibles columnas numéricas sin año al final
        missing_cols = [c for c in amount_col_idxs if c not in ordered_cols]

        if missing_cols:
            if latest_position == "left":
                missing_cols = sorted(missing_cols)
            else:
                missing_cols = sorted(missing_cols, reverse=True)

            ordered_cols += missing_cols

        latest_year_final = latest_year if latest_year is not None else latest_year_detected

    # -----------------------------------------------------
    # CASO FALLBACK: no hay años suficientes o están duplicados
    # -----------------------------------------------------
    else:
        if latest_year is None:
            if detected_years:
                latest_year_final = max(detected_years)
            else:
                latest_year_final = datetime.now().year
        else:
            latest_year_final = latest_year

        if latest_position == "left":
            ordered_cols = sorted(amount_col_idxs)
        elif latest_position == "right":
            ordered_cols = sorted(amount_col_idxs, reverse=True)
        else:
            raise ValueError("latest_position debe ser 'left' o 'right'")

    # -----------------------------------------------------
    # Construir mapping final
    # -----------------------------------------------------
    fixed_mapping = {}
    year_mapping = {}

    for i, col in enumerate(ordered_cols):
        field = "amount" if i == 0 else f"amount_{i}"
        year_field = "amount_year" if i == 0 else f"amount_{i}_year"

        detected_year = detected_col_years.get(col)

        # Si la columna tenía año claro, lo usamos.
        # Si no, inferimos secuencia desde el último año.
        if detected_year is not None:
            year = detected_year
        else:
            year = latest_year_final - i

        fixed_mapping[col] = field
        year_mapping[field] = year
        year_mapping[year_field] = year

    return fixed_mapping, year_mapping

# =========================================================
# 1) LEER HOJA CON OPENPYXL Y EXTRAER FORMATO REAL
# =========================================================

def read_sheet_with_format_openpyxl(file_path, sheet_name):
    wb = load_workbook(file_path, data_only=True)
    ws = wb[sheet_name]

    merged_map = {}

    for merged_range in ws.merged_cells.ranges:
        min_col, min_row, max_col, max_row = merged_range.bounds
        main_cell = ws.cell(row=min_row, column=min_col)

        for r in range(min_row, max_row + 1):
            for c in range(min_col, max_col + 1):
                merged_map[(r, c)] = main_cell

    matrix = []

    for row_idx, row in enumerate(ws.iter_rows(), start=1):
        row_cells = []

        for cell in row:
            if isinstance(cell, MergedCell):
                main_cell = merged_map.get((cell.row, cell.column))
                if main_cell is not None:
                    value = main_cell.value
                    font = main_cell.font
                    alignment = main_cell.alignment
                    number_format = main_cell.number_format
                else:
                    value = None
                    font = None
                    alignment = None
                    number_format = None
            else:
                value = cell.value
                font = cell.font
                alignment = cell.alignment
                number_format = cell.number_format

            font_color = None
            try:
                if font and font.color and hasattr(font.color, "rgb"):
                    font_color = font.color.rgb
            except Exception:
                font_color = None

            try:
                column_letter = cell.column_letter
            except Exception:
                column_letter = None

            row_cells.append({
                "row_index": row_idx,
                "col_index": cell.column,
                "column_letter": column_letter,
                "value": value,
                "is_bold": bool(font.bold) if font else False,
                "font_size": float(font.sz) if font and font.sz is not None else None,
                "font_color": font_color,
                "indent": float(alignment.indent) if alignment and alignment.indent is not None else 0.0,
                "number_format": number_format,
            })

        matrix.append(row_cells)

    return matrix


# =========================================================
# 2) DETECTAR AUTOMÁTICAMENTE COLUMNA DE TEXTO Y COLUMNAS DE IMPORTES
# =========================================================

def detect_main_columns_from_matrix(
    matrix,
    min_numeric_ratio=0.35,
    min_numeric_count=3,
    latest_year=None,
    latest_position="left",
    header_search_rows=15
):
    def text_score(values):
        non_empty = [_clean_text(v) for v in values if _clean_text(v) != ""]
        if not non_empty:
            return -1

        score = 0
        total = 0

        for v in non_empty[:250]:
            total += 1

            has_letters = bool(re.search(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]", v))
            parsed_num = _parse_number(v)
            is_year_like = _is_year_like_cell(v)

            if has_letters:
                score += 2

            if parsed_num is None:
                score += 1

            if is_year_like:
                score -= 1

        return score / max(total, 1)

    col_values = {}

    for row in matrix:
        for cell in row:
            c = cell["col_index"]
            col_values.setdefault(c, []).append(cell["value"])

    text_candidates = []

    for col_idx, values in col_values.items():
        text_candidates.append((col_idx, text_score(values)))

    text_candidates.sort(key=lambda x: x[1], reverse=True)

    text_col_idx = text_candidates[0][0] if text_candidates else 1

    amount_col_idxs, numeric_candidates = _detect_amount_columns_from_matrix(
        matrix=matrix,
        text_col_idx=text_col_idx,
        min_numeric_ratio=min_numeric_ratio,
        min_numeric_count=min_numeric_count,
        max_header_rows=header_search_rows
    )

    detected_years = _detect_years_for_amount_columns_from_matrix(
        matrix=matrix,
        amount_col_idxs=amount_col_idxs,
        header_search_rows=header_search_rows
    )

    amount_fixed_mapping, amount_year_mapping = _assign_amount_columns_to_fixed_periods(
        amount_col_idxs=amount_col_idxs,
        detected_col_years=detected_years,
        latest_year=latest_year,
        latest_position=latest_position
    )

    amount_col_idx = None

    for col_idx, fixed_name in amount_fixed_mapping.items():
        if fixed_name == "amount":
            amount_col_idx = col_idx
            break

    return {
        "text_col_idx": text_col_idx,
        "amount_col_idx": amount_col_idx,
        "amount_col_idxs": amount_col_idxs,
        "amount_fixed_mapping": amount_fixed_mapping,
        "amount_year_mapping": amount_year_mapping,
        "detected_years": detected_years,
        "text_candidates": text_candidates,
        "numeric_candidates": numeric_candidates,
        "latest_year": amount_year_mapping.get("amount_year"),
        "latest_position": latest_position,
    }


# =========================================================
# 3) CONVERTIR LA HOJA A DATAFRAME ESTRUCTURADO
# =========================================================

def build_structured_df_from_matrix(
    matrix,
    text_col_idx=None,
    amount_col_idx=None,
    amount_col_idxs=None,
    latest_year=None,
    latest_position="left",
    min_numeric_ratio=0.35,
    min_numeric_count=3,
    header_search_rows=15
):
    detection = detect_main_columns_from_matrix(
        matrix=matrix,
        min_numeric_ratio=min_numeric_ratio,
        min_numeric_count=min_numeric_count,
        latest_year=latest_year,
        latest_position=latest_position,
        header_search_rows=header_search_rows
    )

    if text_col_idx is None:
        text_col_idx = detection["text_col_idx"]

    if amount_col_idxs is None:
        if amount_col_idx is not None:
            amount_col_idxs = [amount_col_idx]
        else:
            amount_col_idxs = detection["amount_col_idxs"]

    amount_fixed_mapping = detection["amount_fixed_mapping"]
    amount_year_mapping = detection["amount_year_mapping"]

    if amount_col_idx is not None and amount_col_idx not in amount_fixed_mapping:
        amount_fixed_mapping = {amount_col_idx: "amount"}
        year = latest_year if latest_year is not None else datetime.now().year
        amount_year_mapping = {
            "amount": year,
            "amount_year": year
        }

    rows_out = []

    for row in matrix:
        text_cell = None
        amount_cells_by_col = {}

        for cell in row:
            if cell["col_index"] == text_col_idx:
                text_cell = cell

            if cell["col_index"] in amount_col_idxs:
                amount_cells_by_col[cell["col_index"]] = cell

        if text_cell is None:
            continue

        label = _clean_text(text_cell["value"])

        if label == "":
            continue

        row_out = {
            "source_row": text_cell["row_index"],
            "label": label,
            "amount": None,
            "amount_raw": None,
            "is_bold": bool(text_cell["is_bold"]),
            "font_size": text_cell["font_size"],
            "font_color": text_cell["font_color"],
            "indent": text_cell["indent"],
            "x_left": text_cell["col_index"],
        }

        for col_idx, fixed_field in amount_fixed_mapping.items():
            cell = amount_cells_by_col.get(col_idx)
            raw_value = cell["value"] if cell is not None else None

            is_header_temporal_value = (
                cell is not None
                and cell.get("row_index", 999999) <= header_search_rows
                and _is_year_like_cell(raw_value)
            )

            if is_header_temporal_value:
                parsed_value = None
            else:
                parsed_value = _parse_number(raw_value)

            raw_field = f"{fixed_field}_raw"

            row_out[fixed_field] = parsed_value
            row_out[raw_field] = raw_value

            year_field = "amount_year" if fixed_field == "amount" else f"{fixed_field}_year"
            row_out[year_field] = amount_year_mapping.get(year_field)

            source_col_field = "amount_source_col" if fixed_field == "amount" else f"{fixed_field}_source_col"
            row_out[source_col_field] = col_idx

        rows_out.append(row_out)

    df = pd.DataFrame(rows_out)

    if df.empty:
        return df, detection

    amount_cols = []
    amount_raw_cols = []
    amount_year_cols = []
    amount_source_cols = []

    max_amount_idx = -1

    for c in df.columns:
        if c == "amount":
            max_amount_idx = max(max_amount_idx, 0)
        else:
            m = re.fullmatch(r"amount_(\d+)", str(c))
            if m:
                max_amount_idx = max(max_amount_idx, int(m.group(1)))

    for i in range(max_amount_idx + 1):
        if i == 0:
            amount_cols.append("amount")
            amount_raw_cols.append("amount_raw")
            amount_year_cols.append("amount_year")
            amount_source_cols.append("amount_source_col")
        else:
            amount_cols.append(f"amount_{i}")
            amount_raw_cols.append(f"amount_{i}_raw")
            amount_year_cols.append(f"amount_{i}_year")
            amount_source_cols.append(f"amount_{i}_source_col")

    base_cols = [
        "source_row",
        "label",
    ]

    format_cols = [
        "is_bold",
        "font_size",
        "font_color",
        "indent",
        "x_left",
    ]

    ordered_cols = []

    for c in base_cols + amount_cols + amount_year_cols + amount_raw_cols + amount_source_cols + format_cols:
        if c in df.columns and c not in ordered_cols:
            ordered_cols.append(c)

    remaining_cols = [c for c in df.columns if c not in ordered_cols]

    df = df[ordered_cols + remaining_cols]

    amount_value_cols = [c for c in amount_cols if c in df.columns]

    first_real_idx = None

    for i in range(len(df)):
        has_any_amount = False

        for c in amount_value_cols:
            if pd.notna(df.iloc[i].get(c)):
                has_any_amount = True
                break

        if has_any_amount:
            first_real_idx = i
            break

    if first_real_idx is not None:
        df = df.iloc[first_real_idx:].reset_index(drop=True)

    return df, detection


# =========================================================
# 4) REGLAS DE NUMERACIÓN Y TOTALES
# =========================================================

LVL1_ENUM_RES = [
    re.compile(r"^\s*\d+\.\s+\S"),
    re.compile(r"^\s*[IVXLCM]+\.\s+\S", re.I),
    re.compile(r"^\s*[A-Z]\)\s+\S"),
]

LVL2_ENUM_RES = [
    re.compile(r"^\s*\d+\.\d+(\.\d+)?\s+\S"),
    re.compile(r"^\s*[A-Z]\.\d+(\.\d+)?\)?\s+\S", re.I),
    re.compile(r"^\s*[a-z]\)\s+\S"),
    re.compile(r"^\s*\d+\)\s+\S"),
]


TOTAL_RES = [
    re.compile(r"^\s*resultado de explotación(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado explotaci[oó]n(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado operativo(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado de operaciones(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado financiero(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado antes de impuestos(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado antes de impuesto sobre beneficios(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado antes de impuestos sobre beneficios(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado despu[eé]s de impuestos(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado del ejercicio(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado neto(?:\b.*)?$", re.I),
    re.compile(r"^\s*resultado consolidado(?:\b.*)?$", re.I),
    re.compile(r"^\s*beneficio del ejercicio(?:\b.*)?$", re.I),
    re.compile(r"^\s*beneficio neto(?:\b.*)?$", re.I),
    re.compile(r"^\s*beneficio consolidado(?:\b.*)?$", re.I),
    re.compile(r"^\s*p[eé]rdida del ejercicio(?:\b.*)?$", re.I),
    re.compile(r"^\s*p[eé]rdida neta(?:\b.*)?$", re.I),
    re.compile(r"^\s*p[eé]rdida consolidada(?:\b.*)?$", re.I),
    re.compile(r"^\s*margen bruto(?:\b.*)?$", re.I),
    re.compile(r"^\s*margen de explotaci[oó]n(?:\b.*)?$", re.I),
    re.compile(r"^\s*margen operativo(?:\b.*)?$", re.I),

    re.compile(r"^\s*gross margin(?:\b.*)?$", re.I),
    re.compile(r"^\s*gross profit(?:\b.*)?$", re.I),
    re.compile(r"^\s*operating profit(?:\b.*)?$", re.I),
    re.compile(r"^\s*operating income(?:\b.*)?$", re.I),
    re.compile(r"^\s*operating result(?:\b.*)?$", re.I),
    re.compile(r"^\s*profit from operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*income from operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*result from operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*financial result(?:\b.*)?$", re.I),
    re.compile(r"^\s*finance result(?:\b.*)?$", re.I),
    re.compile(r"^\s*net financial result(?:\b.*)?$", re.I),
    re.compile(r"^\s*profit before tax(?:es)?(?:\b.*)?$", re.I),
    re.compile(r"^\s*profit before income tax(?:es)?(?:\b.*)?$", re.I),
    re.compile(r"^\s*earnings before tax(?:es)?(?:\b.*)?$", re.I),
    re.compile(r"^\s*income before tax(?:es)?(?:\b.*)?$", re.I),
    re.compile(r"^\s*profit after tax(?:es)?(?:\b.*)?$", re.I),
    re.compile(r"^\s*income after tax(?:es)?(?:\b.*)?$", re.I),
    re.compile(r"^\s*net profit(?:\b.*)?$", re.I),
    re.compile(r"^\s*net loss(?:\b.*)?$", re.I),
    re.compile(r"^\s*net income(?:\b.*)?$", re.I),
    re.compile(r"^\s*net earnings(?:\b.*)?$", re.I),
    re.compile(r"^\s*(?:consolidated|group|total|adjusted|normalized|recurring)\s+net income(?:\b.*)?$", re.I),
    re.compile(r"^\s*(?:consolidated|group|total|adjusted|normalized|recurring)\s+net profit(?:\b.*)?$", re.I),
    re.compile(r"^\s*(?:consolidated|group|total|adjusted|normalized|recurring)\s+net earnings(?:\b.*)?$", re.I),
    re.compile(r"^\s*(?:consolidated|group|total|adjusted|normalized|recurring)\s+net loss(?:\b.*)?$", re.I),
    re.compile(r"^\s*profit for the year(?:\b.*)?$", re.I),
    re.compile(r"^\s*profit for the period(?:\b.*)?$", re.I),
    re.compile(r"^\s*loss for the year(?:\b.*)?$", re.I),
    re.compile(r"^\s*loss for the period(?:\b.*)?$", re.I),
    re.compile(r"^\s*profit from continuing operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*profit from discontinued operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*income from continuing operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*income from discontinued operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*loss from continuing operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*loss from discontinued operations(?:\b.*)?$", re.I),
    re.compile(r"^\s*comprehensive income(?:\b.*)?$", re.I),
    re.compile(r"^\s*total comprehensive income(?:\b.*)?$", re.I),

    re.compile(r"^\s*ebit(?:\b.*)?$", re.I),
    re.compile(r"^\s*ebita(?:\b.*)?$", re.I),
    re.compile(r"^\s*ebitda(?:\b.*)?$", re.I),
    re.compile(r"^\s*adjusted ebit(?:\b.*)?$", re.I),
    re.compile(r"^\s*adjusted ebita(?:\b.*)?$", re.I),
    re.compile(r"^\s*adjusted ebitda(?:\b.*)?$", re.I),
    re.compile(r"^\s*normalized ebit(?:\b.*)?$", re.I),
    re.compile(r"^\s*normalized ebita(?:\b.*)?$", re.I),
    re.compile(r"^\s*normalized ebitda(?:\b.*)?$", re.I),
    re.compile(r"^\s*recurring ebit(?:\b.*)?$", re.I),
    re.compile(r"^\s*recurring ebita(?:\b.*)?$", re.I),
    re.compile(r"^\s*recurring ebitda(?:\b.*)?$", re.I),

    re.compile(r"^\s*total\b.*$", re.I),
    re.compile(r"^\s*subtotal\b.*$", re.I),
    re.compile(r"^\s*margen\b.*$", re.I),
    re.compile(r"^\s*margin\b.*$", re.I),
]


TOTAL_EXCLUDE_RES = [
    re.compile(r"^\s*other operating income\b.*$", re.I),
    re.compile(r"^\s*other income\b.*$", re.I),
    re.compile(r"^\s*interest income\b.*$", re.I),
    re.compile(r"^\s*finance income\b.*$", re.I),
    re.compile(r"^\s*financial income\b.*$", re.I),
    re.compile(r"^\s*income tax expense\b.*$", re.I),
    re.compile(r"^\s*tax expense\b.*$", re.I),
    re.compile(r"^\s*income taxes\b.*$", re.I),
    re.compile(r"^\s*current income tax\b.*$", re.I),
    re.compile(r"^\s*deferred income tax\b.*$", re.I),
    re.compile(r"^\s*operating expenses\b.*$", re.I),
    re.compile(r"^\s*operating costs\b.*$", re.I),
    re.compile(r"^\s*cost of sales\b.*$", re.I),
    re.compile(r"^\s*cost of goods sold\b.*$", re.I),
    re.compile(r"^\s*staff costs\b.*$", re.I),
    re.compile(r"^\s*personnel expenses\b.*$", re.I),
    re.compile(r"^\s*employee benefits expense\b.*$", re.I),
    re.compile(r"^\s*depreciation\b.*$", re.I),
    re.compile(r"^\s*amortization\b.*$", re.I),
    re.compile(r"^\s*depreciation and amortization\b.*$", re.I),

    re.compile(r"^\s*otros ingresos de explotaci[oó]n\b.*$", re.I),
    re.compile(r"^\s*otros ingresos\b.*$", re.I),
    re.compile(r"^\s*ingresos financieros\b.*$", re.I),
    re.compile(r"^\s*gastos financieros\b.*$", re.I),
    re.compile(r"^\s*gasto por impuesto\b.*$", re.I),
    re.compile(r"^\s*gasto por impuestos\b.*$", re.I),
    re.compile(r"^\s*impuesto sobre beneficios\b.*$", re.I),
    re.compile(r"^\s*impuestos sobre beneficios\b.*$", re.I),
    re.compile(r"^\s*gastos de explotaci[oó]n\b.*$", re.I),
    re.compile(r"^\s*coste de ventas\b.*$", re.I),
    re.compile(r"^\s*coste de las ventas\b.*$", re.I),
    re.compile(r"^\s*costes de ventas\b.*$", re.I),
    re.compile(r"^\s*gastos de personal\b.*$", re.I),
    re.compile(r"^\s*aprovisionamientos\b.*$", re.I),
    re.compile(r"^\s*amortizaci[oó]n\b.*$", re.I),
    re.compile(r"^\s*deterioro\b.*$", re.I),
]


STRONG_TOTAL_PHRASES = [
    r"\bnet income\b",
    r"\bnet profit\b",
    r"\bnet earnings\b",
    r"\bnet loss\b",
    r"\bgross profit\b",
    r"\bgross margin\b",
    r"\boperating profit\b",
    r"\boperating income\b",
    r"\boperating result\b",
    r"\bprofit from operations\b",
    r"\bincome from operations\b",
    r"\bprofit before tax(?:es)?\b",
    r"\bprofit before income tax(?:es)?\b",
    r"\bincome before tax(?:es)?\b",
    r"\bearnings before tax(?:es)?\b",
    r"\bprofit after tax(?:es)?\b",
    r"\bincome after tax(?:es)?\b",
    r"\bprofit for the year\b",
    r"\bprofit for the period\b",
    r"\bloss for the year\b",
    r"\bloss for the period\b",
    r"\bcomprehensive income\b",
    r"\btotal comprehensive income\b",
    r"\bebitda\b",
    r"\bebita\b",
    r"\bebit\b",

    r"\bresultado neto\b",
    r"\bresultado consolidado\b",
    r"\bresultado del ejercicio\b",
    r"\bresultado de explotaci[oó]n\b",
    r"\bresultado explotaci[oó]n\b",
    r"\bresultado operativo\b",
    r"\bresultado financiero\b",
    r"\bresultado antes de impuestos\b",
    r"\bresultado antes de impuesto sobre beneficios\b",
    r"\bresultado antes de impuestos sobre beneficios\b",
    r"\bresultado despu[eé]s de impuestos\b",
    r"\bbeneficio neto\b",
    r"\bbeneficio consolidado\b",
    r"\bbeneficio del ejercicio\b",
    r"\bp[eé]rdida neta\b",
    r"\bp[eé]rdida consolidada\b",
    r"\bp[eé]rdida del ejercicio\b",
    r"\bmargen bruto\b",
    r"\bmargen operativo\b",
    r"\bmargen de explotaci[oó]n\b",
]


def _matches_any(res_list, text):
    return any(r.match(text) for r in res_list)


def _has_strong_total_phrase(text_norm):
    s = _norm_text(text_norm)

    if not s:
        return False

    if len(s.split()) > 18:
        return False

    return any(re.search(pattern, s, re.I) for pattern in STRONG_TOTAL_PHRASES)


def _is_total(text_norm):
    if not text_norm:
        return False

    text_norm = _norm_text(text_norm)

    if _matches_any(TOTAL_EXCLUDE_RES, text_norm):
        return False

    if _matches_any(TOTAL_RES, text_norm):
        return True

    if _has_strong_total_phrase(text_norm):
        return True

    return False


def _is_lvl1_numbering(text):
    return _matches_any(LVL1_ENUM_RES, text)


def _is_lvl2_numbering(text):
    return _matches_any(LVL2_ENUM_RES, text)


def _detect_enum_depth(text):
    s = _clean_text(text)

    if re.match(r"^\d+\.\d+\.\d+(\.\d+)*\s+\S", s):
        return 3
    if re.match(r"^\d+\.\d+\s+\S", s):
        return 2
    if re.match(r"^[A-Z]\.\d+\.\d+(\.\d+)*\)?\s+\S", s, re.I):
        return 3
    if re.match(r"^[A-Z]\.\d+(\.\d+)?\)?\s+\S", s, re.I):
        return 2
    if _is_lvl1_numbering(s):
        return 1
    if _is_lvl2_numbering(s):
        return 2

    return None


# =========================================================
# 5) ASIGNAR NIVELES
# =========================================================

def assign_levels_with_orientation_from_structured_df(
    df: pd.DataFrame,
    statement_type: str = "pyg",
    lookaround_lines: int = 6,
    strong_lvl1_threshold: int = 2,
    verbose: bool = False
):
    out = df.copy()

    if out.empty:
        return out

    out["clean_label"] = out["label"].apply(_clean_text)
    out["norm_label"] = out["label"].apply(_norm_text)
    out["line_role"] = "normal"
    out["level"] = pd.NA
    out["parent_name"] = pd.NA
    out["parent_index"] = pd.NA
    out["document_structure"] = pd.NA
    out["visual_signal_count"] = 0
    out["enum_depth"] = out["clean_label"].apply(_detect_enum_depth)
    out["level_debug"] = ""

    def is_mostly_upper(label):
        return _uppercase_ratio(label) >= 0.75

    def row_signals(row):
        label = row["clean_label"]
        is_bold = bool(row.get("is_bold", False))
        main_enum = _is_lvl1_numbering(label)
        sub_enum = _is_lvl2_numbering(label)
        mostly_upper = is_mostly_upper(label)

        lvl1_signals = 0
        if is_bold:
            lvl1_signals += 1
        if mostly_upper:
            lvl1_signals += 1
        if main_enum:
            lvl1_signals += 1

        return {
            "is_bold": is_bold,
            "mostly_upper": mostly_upper,
            "main_enum": main_enum,
            "sub_enum": sub_enum,
            "lvl1_signals": lvl1_signals,
        }

    def is_total_row(row):
        st = _norm_text(statement_type)

        if st != "pyg":
            return False

        return _is_total(row["norm_label"])

    def is_lvl1(row, has_weak_group):
        s = row_signals(row)

        if s["main_enum"]:
            return True

        if s["sub_enum"]:
            return False

        if has_weak_group and s["lvl1_signals"] >= 1:
            return True

        if s["lvl1_signals"] >= strong_lvl1_threshold:
            return True

        return False

    def is_lvl2(row, has_strong_group):
        s = row_signals(row)

        if s["sub_enum"]:
            return True

        if has_strong_group and (not s["is_bold"]) and (not s["mostly_upper"]):
            return True

        return False

    for idx, row in out.iterrows():
        if is_total_row(row):
            out.at[idx, "line_role"] = "total"
            out.at[idx, "level"] = pd.NA
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
            out.at[idx, "level_debug"] = "strict_total_no_parent"

    non_total = out[out["line_role"] != "total"].copy()

    strong_count = 0
    weak_count = 0

    for _, row in non_total.iterrows():
        s = row_signals(row)

        if s["main_enum"] or s["is_bold"] or s["mostly_upper"]:
            strong_count += 1

        if s["sub_enum"] or ((not s["is_bold"]) and (not s["mostly_upper"])):
            weak_count += 1

    has_strong_group = strong_count >= 2
    has_weak_group = weak_count >= 2

    for idx, row in out.iterrows():
        if row["line_role"] == "total":
            continue

        s = row_signals(row)
        out.at[idx, "visual_signal_count"] = s["lvl1_signals"]

        if is_lvl1(row, has_weak_group):
            out.at[idx, "level"] = 1
            out.at[idx, "level_debug"] = f"lvl1_{s}"

        elif is_lvl2(row, has_strong_group):
            out.at[idx, "level"] = 2
            out.at[idx, "level_debug"] = f"lvl2_{s}"

        else:
            if s["lvl1_signals"] >= 1:
                out.at[idx, "level"] = 1
                out.at[idx, "level_debug"] = f"lvl1_fallback_{s}"
            else:
                out.at[idx, "level"] = 2
                out.at[idx, "level_debug"] = f"lvl2_fallback_{s}"

    non_total_mask = out["line_role"] != "total"

    has_lvl1 = (
        out.loc[non_total_mask, "level"]
        .dropna()
        .apply(lambda x: int(x) == 1)
        .any()
    )

    has_lvl2 = (
        out.loc[non_total_mask, "level"]
        .dropna()
        .apply(lambda x: int(x) == 2)
        .any()
    )

    if (not has_lvl1) and has_lvl2:
        for idx, row in out.loc[non_total_mask].iterrows():
            if pd.notna(row["level"]) and int(row["level"]) == 2:
                out.at[idx, "level"] = 1
                out.at[idx, "parent_name"] = pd.NA
                out.at[idx, "parent_index"] = pd.NA
                out.at[idx, "level_debug"] = str(out.at[idx, "level_debug"]) + " | promoted_lvl2_to_lvl1_no_lvl1_exists"

    lvl1_positions = [
        i for i in range(len(out))
        if out.iloc[i]["line_role"] != "total"
        and pd.notna(out.iloc[i]["level"])
        and int(out.iloc[i]["level"]) == 1
    ]

    document_structure = "forward"

    if lvl1_positions:
        forward_score = 0
        backward_score = 0

        for pos in lvl1_positions:
            j = pos + 1
            steps = 0

            while j < len(out) and steps < lookaround_lines:
                rowj = out.iloc[j]

                if rowj["line_role"] == "total":
                    j += 1
                    steps += 1
                    continue

                if pd.notna(rowj["level"]) and int(rowj["level"]) == 1:
                    break

                if pd.notna(rowj["level"]) and int(rowj["level"]) == 2:
                    forward_score += max(1, lookaround_lines - steps)

                j += 1
                steps += 1

            j = pos - 1
            steps = 0

            while j >= 0 and steps < lookaround_lines:
                rowj = out.iloc[j]

                if rowj["line_role"] == "total":
                    j -= 1
                    steps += 1
                    continue

                if pd.notna(rowj["level"]) and int(rowj["level"]) == 1:
                    break

                if pd.notna(rowj["level"]) and int(rowj["level"]) == 2:
                    backward_score += max(1, lookaround_lines - steps)

                j -= 1
                steps += 1

        if backward_score > forward_score:
            document_structure = "backward"
        else:
            document_structure = "forward"

        out["orientation_debug"] = (
            f"forward_score={forward_score}; "
            f"backward_score={backward_score}; "
            f"lookaround_lines={lookaround_lines}"
        )
    else:
        out["orientation_debug"] = "no_lvl1_positions"

    out["document_structure"] = document_structure

    out["parent_name"] = pd.NA
    out["parent_index"] = pd.NA

    if document_structure == "forward":
        current_parent_idx = None
        current_parent_name = None

        for idx, row in out.iterrows():
            if row["line_role"] == "total":
                continue

            if pd.isna(row["level"]):
                continue

            lvl = int(row["level"])

            if lvl == 1:
                current_parent_idx = idx
                current_parent_name = row["clean_label"]
                out.at[idx, "parent_name"] = pd.NA
                out.at[idx, "parent_index"] = pd.NA

            elif lvl == 2:
                if current_parent_idx is not None:
                    out.at[idx, "parent_name"] = current_parent_name
                    out.at[idx, "parent_index"] = current_parent_idx
                else:
                    out.at[idx, "parent_name"] = pd.NA
                    out.at[idx, "parent_index"] = pd.NA

    else:
        next_parent_idx = None
        next_parent_name = None

        for idx in range(len(out) - 1, -1, -1):
            row = out.iloc[idx]

            if row["line_role"] == "total":
                continue

            if pd.isna(row["level"]):
                continue

            lvl = int(row["level"])

            if lvl == 1:
                next_parent_idx = idx
                next_parent_name = row["clean_label"]
                out.at[idx, "parent_name"] = pd.NA
                out.at[idx, "parent_index"] = pd.NA

            elif lvl == 2:
                if next_parent_idx is not None:
                    out.at[idx, "parent_name"] = next_parent_name
                    out.at[idx, "parent_index"] = next_parent_idx
                else:
                    out.at[idx, "parent_name"] = pd.NA
                    out.at[idx, "parent_index"] = pd.NA

    return out


# =========================================================
# 6) ORIENTACIÓN + PADRES
# =========================================================

def _assign_parents_by_detected_orientation(df_levels: pd.DataFrame, lookaround_lines: int = 20):
    out = df_levels.copy().reset_index(drop=True)

    if out.empty:
        out["document_structure"] = "forward"
        out["parent_name"] = pd.NA
        out["parent_index"] = pd.NA
        out["orientation_debug"] = "empty_df"
        return out

    out["parent_name"] = pd.NA
    out["parent_index"] = pd.NA
    out["document_structure"] = "forward"
    out["orientation_debug"] = ""

    def _level_at(pos):
        if pos < 0 or pos >= len(out):
            return None

        row = out.iloc[pos]

        if row.get("line_role") == "total":
            return None

        if pd.isna(row.get("level")):
            return None

        try:
            return int(row["level"])
        except Exception:
            return None

    def _positions_by_level(target_level):
        return [
            i for i in range(len(out))
            if _level_at(i) == target_level
        ]

    def _count_children_after_parent(parent_pos, child_level, max_lines):
        parent_level = _level_at(parent_pos)
        if parent_level is None:
            return 0

        count = 0
        steps = 0
        j = parent_pos + 1

        while j < len(out) and steps < max_lines:
            lvl_j = _level_at(j)

            if lvl_j == parent_level:
                break

            if lvl_j == child_level:
                count += 1

            j += 1
            steps += 1

        return count

    def _count_children_before_parent(parent_pos, child_level, max_lines):
        parent_level = _level_at(parent_pos)
        if parent_level is None:
            return 0

        count = 0
        steps = 0
        j = parent_pos - 1

        while j >= 0 and steps < max_lines:
            lvl_j = _level_at(j)

            if lvl_j == parent_level:
                break

            if lvl_j == child_level:
                count += 1

            j -= 1
            steps += 1

        return count

    lvl1_positions = _positions_by_level(1)
    lvl2_positions = _positions_by_level(2)

    if not lvl1_positions or not lvl2_positions:
        out["document_structure"] = "forward"
        out["orientation_debug"] = (
            f"not_enough_levels; "
            f"lvl1_count={len(lvl1_positions)}; "
            f"lvl2_count={len(lvl2_positions)}"
        )
        return out

    first_lvl1 = min(lvl1_positions)
    last_lvl1 = max(lvl1_positions)

    lvl2_before_first_lvl1 = sum(1 for i in lvl2_positions if i < first_lvl1)
    lvl2_after_last_lvl1 = sum(1 for i in lvl2_positions if i > last_lvl1)

    section_forward_score = 0
    section_backward_score = 0

    for pos in lvl1_positions:
        section_forward_score += _count_children_after_parent(
            parent_pos=pos,
            child_level=2,
            max_lines=lookaround_lines
        )

        section_backward_score += _count_children_before_parent(
            parent_pos=pos,
            child_level=2,
            max_lines=lookaround_lines
        )

    strong_backward_score = lvl2_before_first_lvl1 * 10
    strong_forward_score = lvl2_after_last_lvl1 * 10

    forward_score = section_forward_score + strong_forward_score
    backward_score = section_backward_score + strong_backward_score

    if lvl2_before_first_lvl1 > 0 and lvl2_after_last_lvl1 == 0:
        structure = "backward"
    elif lvl2_after_last_lvl1 > 0 and lvl2_before_first_lvl1 == 0:
        structure = "forward"
    else:
        structure = "backward" if backward_score > forward_score else "forward"

    if structure == "forward":
        out = _assign_parents_forward(out)
    else:
        out = _assign_parents_backward(out)

    out["document_structure"] = structure
    out["orientation_debug"] = (
        f"structure={structure}; "
        f"forward_score={forward_score}; "
        f"backward_score={backward_score}; "
        f"section_forward_score={section_forward_score}; "
        f"section_backward_score={section_backward_score}; "
        f"strong_forward_score={strong_forward_score}; "
        f"strong_backward_score={strong_backward_score}; "
        f"lvl2_before_first_lvl1={lvl2_before_first_lvl1}; "
        f"lvl2_after_last_lvl1={lvl2_after_last_lvl1}; "
        f"lookaround_lines={lookaround_lines}"
    )

    return out


def _assign_parents_forward(df_levels: pd.DataFrame):
    out = df_levels.copy().reset_index(drop=True)

    out["parent_name"] = pd.NA
    out["parent_index"] = pd.NA

    last_parent_by_level = {}

    for idx, row in out.iterrows():
        if row.get("line_role") == "total":
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
            continue

        if pd.isna(row.get("level")):
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
            continue

        try:
            lvl = int(row["level"])
        except Exception:
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
            continue

        if lvl == 1:
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
        else:
            parent_level = lvl - 1
            parent_idx = last_parent_by_level.get(parent_level)

            if parent_idx is not None:
                out.at[idx, "parent_name"] = out.at[parent_idx, "clean_label"]
                out.at[idx, "parent_index"] = parent_idx
            else:
                out.at[idx, "parent_name"] = pd.NA
                out.at[idx, "parent_index"] = pd.NA

        last_parent_by_level[lvl] = idx

        for k in list(last_parent_by_level.keys()):
            if k > lvl:
                last_parent_by_level[k] = None

    return out


def _assign_parents_backward(df_levels: pd.DataFrame):
    out = df_levels.copy().reset_index(drop=True)

    out["parent_name"] = pd.NA
    out["parent_index"] = pd.NA

    next_parent_by_level = {}

    for idx in range(len(out) - 1, -1, -1):
        row = out.iloc[idx]

        if row.get("line_role") == "total":
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
            continue

        if pd.isna(row.get("level")):
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
            continue

        try:
            lvl = int(row["level"])
        except Exception:
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
            continue

        if lvl == 1:
            out.at[idx, "parent_name"] = pd.NA
            out.at[idx, "parent_index"] = pd.NA
        else:
            parent_level = lvl - 1
            parent_idx = next_parent_by_level.get(parent_level)

            if parent_idx is not None:
                out.at[idx, "parent_name"] = out.at[parent_idx, "clean_label"]
                out.at[idx, "parent_index"] = parent_idx
            else:
                out.at[idx, "parent_name"] = pd.NA
                out.at[idx, "parent_index"] = pd.NA

        next_parent_by_level[lvl] = idx

        for k in list(next_parent_by_level.keys()):
            if k > lvl:
                next_parent_by_level[k] = None

    return out


# =========================================================
# 7) WRAPPER FINAL
# =========================================================

def assign_levels_from_excel_openpyxl(
    file_path,
    sheet_name,
    lookaround_lines=20,
    strong_lvl1_threshold=4,
    verbose=False,
    latest_year=None,
    latest_position="left"
):
    """
    Lee una hoja de Excel, detecta:
    - columna de texto
    - columnas numéricas de importes
    - años asociados a esas columnas

    Devuelve un DataFrame con:
    - amount      = último año / periodo más reciente
    - amount_1    = año anterior
    - amount_2    = dos años antes
    - amount_year, amount_1_year, amount_2_year...
    - niveles jerárquicos
    - padres
    """
    matrix = read_sheet_with_format_openpyxl(file_path, sheet_name)

    df_structured, detected = build_structured_df_from_matrix(
        matrix,
        latest_year=latest_year,
        latest_position=latest_position
    )

    if verbose:
        print(f"\n--- {sheet_name} ---")
        print("Detected columns:")
        print(f"  text_col_idx       = {detected['text_col_idx']}")
        print(f"  amount_col_idx     = {detected['amount_col_idx']}")
        print(f"  amount_col_idxs    = {detected['amount_col_idxs']}")
        print(f"  detected_years     = {detected['detected_years']}")
        print(f"  amount_mapping     = {detected['amount_fixed_mapping']}")
        print(f"  amount_years       = {detected['amount_year_mapping']}")
        print(f"  latest_position    = {detected['latest_position']}")
        print("Top text candidates:", detected["text_candidates"][:5])
        print("Top numeric candidates:", detected["numeric_candidates"][:5])

    statement_type = "balance" if "balance" in _norm_text(sheet_name) else "pyg"

    df_levels = assign_levels_with_orientation_from_structured_df(
        df_structured,
        statement_type=statement_type,
        lookaround_lines=lookaround_lines,
        strong_lvl1_threshold=strong_lvl1_threshold,
        verbose=verbose
    )

    df_levels = _assign_parents_by_detected_orientation(
        df_levels,
        lookaround_lines=lookaround_lines
    )

    if verbose and not df_levels.empty:
        preview_cols = [
            c for c in [
                "source_row",
                "label",
                "amount",
                "amount_year",
                "amount_1",
                "amount_1_year",
                "amount_2",
                "amount_2_year",
                "line_role",
                "level",
                "parent_name",
                "parent_index",
                "document_structure",
                "orientation_debug",
                "visual_key",
                "visual_score",
                "visual_signal_count",
                "enum_depth",
                "level_debug",
            ] if c in df_levels.columns
        ]

        print(df_levels[preview_cols].head(100).to_string(index=False))

    return df_levels


# =========================================================
# BLOQUE X - HELPERS DE NORMALIZACIÓN Y MATCHING DE CONCEPTOS
# =========================================================

def normalize_text(text):
    if text is None:
        return ""
    text = str(text).strip().lower()
    text = unidecode(text)
    text = re.sub(r"[^a-z0-9\s%]", " ", text)
    text = " ".join(text.split())
    return text


def token_set(text):
    return set(normalize_text(text).split())


def is_exact_match(raw, alias):
    return normalize_text(raw) == normalize_text(alias)


def is_token_subset_match(raw, alias):
    raw_tokens = token_set(raw)
    alias_tokens = token_set(alias)
    return len(alias_tokens) > 0 and alias_tokens.issubset(raw_tokens)

# =====================================================
# DICCIONARIO GRANDE - VERSION MAS COMPLETA
# Español + Inglés
# =====================================================

LINE_MAPPING_RULES = [
    # =================================================
    # PYG
    # =================================================

    {
        "canonical": "operating_revenue",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "ingresos_operativos",
        "aliases": [
            # Español
            "importe neto cifra de negocios",
            "importe neto de la cifra de negocios",
            "cifra de negocios",
            "ventas",
            "ventas netas",
            "ingresos de explotacion",
            "ingresos de explotación",
            "ingresos ordinarios",
            "ingresos operativos",

            # Inglés
            "operating revenue",
            "revenue",
            "revenues",
            "net revenue",
            "net revenues",
            "net sales",
            "sales",
            "turnover",
            "operating income revenue"
        ]
    },

    {
        "canonical": "financial_income",
        "statement": "pyg",
        "section": "financiero",
        "aggregate": "ingresos_financieros",
        "aliases": [
            # Español
            "ingresos financieros",
            "ingresos por intereses",
            "ingresos financieros por intereses",
            "intereses a favor",

            # Inglés
            "financial income",
            "finance income",
            "interest income",
            "financial revenues",
            "income from financial assets"
        ]
    },

    {
        "canonical": "cost_of_goods_sold",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "coste_ventas",
        "aliases": [
            # Español
            "coste de ventas",
            "coste ventas",
            "costes de ventas",
            "coste de la mercancia",
            "coste de la mercancía",
            "coste mercancia",
            "coste mercancía",
            "aprovisionamientos",
            "consumo de mercaderias",
            "consumo de mercaderías",

            # Inglés
            "cost of sales",
            "cost of goods sold",
            "cogs",
            "cost of revenue",
            "cost of merchandise"
        ]
    },

    {
        "canonical": "gross_profit",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "margen_bruto",
        "aliases": [
            # Español
            "margen bruto",
            "resultado bruto",
            "beneficio bruto",

            # Inglés
            "gross profit",
            "gross margin",
            "gross result"
        ]
    },

    {
        "canonical": "gross_margin_pct",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "margen_bruto_pct",
        "aliases": [
            # Español
            "margen bruto porcentual",
            "margen bruto %",
            "porcentaje margen bruto",

            # Inglés
            "gross margin %",
            "gross margin percentage",
            "gross profit margin"
        ]
    },

    {
        "canonical": "personnel_expenses",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "gastos_personal",
        "aliases": [
            # Español
            "gastos de personal",
            "gastos personal",
            "costes de personal",
            "sueldos y salarios",
            "salarios",
            "cargas sociales",

            # Inglés
            "staff costs",
            "personnel expenses",
            "employee expenses",
            "employee benefits expense",
            "payroll",
            "salary expenses",
            "wages and salaries"
        ]
    },

    {
        "canonical": "operating_expenses",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "gastos_operativos",
        "aliases": [
            # Español
            "gastos de explotacion",
            "gastos de explotación",
            "otros gastos de explotacion",
            "otros gastos de explotación",
            "gastos operativos",
            "gastos generales",

            # Inglés
            "operating expenses",
            "operating costs",
            "administrative expenses",
            "opex",
            "selling general administrative",
            "selling general and administrative expenses",
            "sg&a",
            "sga"
        ]
    },

    {
        "canonical": "other_gains_losses",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "otros_resultados_operativos",
        "aliases": [
            # Español
            "otras perdidas y ganancias",
            "otras pérdidas y ganancias",
            "otras perdidas y ganancias netas",
            "otras pérdidas y ganancias netas",
            "otros resultados",
            "otros resultados operativos",
            "otros ingresos y gastos",

            # Inglés
            "other gains and losses",
            "other operating gains losses",
            "other operating result",
            "other operating income and expenses",
            "other income and expenses"
        ]
    },

    {
        "canonical": "ebitda",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "ebitda",
        "aliases": [
            # Español
            "ebitda",
            "resultado operativo ebitda",
            "resultado bruto de explotacion",
            "resultado bruto de explotación",

            # Inglés
            "ebitda",
            "adjusted ebitda",
            "normalized ebitda",
            "recurring ebitda",
            "earnings before interest taxes depreciation amortization",
            "earnings before interest, taxes, depreciation and amortization",
            "earnings before interest tax depreciation amortization",
            "earnings before interest, tax, depreciation and amortization",
            "earnings before interest taxes depreciation and amortization"
        ]
    },

    {
        "canonical": "ebita",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "ebita",
        "aliases": [
            "ebita",
            "adjusted ebita",
            "normalized ebita",
            "recurring ebita",
            "earnings before interest taxes and amortization",
            "earnings before interest, taxes and amortization",
            "earnings before interest tax and amortization",
            "earnings before interest, tax and amortization"
        ]
    },

    {
        "canonical": "depreciation_amortization",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "amortizacion",
        "aliases": [
            # Español
            "amortizacion",
            "amortización",
            "amortizaciones",
            "amortizaciones y depreciaciones",
            "depreciacion",
            "depreciación",
            "depreciaciones",
            "dotaciones para amortizaciones",

            # Inglés
            "depreciation",
            "amortization",
            "depreciation and amortization",
            "depreciation amortization",
            "d&a"
        ]
    },

    {
        "canonical": "ebit",
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "ebit",
        "aliases": [
            # Español
            "ebit",
            "resultado de explotacion",
            "resultado de explotación",
            "resultados de explotacion",
            "resultados de explotación",
            "resultado explotacion",
            "resultado explotación",
            "resultado operativo",
            "beneficio de explotacion",
            "beneficio de explotación",
            "beneficio operativo",

            # Inglés
            "ebit",
            "operating income",
            "operating profit",
            "operating result",
            "profit from operations",
            "income from operations",
            "earnings before interest and taxes",
            "earnings before interest taxes",
            "earnings before interest and tax",
            "earnings before interest tax"
        ]
    },

    {
        "canonical": "interest_expense",
        "statement": "pyg",
        "section": "financiero",
        "aggregate": "gastos_financieros",
        "aliases": [
            # Español
            "gastos financieros",
            "costes financieros",
            "gastos por intereses",
            "intereses de deuda",

            # Inglés
            "interest expense",
            "financial expenses",
            "finance costs",
            "finance cost",
            "cost of debt",
            "interest costs",
            "interest expense on debt"
        ]
    },

    {
        "canonical": "equity_method_result",
        "statement": "pyg",
        "section": "financiero",
        "aggregate": "resultado_participadas",
        "aliases": [
            # Español
            "resultados de inversiones contabilizadas por el metodo de la participacion",
            "resultados de inversiones contabilizadas por el método de la participación",
            "resultado por participaciones",
            "resultado de participadas",

            # Inglés
            "equity method result",
            "share of profit of associates",
            "share of results of associates",
            "share of profit from associates",
            "share of profit of equity accounted investees"
        ]
    },

    {
        "canonical": "profit_before_tax",
        "statement": "pyg",
        "section": "fiscal",
        "aggregate": "bai",
        "aliases": [
            # Español
            "beneficio antes de impuestos",
            "beneficio antes de impuesto",
            "resultado antes de impuestos",
            "resultado antes de impuesto",
            "resultado antes de impuestos sobre beneficios",
            "resultado antes de impuesto sobre beneficios",
            "resultado antes de impuestos de actividades continuadas",
            "resultado antes de impuesto de actividades continuadas",
            "bai",

            # Inglés - profit
            "profit before tax",
            "profit before taxes",
            "profit before income tax",
            "profit before income taxes",
            "profit before taxation",
            "pre tax profit",
            "pre-tax profit",
            "pbt",

            # Inglés - income
            "income before tax",
            "income before taxes",
            "income before income tax",
            "income before income taxes",
            "income before taxation",
            "pre tax income",
            "pre-tax income",

            # Inglés - earnings
            "earnings before tax",
            "earnings before taxes",
            "earnings before income tax",
            "earnings before income taxes",
            "earnings before taxation",
            "pre tax earnings",
            "pre-tax earnings",
            "ebt",
            "earnings before taxes ebt"
        ]
    },

    {
        "canonical": "income_tax",
        "statement": "pyg",
        "section": "fiscal",
        "aggregate": "tax",
        "aliases": [
            # Español
            "impuesto sobre beneficios",
            "impuestos sobre beneficios",
            "gasto por impuesto",
            "gasto por impuestos",
            "gasto por impuesto sobre beneficios",
            "impuesto corriente",
            "impuesto diferido",

            # Inglés
            "income tax",
            "income taxes",
            "income tax expense",
            "tax expense",
            "current tax",
            "deferred tax",
            "current income tax",
            "deferred income tax"
        ]
    },

    {
        "canonical": "net_income",
        "statement": "pyg",
        "section": "neto",
        "aggregate": "beneficio_neto",
        "aliases": [
            # Español
            "resultado neto del ejercicio",
            "resultado neto",
            "resultado neto consolidado",
            "resultado consolidado",
            "beneficio neto",
            "beneficio neto del ejercicio",
            "beneficio consolidado",
            "perdida neta",
            "pérdida neta",
            "perdida del ejercicio",
            "pérdida del ejercicio",
            "resultado despues de impuestos",
            "resultado después de impuestos",
            "beneficio despues de impuestos",
            "beneficio después de impuestos",

            # Inglés - net income / profit
            "net income",
            "consolidated net income",
            "net income on continuing operations",
            "net income on discontinuing operations",
            "net income from continuing operations",
            "net income from discontinued operations",
            "net profit",
            "consolidated net profit",
            "net loss",
            "consolidated net loss",

            # Inglés - profit/loss for period
            "profit for the year",
            "profit for the period",
            "loss for the year",
            "loss for the period",
            "profit after tax",
            "profit after taxes",
            "profit after income tax",
            "profit after income taxes",
            "income after tax",
            "income after taxes",
            "income after income tax",
            "income after income taxes",

            # Inglés - earnings
            "net earnings",
            "consolidated net earnings",
            "earnings after tax",
            "earnings after taxes",
            "earnings after income tax",
            "earnings after income taxes",
            "earnings for the year",
            "earnings for the period",
            "profit attributable to shareholders",
            "earnings attributable to shareholders"
        ]
    },

    {
        "canonical": "minority_income",
        "statement": "pyg",
        "section": "neto",
        "aggregate": "resultado_minoritarios",
        "aliases": [
            # Español
            "resultado atribuido a accionistas minoritarios",
            "resultado atribuido a minoritarios",
            "resultado atribuido a intereses minoritarios",

            # Inglés
            "minority interests result",
            "minority interest",
            "profit attributable to non controlling interests",
            "profit attributable to non-controlling interests",
            "net income attributable to non controlling interests",
            "net income attributable to non-controlling interests"
        ]
    },

    {
        "canonical": "net_income_parent",
        "statement": "pyg",
        "section": "neto",
        "aggregate": "beneficio_neto_dominante",
        "aliases": [
            # Español
            "resultado neto atribuido a la sociedad dominante",
            "resultado atribuido a la dominante",
            "resultado atribuido a los propietarios de la dominante",
            "beneficio atribuido a la dominante",

            # Inglés
            "net income attributable to parent",
            "net income attributable to owners of the parent",
            "profit attributable to parent",
            "profit attributable to owners of the parent",
            "earnings attributable to parent",
            "earnings attributable to owners of the parent"
        ]
    },

    {
        "canonical": "eps",
        "statement": "pyg",
        "section": "neto",
        "aggregate": "bpa",
        "aliases": [
            # Español
            "beneficio por accion",
            "beneficio por acción",
            "bpa",

            # Inglés
            "eps",
            "earnings per share",
            "basic earnings per share",
            "diluted earnings per share"
        ]
    },

    # =================================================
    # BALANCE - ACTIVO NO CORRIENTE
    # =================================================

    {
        "canonical": "non_current_assets_total",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "activo_no_corriente",
        "aliases": [
            "activos no corrientes",
            "activo no corriente",
            "non current assets",
            "non-current assets",
            "fixed assets total"
        ]
    },

    {
        "canonical": "right_of_use_assets",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "right_of_use_assets",
        "aliases": [
            "derecho de uso",
            "activos por derecho de uso",
            "right of use assets",
            "right-of-use assets",
            "right of use"
        ]
    },

    {
        "canonical": "intangible_assets",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "intangibles",
        "aliases": [
            "activos intangibles",
            "otros activos intangibles",
            "inmovilizado intangible",
            "intangible assets",
            "other intangible assets",
            "intangibles"
        ]
    },

    {
        "canonical": "goodwill",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "goodwill",
        "aliases": [
            "fondo de comercio",
            "goodwill"
        ]
    },

    {
        "canonical": "property_plant_equipment",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "inmovilizado_material",
        "aliases": [
            "inmovilizado",
            "inmovilizado material",
            "inmovilizado tangible",
            "activo fijo",
            "property plant equipment",
            "property plant and equipment",
            "ppe",
            "fixed assets",
            "tangible fixed assets"
        ]
    },

    {
        "canonical": "investment_properties",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "propiedades_inversion",
        "aliases": [
            "propiedades de inversion",
            "propiedades de inversión",
            "inmuebles de inversion",
            "inmuebles de inversión",
            "investment properties"
        ]
    },

    {
        "canonical": "financial_assets_lp",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "activos_financieros_lp",
        "aliases": [
            "inversiones financieras",
            "inversiones financieras lp",
            "inversiones financieras a largo plazo",
            "activos financieros no corrientes",
            "financial investments lp",
            "long term financial assets",
            "long-term financial assets",
            "non current financial assets",
            "non-current financial assets"
        ]
    },

    {
        "canonical": "other_non_current_assets",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "otros_activos_no_corrientes",
        "aliases": [
            "otros activos no corrientes",
            "other non current assets",
            "other non-current assets"
        ]
    },

    {
        "canonical": "deferred_tax_assets",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "activos_impuesto_diferido",
        "aliases": [
            "activos por impuesto diferido",
            "activos por impuestos diferidos",
            "deferred tax assets"
        ]
    },

    # =================================================
    # BALANCE - ACTIVO CORRIENTE
    # =================================================

    {
        "canonical": "current_assets_total",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "activo_corriente",
        "aliases": [
            "activos corrientes",
            "activo corriente",
            "non fixed assets",
            "current assets"
        ]
    },

    {
        "canonical": "inventory",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "existencias",
        "aliases": [
            "existencias",
            "inventario",
            "inventarios",
            "inventory",
            "inventories",
            "stock"
        ]
    },

    {
        "canonical": "trade_receivables",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "clientes",
        "aliases": [
            "deudores",
            "clientes",
            "clientes comerciales",
            "cuentas a cobrar",
            "accounts receivable",
            "trade receivables",
            "receivables",
            "trade and other receivables"
        ]
    },

    {
        "canonical": "income_tax_receivable",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "impuesto_corriente_activo",
        "aliases": [
            "activos por impuesto sobre beneficios corriente",
            "impuesto corriente activo",
            "activo por impuesto corriente",
            "current tax asset",
            "income tax receivable",
            "tax receivable"
        ]
    },

    {
        "canonical": "other_current_assets",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "otros_activos_corrientes",
        "aliases": [
            "otros activos corrientes",
            "other current assets"
        ]
    },

    {
        "canonical": "other_financial_assets_cp",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "otros_activos_financieros_cp",
        "aliases": [
            "otros activos financieros",
            "otros activos financieros cp",
            "activos financieros corrientes",
            "other financial assets",
            "current financial assets"
        ]
    },

    {
        "canonical": "temporary_financial_investments",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "inversiones_financieras_cp",
        "aliases": [
            "inversiones financieras temporales",
            "inversiones financieras a corto plazo",
            "temporary financial investments",
            "short term financial investments",
            "short-term financial investments"
        ]
    },

    {
        "canonical": "cash_and_equivalents",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "caja",
        "aliases": [
            "caja",
            "efectivo",
            "tesoreria",
            "tesorería",
            "efectivo y equivalentes",
            "cash",
            "cash equivalents",
            "cash and cash equivalents",
            "cash and short term deposits",
            "cash and short-term deposits"
        ]
    },

    {
        "canonical": "total_assets",
        "statement": "balance",
        "section": "activo_total",
        "aggregate": "activo_total",
        "aliases": [
            "total activo",
            "activo total",
            "total assets",
            "assets total"
        ]
    },

    # =================================================
    # BALANCE - PATRIMONIO NETO
    # =================================================

    {
        "canonical": "equity_total",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "patrimonio_neto",
        "aliases": [
            "patrimonio neto",
            "fondos propios",
            "fundos propios",
            "equity",
            "shareholders equity",
            "shareholders' equity",
            "total equity",
            "net equity"
        ]
    },

    {
        "canonical": "equity_parent",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "patrimonio_neto_dominante",
        "aliases": [
            "patrimonio neto atribuido a la sociedad dominante",
            "patrimonio neto atribuido al accionista dominante",
            "equity attributable to parent",
            "equity attributable to owners of the parent"
        ]
    },

    {
        "canonical": "equity_minorities",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "patrimonio_neto_minoritarios",
        "aliases": [
            "patrimonio neto atribuido a los minoritarios",
            "intereses minoritarios",
            "minority interests equity",
            "minority interests",
            "non controlling interests",
            "non-controlling interests"
        ]
    },

    # =================================================
    # BALANCE - PASIVO NO CORRIENTE
    # =================================================

    {
        "canonical": "non_current_liabilities_total",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "pasivo_no_corriente",
        "aliases": [
            "pasivos no corrientes",
            "pasivo no corriente",
            "non current liabilities",
            "non-current liabilities"
        ]
    },

    {
        "canonical": "provisions_lp",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "provisiones_lp",
        "aliases": [
            "provisiones a largo plazo",
            "provisiones no corrientes",
            "provisiones largo plazo",
            "long term provisions",
            "long-term provisions",
            "non current provisions",
            "non-current provisions"
        ]
    },

    {
        "canonical": "other_liabilities_lp",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "otros_pasivos_lp",
        "aliases": [
            "otros pasivos a largo plazo",
            "otros pasivos lp",
            "other long term liabilities",
            "other long-term liabilities",
            "other non current liabilities",
            "other non-current liabilities"
        ]
    },

    {
        "canonical": "financial_debt_lp",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "deuda_financiera_lp",
        "aliases": [
            "deuda lp",
            "deuda l p",
            "deuda largo plazo",
            "deuda a largo plazo",
            "deuda financiera lp",
            "deuda financiera a largo plazo",
            "prestamos a largo plazo",
            "préstamos a largo plazo",
            "long term debt",
            "long-term debt",
            "financial debt lp",
            "non current debt",
            "non-current debt",
            "non current financial debt",
            "non-current financial debt"
        ]
    },

    {
        "canonical": "lease_liability_lp",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "arrendamientos_lp",
        "aliases": [
            "pasivo por arrendamiento a largo plazo",
            "arrendamientos a largo plazo",
            "lease liability lp",
            "lease liabilities long term",
            "long term lease liability",
            "long-term lease liability",
            "non current lease liabilities",
            "non-current lease liabilities"
        ]
    },

    {
        "canonical": "deferred_tax_liabilities",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "pasivos_impuesto_diferido",
        "aliases": [
            "pasivos por impuesto diferidos",
            "pasivos por impuesto diferido",
            "deferred tax liabilities"
        ]
    },

    # =================================================
    # BALANCE - PASIVO CORRIENTE
    # =================================================

    {
        "canonical": "current_liabilities_total",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "pasivo_corriente",
        "aliases": [
            "pasivos corrientes",
            "pasivo corriente",
            "current liabilities"
        ]
    },

    {
        "canonical": "financial_debt_cp",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "deuda_financiera_cp",
        "aliases": [
            "deuda cp",
            "deuda c p",
            "deuda corto plazo",
            "deuda a corto plazo",
            "deuda financiera cp",
            "deuda financiera a corto plazo",
            "prestamos a corto plazo",
            "préstamos a corto plazo",
            "short term debt",
            "short-term debt",
            "current debt",
            "current financial debt"
        ]
    },

    {
        "canonical": "other_financial_liabilities_cp",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "otros_pasivos_financieros_cp",
        "aliases": [
            "otros pasivos financieros",
            "otros pasivos financieros cp",
            "pasivos financieros corrientes",
            "other financial liabilities",
            "current financial liabilities"
        ]
    },

    {
        "canonical": "provisions_cp",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "provisiones_cp",
        "aliases": [
            "provisiones a corto plazo",
            "provisiones corrientes",
            "provisiones corto plazo",
            "short term provisions",
            "short-term provisions",
            "current provisions"
        ]
    },

    {
        "canonical": "lease_liability_cp",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "arrendamientos_cp",
        "aliases": [
            "pasivo por arrendamiento a corto plazo",
            "arrendamientos a corto plazo",
            "lease liability cp",
            "lease liabilities short term",
            "short term lease liability",
            "short-term lease liability",
            "current lease liabilities"
        ]
    },

    {
        "canonical": "income_tax_payable",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "impuesto_corriente_pasivo",
        "aliases": [
            "pasivos por impuesto sobre beneficios corriente",
            "impuesto corriente pasivo",
            "pasivo por impuesto corriente",
            "current tax liability",
            "income tax payable",
            "tax payable"
        ]
    },

    {
        "canonical": "trade_payables",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "proveedores",
        "aliases": [
            "acreedores",
            "proveedores",
            "acreedores comerciales",
            "cuentas a pagar",
            "accounts payable",
            "trade payables",
            "payables",
            "trade and other payables"
        ]
    },

    # =================================================
    # BALANCE - TOTALES FINALES
    # =================================================

    {
        "canonical": "total_liabilities",
        "statement": "balance",
        "section": "pasivo_total",
        "aggregate": "pasivo_total",
        "aliases": [
            "pasivo total",
            "total pasivo",
            "total liabilities",
            "liabilities total"
        ]
    },

    {
        "canonical": "total_liabilities_and_equity",
        "statement": "balance",
        "section": "pasivo_total_pn",
        "aggregate": "pasivo_total_pn",
        "aliases": [
            "total pasivo y patrimonio neto",
            "total liabilities and equity",
            "total liabilities and shareholders equity",
            "total liabilities and shareholders' equity",
            "total pasivo + pn",
            "pasivo y patrimonio neto total",
            "total pasivo y pn"
        ]
    }
]


# =====================================================
# EXTENSIÓN LINE_MAPPING_RULES - MÁS DETALLE BALANCE
# Pegar justo después de tu LINE_MAPPING_RULES actual
# =====================================================

EXTRA_BALANCE_MAPPING_RULES = [

    # =================================================
    # BALANCE - ACTIVO NO CORRIENTE / MÁS DETALLE
    # =================================================

    {
        "canonical": "investments_in_associates",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "inversiones_en_asociadas",
        "aliases": [
            "inversiones en empresas asociadas",
            "inversiones en asociadas",
            "participaciones en asociadas",
            "participaciones puestas en equivalencia",
            "investments in associates",
            "investments in associated companies",
            "equity accounted investments",
            "investments accounted for using the equity method"
        ]
    },

    {
        "canonical": "loans_receivable_lp",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "prestamos_cobrar_lp",
        "aliases": [
            "prestamos a largo plazo",
            "préstamos a largo plazo",
            "creditos a largo plazo",
            "créditos a largo plazo",
            "prestamos concedidos a largo plazo",
            "long term loans receivable",
            "long-term loans receivable",
            "non current loans receivable",
            "non-current loans receivable"
        ]
    },

    {
        "canonical": "derivative_assets_lp",
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "derivados_activo_lp",
        "aliases": [
            "derivados financieros no corrientes",
            "activos derivados no corrientes",
            "non current derivative assets",
            "non-current derivative assets",
            "derivative financial assets non current",
            "derivative financial assets non-current"
        ]
    },

    # =================================================
    # BALANCE - ACTIVO CORRIENTE / MÁS DETALLE
    # =================================================

    {
        "canonical": "other_receivables",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "otros_deudores",
        "aliases": [
            "otros deudores",
            "otros deudores comerciales",
            "otros creditos a cobrar",
            "otros créditos a cobrar",
            "other receivables",
            "other accounts receivable",
            "other trade receivables",
            "trade and other receivables"
        ]
    },

    {
        "canonical": "prepayments",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "gastos_anticipados",
        "aliases": [
            "periodificaciones a corto plazo",
            "gastos anticipados",
            "pagos anticipados",
            "prepayments",
            "prepaid expenses",
            "prepaid expenses and other current assets"
        ]
    },

    {
        "canonical": "derivative_assets_cp",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "derivados_activo_cp",
        "aliases": [
            "derivados financieros corrientes",
            "activos derivados corrientes",
            "current derivative assets",
            "derivative financial assets current",
            "current financial derivative assets"
        ]
    },

    {
        "canonical": "assets_held_for_sale",
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "activos_mantenidos_venta",
        "aliases": [
            "activos mantenidos para la venta",
            "activos no corrientes mantenidos para la venta",
            "assets held for sale",
            "non current assets held for sale",
            "non-current assets held for sale"
        ]
    },

    # =================================================
    # BALANCE - PATRIMONIO NETO / DETALLE
    # =================================================

    {
        "canonical": "share_capital",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "capital_social",
        "aliases": [
            "capital social",
            "capital",
            "capital escriturado",
            "share capital",
            "issued capital",
            "subscribed capital",
            "ordinary share capital"
        ]
    },

    {
        "canonical": "share_premium",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "prima_emision",
        "aliases": [
            "prima de emision",
            "prima de emisión",
            "share premium",
            "share premium account",
            "additional paid in capital",
            "additional paid-in capital"
        ]
    },

    {
        "canonical": "reserves",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "reservas",
        "aliases": [
            "reservas",
            "reserva legal",
            "otras reservas",
            "reserves",
            "legal reserve",
            "other reserves",
            "capital reserves"
        ]
    },

    {
        "canonical": "retained_earnings",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "resultados_acumulados",
        "aliases": [
            "resultados acumulados",
            "resultados de ejercicios anteriores",
            "remanente",
            "retained earnings",
            "accumulated earnings",
            "accumulated losses",
            "accumulated profit",
            "accumulated deficit"
        ]
    },

    {
        "canonical": "treasury_shares",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "acciones_propias",
        "aliases": [
            "acciones propias",
            "participaciones propias",
            "autocartera",
            "treasury shares",
            "treasury stock",
            "own shares"
        ]
    },

    {
        "canonical": "valuation_adjustments",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "ajustes_valoracion",
        "aliases": [
            "ajustes por valoracion",
            "ajustes por valoración",
            "ajustes de valoracion",
            "ajustes de valoración",
            "valuation adjustments",
            "fair value adjustments",
            "revaluation reserve",
            "other comprehensive income reserve",
            "oci reserve"
        ]
    },

    {
        "canonical": "grants_donations_bequests",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "subvenciones_donaciones_legados",
        "aliases": [
            "subvenciones donaciones y legados",
            "subvenciones, donaciones y legados",
            "subvenciones de capital",
            "grants donations and bequests",
            "grants, donations and bequests",
            "capital grants",
            "government grants"
        ]
    },

    {
        "canonical": "profit_loss_for_period_equity",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "resultado_ejercicio_pn",
        "aliases": [
            "resultado del ejercicio",
            "resultado del periodo",
            "beneficio del ejercicio",
            "perdida del ejercicio",
            "pérdida del ejercicio",
            "profit for the year",
            "profit for the period",
            "loss for the year",
            "loss for the period",
            "net income for the year"
        ]
    },

    {
        "canonical": "other_equity_items",
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "otros_patrimonio_neto",
        "aliases": [
            "otros instrumentos de patrimonio neto",
            "otros componentes del patrimonio neto",
            "otros elementos de patrimonio neto",
            "other equity instruments",
            "other components of equity",
            "other equity items"
        ]
    },

    # =================================================
    # BALANCE - PASIVO NO CORRIENTE / MÁS DETALLE
    # =================================================

    {
        "canonical": "other_financial_liabilities_lp",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "otros_pasivos_financieros_lp",
        "aliases": [
            "otros pasivos financieros no corrientes",
            "otros pasivos financieros a largo plazo",
            "pasivos financieros no corrientes",
            "non current financial liabilities",
            "non-current financial liabilities",
            "other non current financial liabilities",
            "other non-current financial liabilities"
        ]
    },

    {
        "canonical": "derivative_liabilities_lp",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "derivados_pasivo_lp",
        "aliases": [
            "derivados financieros no corrientes",
            "pasivos derivados no corrientes",
            "non current derivative liabilities",
            "non-current derivative liabilities",
            "derivative financial liabilities non current",
            "derivative financial liabilities non-current"
        ]
    },

    {
        "canonical": "pension_obligations_lp",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "obligaciones_pensiones_lp",
        "aliases": [
            "obligaciones por pensiones",
            "compromisos por pensiones",
            "obligaciones por prestaciones a empleados",
            "employee benefit obligations",
            "pension obligations",
            "retirement benefit obligations",
            "post employment benefit obligations"
        ]
    },

    {
        "canonical": "deferred_income_lp",
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "ingresos_diferidos_lp",
        "aliases": [
            "ingresos diferidos no corrientes",
            "ingresos diferidos a largo plazo",
            "deferred income non current",
            "deferred income non-current",
            "non current deferred income",
            "non-current deferred income"
        ]
    },

    # =================================================
    # BALANCE - PASIVO CORRIENTE / MÁS DETALLE
    # =================================================

    {
        "canonical": "other_current_liabilities",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "otros_pasivos_corrientes",
        "aliases": [
            "otros pasivos corrientes",
            "otros acreedores",
            "otras cuentas a pagar",
            "other current liabilities",
            "other payables",
            "other accounts payable",
            "other liabilities"
        ]
    },

    {
        "canonical": "accruals",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "periodificaciones_pasivo",
        "aliases": [
            "periodificaciones a corto plazo",
            "gastos devengados",
            "accruals",
            "accrued expenses",
            "accrued liabilities",
            "accrued expenses and other liabilities"
        ]
    },

    {
        "canonical": "contract_liabilities",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "pasivos_contrato",
        "aliases": [
            "pasivos por contratos",
            "anticipos de clientes",
            "ingresos anticipados",
            "contract liabilities",
            "customer advances",
            "advances from customers",
            "deferred revenue",
            "deferred income current"
        ]
    },

    {
        "canonical": "derivative_liabilities_cp",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "derivados_pasivo_cp",
        "aliases": [
            "derivados financieros corrientes",
            "pasivos derivados corrientes",
            "current derivative liabilities",
            "derivative financial liabilities current",
            "current financial derivative liabilities"
        ]
    },

    {
        "canonical": "payroll_social_security_payable",
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "personal_ss_acreedores",
        "aliases": [
            "remuneraciones pendientes de pago",
            "personal acreedor",
            "seguridad social acreedora",
            "payroll payable",
            "salaries payable",
            "social security payable",
            "employee benefits payable"
        ]
    }
]

# Añadir extensión al diccionario principal
LINE_MAPPING_RULES.extend(EXTRA_BALANCE_MAPPING_RULES)


# =====================================================
# AJUSTES FINALES DICCIONARIO / REGLAS DE NEGOCIO
#
# Pegar justo después de:
# LINE_MAPPING_RULES.extend(EXTRA_BALANCE_MAPPING_RULES)
# =====================================================


# =====================================================
# 1) CANONICALS DE PYG QUE SOLO PUEDEN SER TOTALES
# =====================================================

PYG_TOTAL_ONLY_CANONICALS = {
    "gross_profit",
    "gross_margin_pct",
    "ebitda",
    "ebit",
    "profit_before_tax",
    "net_income",
    "net_income_parent",
    "minority_income",
}


# =====================================================
# 2) ELIMINAR EBITA COMO CANONICAL SEPARADO
#
# Decisión:
# - EBITA se unifica dentro de EBITDA.
# - Ya no existirá canonical = "ebita".
# =====================================================

LINE_MAPPING_RULES = [
    rule for rule in LINE_MAPPING_RULES
    if rule.get("canonical") != "ebita"
]


# =====================================================
# 3) HELPER PARA AÑADIR ALIASES SIN DUPLICAR
# =====================================================

def add_aliases_to_rule(canonical, new_aliases):
    found = False

    for rule in LINE_MAPPING_RULES:
        if rule.get("canonical") == canonical:
            found = True

            existing = set(
                str(a).strip().lower()
                for a in rule.get("aliases", [])
            )

            for alias in new_aliases:
                alias_norm = str(alias).strip().lower()

                if alias_norm not in existing:
                    rule["aliases"].append(alias)
                    existing.add(alias_norm)

    if not found:
        print(f"[WARNING] No se encontró canonical para añadir aliases: {canonical}")


# =====================================================
# 4) UNIFICAR EBITA DENTRO DE EBITDA
# =====================================================

add_aliases_to_rule("ebitda", [
    "ebita",
    "adjusted ebita",
    "normalized ebita",
    "recurring ebita",
    "earnings before interest taxes and amortization",
    "earnings before interest, taxes and amortization",
    "earnings before interest tax and amortization",
    "earnings before interest, tax and amortization",
])


# =====================================================
# 5) AÑADIR ALIASES DE PERSONAL
#
# Para que "Staff wages" pueda mapearse a personnel_expenses.
# =====================================================

add_aliases_to_rule("personnel_expenses", [
    # Español extra
    "sueldos",
    "sueldos salarios",
    "sueldos y cargas sociales",
    "coste salarial",
    "costes salariales",
    "gasto salarial",
    "gastos salariales",

    # Inglés extra
    "staff wages",
    "staff salaries",
    "staff expenses",
    "personnel costs",
    "employee wages",
    "employee salaries",
    "employee costs",
    "wages",
    "salaries",
    "wages salaries",
    "wages and salaries",
    "payroll costs",
    "payroll expenses",
])


# =====================================================
# 6) AÑADIR OTHER OPERATING INCOME
#
# Para evitar que caiga mal como EBIT.
# =====================================================

add_aliases_to_rule("other_gains_losses", [
    # Español extra
    "otros ingresos de explotacion",
    "otros ingresos de explotación",
    "otros ingresos operativos",
    "otros ingresos y gastos de explotacion",
    "otros ingresos y gastos de explotación",

    # Inglés extra
    "other operating income",
    "other operating revenue",
    "other operating gains",
    "other operating gains and losses",
    "other income",
])


# =====================================================
# 7) REGLA TOTAL-ONLY PARA PYG
#
# Regla:
# - Solo aplica si statement = pyg.
# - Si el canonical candidato es de tipo total/resultado,
#   la línea debe cumplir _is_total(raw_concept).
#
# Ejemplos:
# - "Other operating income" NO puede ser EBIT.
# - "Operating profit" SÍ puede ser EBIT.
# - "EBITDA" SÍ puede ser EBITDA.
# - "Staff wages" no se ve afectado, porque no es total-only.
# =====================================================

def is_pyg_total_only_mapping_allowed(raw_concept, statement_type, candidate_canonical):
    statement_norm = normalize_text(statement_type)

    # En balance NO aplicamos esta regla
    if statement_norm != "pyg":
        return True

    # Si no es un canonical de total, se permite
    if candidate_canonical not in PYG_TOTAL_ONLY_CANONICALS:
        return True

    # Si es canonical de total, la línea debe parecer total/resultado
    return _is_total(raw_concept)

# =====================================================
# ENCABEZADOS DE SECCION
# Solo estas filas actualizan el contexto
# Usa niveles para evitar falsos positivos
# =====================================================

SECTION_HEADERS = {
    "activo_no_corriente": [
        "activos no corrientes",
        "activo no corriente",
        "non current assets",
        "non-current assets",
        "fixed assets"
    ],
    "activo_corriente": [
        "activos corrientes",
        "activo corriente",
        "current assets"
    ],
    "patrimonio_neto": [
        "patrimonio neto",
        "fondos propios",
        "equity",
        "shareholders equity"
    ],
    "pasivo_no_corriente": [
        "pasivos no corrientes",
        "pasivo no corriente",
        "non current liabilities",
        "non-current liabilities"
    ],
    "pasivo_corriente": [
        "pasivos corrientes",
        "pasivo corriente",
        "current liabilities"
    ]
}


def detect_section_from_header(text, level=None):
    """
    Detecta si una fila es cabecera real de sección.

    Importante:
    - Solo permite actualizar contexto si la fila es de nivel alto.
    - Evita que partidas como 'otros activos corrientes' cambien la sección.
    """

    if text is None:
        return None

    t = normalize_text(text)

    if not t:
        return None

    # -------------------------------------------------
    # 1. Control por nivel
    # Solo cabeceras de nivel alto pueden cambiar sección
    # -------------------------------------------------
    if level is not None:
        try:
            level_num = int(level)
            if level_num > 1:
                return None
        except:
            pass

    # -------------------------------------------------
    # 2. Evitar textos demasiado largos
    # Una cabecera suele ser corta
    # -------------------------------------------------
    words = t.split()
    if len(words) > 6:
        return None

    # -------------------------------------------------
    # 3. Match controlado
    # Primero exacto, luego flexible muy limitado
    # -------------------------------------------------
    for section, aliases in SECTION_HEADERS.items():
        for alias in aliases:
            alias_norm = normalize_text(alias)

            if t == alias_norm:
                return section

            # Permite casos tipo:
            # "total activo corriente"
            # "activo corriente total"
            if alias_norm in t:
                extra_words = set(t.split()) - set(alias_norm.split())

                allowed_extra_words = {
                    "total",
                    "totales",
                    "nota",
                    "notas",
                    "eur",
                    "euros",
                    "miles",
                    "mileseur",
                    "grupo"
                }

                if extra_words.issubset(allowed_extra_words):
                    return section

    return None

# =====================================================
# REGLAS FUERTES
# Sirven para priorizar conceptos que se entienden solos
# aunque no exista cabecera explícita
# =====================================================

def is_strong_match(raw_concept, alias):
    """
    Una coincidencia fuerte es:
    - exacta
    - o por tokens completos incluidos
    """
    if is_exact_match(raw_concept, alias):
        return True
    if is_token_subset_match(raw_concept, alias):
        return True
    return False

# =====================================================
# BLOQUE 3 - CONCEPTOS AMBIGUOS
# Solo estas líneas deberían apoyarse mucho en el contexto
# =====================================================

AMBIGUOUS_TERMS = [
    "otros activos",
    "otros pasivos",
    "otros activos corrientes",
    "otros activos no corrientes",
    "otros pasivos financieros",
    "otros pasivos lp",
    "otras cuentas",
    "periodificaciones",
    "ajustes",
    "administraciones publicas",
    "inversiones financieras"
]

def is_ambiguous_concept(text):
    t = normalize_text(text)
    for term in AMBIGUOUS_TERMS:
        if term in t:
            return True
    return False

# =====================================================
# BLOQUE EXTRA - OVERRIDES CRÍTICOS PyG / BALANCE
# Blindar conceptos que no deben competir con reglas genéricas
#
# Reglas:
# - BALANCE: solo aplica si level == 1
# - PyG: solo aplica si level está vacío / None / NaN
# =====================================================

import pandas as pd

CRITICAL_LABEL_OVERRIDES = {
    # =================================================
    # PyG - Totales / subtotales sin nivel
    # =================================================

    "ebitda": {
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "ebitda",
        "aliases": {
            # Español
            "ebitda",
            "resultado bruto de explotacion",
            "resultado bruto de explotación",
            "resultado operativo ebitda",

            # Inglés
            "adjusted ebitda",
            "normalized ebitda",
            "recurring ebitda",
            "earnings before interest taxes depreciation amortization",
            "earnings before interest, taxes, depreciation and amortization",
            "earnings before interest tax depreciation amortization",
            "earnings before interest, tax, depreciation and amortization",
            "earnings before interest taxes depreciation and amortization"
        }
    },

    "ebita": {
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "ebita",
        "aliases": {
            # Inglés
            "ebita",
            "adjusted ebita",
            "normalized ebita",
            "recurring ebita",
            "earnings before interest taxes and amortization",
            "earnings before interest, taxes and amortization",
            "earnings before interest tax and amortization",
            "earnings before interest, tax and amortization"
        }
    },

    "ebit": {
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "ebit",
        "aliases": {
            # Español
            "ebit",
            "resultado de explotacion",
            "resultado de explotación",
            "resultados de explotacion",
            "resultados de explotación",
            "resultado explotacion",
            "resultado explotación",
            "resultado operativo",
            "beneficio de explotacion",
            "beneficio de explotación",
            "beneficio operativo",

            # Inglés
            "operating income",
            "operating profit",
            "operating result",
            "profit from operations",
            "income from operations",
            "earnings before interest and taxes",
            "earnings before interest taxes",
            "earnings before interest and tax",
            "earnings before interest tax"
        }
    },

    "gross_profit": {
        "statement": "pyg",
        "section": "operativo",
        "aggregate": "margen_bruto",
        "aliases": {
            # Español
            "margen bruto",
            "resultado bruto",
            "beneficio bruto",

            # Inglés
            "gross profit",
            "gross margin",
            "gross result"
        }
    },

    "financial_result": {
        "statement": "pyg",
        "section": "financiero",
        "aggregate": "resultado_financiero",
        "aliases": {
            # Español
            "resultado financiero",
            "resultados financieros",
            "resultado neto financiero",

            # Inglés
            "financial result",
            "finance result",
            "net financial result"
        }
    },

    "profit_before_tax": {
        "statement": "pyg",
        "section": "fiscal",
        "aggregate": "bai",
        "aliases": {
            # Español
            "beneficio antes de impuestos",
            "beneficio antes de impuesto",
            "resultado antes de impuestos",
            "resultado antes de impuesto",
            "resultado antes de impuestos sobre beneficios",
            "resultado antes de impuesto sobre beneficios",
            "resultado antes de impuestos de actividades continuadas",
            "resultado antes de impuesto de actividades continuadas",
            "bai",

            # Inglés - profit
            "profit before tax",
            "profit before taxes",
            "profit before income tax",
            "profit before income taxes",
            "profit before taxation",
            "pre tax profit",
            "pre-tax profit",
            "pbt",

            # Inglés - income
            "income before tax",
            "income before taxes",
            "income before income tax",
            "income before income taxes",
            "income before taxation",
            "pre tax income",
            "pre-tax income",

            # Inglés - earnings
            "earnings before tax",
            "earnings before taxes",
            "earnings before income tax",
            "earnings before income taxes",
            "earnings before taxation",
            "pre tax earnings",
            "pre-tax earnings",
            "ebt",
            "earnings before taxes ebt"
        }
    },

    "income_tax": {
        "statement": "pyg",
        "section": "fiscal",
        "aggregate": "tax",
        "aliases": {
            # Español
            "impuesto sobre beneficios",
            "impuestos sobre beneficios",
            "gasto por impuesto",
            "gasto por impuestos",
            "gasto por impuesto sobre beneficios",
            "impuesto corriente",
            "impuesto diferido",

            # Inglés
            "income tax",
            "income taxes",
            "income tax expense",
            "tax expense",
            "current tax",
            "deferred tax",
            "current income tax",
            "deferred income tax"
        }
    },

    "net_income": {
        "statement": "pyg",
        "section": "neto",
        "aggregate": "beneficio_neto",
        "aliases": {
            # Español
            "resultado neto del ejercicio",
            "resultado neto",
            "resultado neto consolidado",
            "resultado consolidado",
            "beneficio neto",
            "beneficio neto del ejercicio",
            "beneficio consolidado",
            "perdida neta",
            "pérdida neta",
            "perdida del ejercicio",
            "pérdida del ejercicio",
            "resultado despues de impuestos",
            "resultado después de impuestos",
            "beneficio despues de impuestos",
            "beneficio después de impuestos",

            # Inglés - net income / profit
            "net income",
            "consolidated net income",
            "net income on continuing operations",
            "net income on discontinuing operations",
            "net income from continuing operations",
            "net income from discontinued operations",
            "net profit",
            "consolidated net profit",
            "net loss",
            "consolidated net loss",

            # Inglés - profit/loss for period
            "profit for the year",
            "profit for the period",
            "loss for the year",
            "loss for the period",
            "profit after tax",
            "profit after taxes",
            "profit after income tax",
            "profit after income taxes",
            "income after tax",
            "income after taxes",
            "income after income tax",
            "income after income taxes",

            # Inglés - earnings
            "net earnings",
            "consolidated net earnings",
            "earnings after tax",
            "earnings after taxes",
            "earnings after income tax",
            "earnings after income taxes",
            "earnings for the year",
            "earnings for the period",
            "profit attributable to shareholders",
            "earnings attributable to shareholders"
        }
    },

    "net_income_parent": {
        "statement": "pyg",
        "section": "neto",
        "aggregate": "beneficio_neto_dominante",
        "aliases": {
            # Español
            "resultado neto atribuido a la sociedad dominante",
            "resultado atribuido a la dominante",
            "resultado atribuido a los propietarios de la dominante",
            "beneficio atribuido a la dominante",

            # Inglés
            "net income attributable to parent",
            "net income attributable to owners of the parent",
            "profit attributable to parent",
            "profit attributable to owners of the parent",
            "earnings attributable to parent",
            "earnings attributable to owners of the parent"
        }
    },

    "minority_income": {
        "statement": "pyg",
        "section": "neto",
        "aggregate": "resultado_minoritarios",
        "aliases": {
            # Español
            "resultado atribuido a accionistas minoritarios",
            "resultado atribuido a minoritarios",
            "resultado atribuido a intereses minoritarios",

            # Inglés
            "minority interests result",
            "minority interest",
            "profit attributable to non controlling interests",
            "profit attributable to non-controlling interests",
            "net income attributable to non controlling interests",
            "net income attributable to non-controlling interests"
        }
    },

    # =================================================
    # Balance - Solo nivel 1
    # =================================================

    "non_current_assets_total": {
        "statement": "balance",
        "section": "activo_no_corriente",
        "aggregate": "activo_no_corriente",
        "aliases": {
            "activo no corriente",
            "activos no corrientes",
            "non current assets",
            "non-current assets",
            "fixed assets total"
        }
    },

    "current_assets_total": {
        "statement": "balance",
        "section": "activo_corriente",
        "aggregate": "activo_corriente",
        "aliases": {
            "activo corriente",
            "activos corrientes",
            "current assets"
        }
    },

    "total_assets": {
        "statement": "balance",
        "section": "activo_total",
        "aggregate": "activo_total",
        "aliases": {
            "total activo",
            "activo total",
            "total assets",
            "assets total"
        }
    },

    "equity_total": {
        "statement": "balance",
        "section": "patrimonio_neto",
        "aggregate": "patrimonio_neto",
        "aliases": {
            "patrimonio neto",
            "fondos propios",
            "equity",
            "shareholders equity",
            "shareholders' equity",
            "total equity",
            "net equity"
        }
    },

    "non_current_liabilities_total": {
        "statement": "balance",
        "section": "pasivo_no_corriente",
        "aggregate": "pasivo_no_corriente",
        "aliases": {
            "pasivo no corriente",
            "pasivos no corrientes",
            "non current liabilities",
            "non-current liabilities"
        }
    },

    "current_liabilities_total": {
        "statement": "balance",
        "section": "pasivo_corriente",
        "aggregate": "pasivo_corriente",
        "aliases": {
            "pasivo corriente",
            "pasivos corrientes",
            "current liabilities"
        }
    },

    "total_liabilities": {
        "statement": "balance",
        "section": "pasivo_total",
        "aggregate": "pasivo_total",
        "aliases": {
            "total pasivo",
            "pasivo total",
            "total liabilities",
            "liabilities total"
        }
    },

    "total_liabilities_and_equity": {
        "statement": "balance",
        "section": "pasivo_total_pn",
        "aggregate": "pasivo_total_pn",
        "aliases": {
            "total pasivo y patrimonio neto",
            "pasivo y patrimonio neto total",
            "total liabilities and equity",
            "total liabilities and shareholders equity",
            "total liabilities and shareholders' equity",
            "total pasivo y pn"
        }
    }
}


def is_empty_level(level):
    """
    Detecta si el nivel está vacío.
    Lo usamos para PyG, donde los totales suelen venir sin nivel.
    """
    if level is None:
        return True

    try:
        if pd.isna(level):
            return True
    except Exception:
        pass

    if str(level).strip() in ["", "<NA>", "None", "nan", "NaN"]:
        return True

    return False


def is_level_1(level):
    """
    Detecta si una fila es nivel 1.
    Lo usamos para Balance.
    """
    if is_empty_level(level):
        return False

    try:
        return int(float(level)) == 1
    except Exception:
        return False


def can_apply_critical_override(statement, level):
    """
    Reglas de aplicación:
    - Balance: solo nivel 1
    - PyG: solo sin nivel
    """
    statement = normalize_text(statement)

    if statement == "balance":
        return is_level_1(level)

    if statement == "pyg":
        return is_empty_level(level)

    return False


def detect_critical_label_override(text, statement, level=None):
    """
    Detecta etiquetas críticas con bajo riesgo de falsos positivos.

    Devuelve dict con canonical/section/aggregate si encuentra match.
    Si no, devuelve None.

    Lógica:
    1) Match exacto.
    2) Match controlado:
       permite alias + palabras neutras habituales.
    3) Match fuerte de frases críticas:
       permite detectar casos tipo:
       - Consolidated net income
       - Net income on continuing operations
       - Earnings before taxes
       - Resultado neto consolidado
    """

    if text is None:
        return None

    if not can_apply_critical_override(statement, level):
        return None

    t = normalize_text(text)

    if not t:
        return None

    statement_norm = normalize_text(statement)

    allowed_extra_words = {
        # genéricas
        "total",
        "totales",
        "subtotal",
        "nota",
        "notas",
        "eur",
        "euros",
        "miles",
        "thousand",
        "thousands",
        "million",
        "millions",

        # consolidación / grupo
        "grupo",
        "group",
        "consolidado",
        "consolidada",
        "consolidated",

        # periodo
        "ejercicio",
        "periodo",
        "período",
        "year",
        "period",
        "annual",

        # operaciones continuadas/discontinuadas
        "continuing",
        "continued",
        "discontinuing",
        "discontinued",
        "operations",
        "operation",
        "actividades",
        "continuadas",
        "discontinuadas",
        "interrumpidas",

        # atribución
        "attributable",
        "attributed",
        "owners",
        "owner",
        "parent",
        "shareholders",
        "shareholder",
        "dominante",
        "propietarios",
        "sociedad",

        # conectores frecuentes
        "on",
        "from",
        "for",
        "of",
        "to",
        "the",
        "and",
        "de",
        "del",
        "la",
        "el",
        "los",
        "las",
        "a",
        "y"
    }

    for canonical, rule in CRITICAL_LABEL_OVERRIDES.items():

        if normalize_text(rule["statement"]) != statement_norm:
            continue

        for alias in rule["aliases"]:
            alias_norm = normalize_text(alias)

            # 1) Match exacto: máxima seguridad
            if t == alias_norm:
                return {
                    "canonical": canonical,
                    "statement": rule["statement"],
                    "section": rule["section"],
                    "aggregate": rule["aggregate"],
                    "match_type": "critical_exact",
                    "matched_alias": alias
                }

            # 2) Match controlado:
            # Solo permite alias + palabras neutras
            text_words = set(t.split())
            alias_words = set(alias_norm.split())

            if alias_words.issubset(text_words):
                extra_words = text_words - alias_words

                if extra_words.issubset(allowed_extra_words):
                    return {
                        "canonical": canonical,
                        "statement": rule["statement"],
                        "section": rule["section"],
                        "aggregate": rule["aggregate"],
                        "match_type": "critical_controlled",
                        "matched_alias": alias
                    }

    return None


# Compatibilidad con tu función antigua
def is_profit_before_tax_label(text, level=None):
    override = detect_critical_label_override(
        text=text,
        statement="pyg",
        level=level
    )
    return override is not None and override["canonical"] == "profit_before_tax"

# =====================================================
# BLOQUE 4 - MAPPING PRINCIPAL CON CONTEXTO
#
# Incluye:
# 1) detección de headers con level
# 2) overrides críticos con statement + level
# 3) mapping normal
# 4) contexto de sección
# 5) validación activo/pasivo
# 6) coherencia padre-hijo:
#    - si el hijo NO tiene canonical, hereda canonical/aggregate/section del padre
#    - si el hijo SÍ tiene canonical, mantiene canonical pero hereda/alinea sección del padre
# 7) soporte multi-periodo:
#    - amount      = último año / periodo más reciente
#    - amount_1    = año anterior
#    - amount_2    = dos años antes
#    - amount_year, amount_1_year, amount_2_year...
#
# Importante:
# - Usa parent_index / parent_name generado en el bloque de niveles.
# - Funciona aunque el padre esté arriba o debajo.
# =====================================================

import re
import pandas as pd


def safe_get(row, col_name, default=None):
    """
    Lee columna por nombre si existe.
    """
    try:
        if col_name in row.index:
            return row[col_name]
    except Exception:
        pass
    return default


def _is_valid_parent_index(x, max_len):
    """
    Valida parent_index aunque venga como int, float, str, pd.NA, etc.
    """
    if x is None:
        return False

    try:
        if pd.isna(x):
            return False
    except Exception:
        pass

    try:
        idx = int(float(x))
    except Exception:
        return False

    return 0 <= idx < max_len


def _to_int_index(x):
    try:
        return int(float(x))
    except Exception:
        return None


def _build_amounts_by_period_from_row(row):
    """
    Construye un diccionario con todos los periodos detectados.

    Entrada esperada:
    - amount, amount_year, amount_raw, amount_source_col
    - amount_1, amount_1_year, amount_1_raw, amount_1_source_col
    - amount_2, amount_2_year, amount_2_raw, amount_2_source_col
    - etc.

    Salida:
    {
        "amount": {
            "value": ...,
            "year": ...,
            "raw": ...,
            "source_col": ...
        },
        "amount_1": {...},
        "amount_2": {...}
    }
    """

    amounts_by_period = {}

    # -------------------------------------------------
    # Periodo actual / más reciente
    # -------------------------------------------------
    if "amount" in row.index:
        amounts_by_period["amount"] = {
            "value": safe_get(row, "amount", None),
            "year": safe_get(row, "amount_year", None),
            "raw": safe_get(row, "amount_raw", None),
            "source_col": safe_get(row, "amount_source_col", None),
        }

    # -------------------------------------------------
    # Periodos anteriores: amount_1, amount_2, amount_3...
    # -------------------------------------------------
    amount_indexes = []

    for col in row.index:
        m = re.fullmatch(r"amount_(\d+)", str(col))
        if m:
            amount_indexes.append(int(m.group(1)))

    for idx_amount in sorted(set(amount_indexes)):
        field = f"amount_{idx_amount}"

        amounts_by_period[field] = {
            "value": safe_get(row, field, None),
            "year": safe_get(row, f"{field}_year", None),
            "raw": safe_get(row, f"{field}_raw", None),
            "source_col": safe_get(row, f"{field}_source_col", None),
        }

    return amounts_by_period


def map_lines_with_context(df, statement_type):
    results = []
    statement_norm = normalize_text(statement_type)

    # Reset index para que parent_index coincida con posición real
    df_work = df.copy().reset_index(drop=True)

    current_section = None

    # =====================================================
    # HELPER INTERNO: BUSCAR MEJOR MAPPING DE UNA LÍNEA
    # =====================================================

    def find_best_mapping(raw_concept, level, current_section):
        raw_norm = normalize_text(raw_concept)

        # -------------------------------------------------
        # 1. Overrides críticos
        # -------------------------------------------------
        override = detect_critical_label_override(
            text=raw_concept,
            statement=statement_type,
            level=level
        )

        best_rule = None
        best_score = -1
        best_match_type = None
        matched_alias = None

        if override:
            best_rule = {
                "canonical": override["canonical"],
                "statement": override["statement"],
                "section": override["section"],
                "aggregate": override["aggregate"],
                "aliases": []
            }
            best_score = 9999
            best_match_type = override.get("match_type")
            matched_alias = override.get("matched_alias")

        else:
            # -------------------------------------------------
            # 2. Mapping normal
            # -------------------------------------------------
            for rule in LINE_MAPPING_RULES:
                if normalize_text(rule["statement"]) != statement_norm:
                    continue

                # =================================================
                # BLOQUEO DE TOTALES SOLO PARA PYG
                #
                # Ejemplo:
                # "Other operating income" no puede ser EBIT
                # porque no es una línea de total/resultado.
                # =================================================
                if not is_pyg_total_only_mapping_allowed(
                    raw_concept=raw_concept,
                    statement_type=statement_type,
                    candidate_canonical=rule["canonical"]
                ):
                    continue

                # Blindaje fiscal:
                # evita que "resultado antes de impuestos"
                # caiga como income_tax.
                if statement_norm == "pyg" and rule["canonical"] == "income_tax":
                    if is_profit_before_tax_label(raw_concept, level=level):
                        continue

                for alias in rule["aliases"]:
                    score = -1
                    alias_len = len(token_set(alias))
                    match_type = None

                    if is_exact_match(raw_concept, alias):
                        score = 1000 + alias_len
                        match_type = "exact"

                    elif is_token_subset_match(raw_concept, alias):

                        # Endurecer fiscal
                        if statement_norm == "pyg" and rule["canonical"] in (
                            "profit_before_tax",
                            "income_tax"
                        ):
                            if alias_len < 3:
                                score = -1
                                match_type = None
                            else:
                                score = 100 + alias_len
                                match_type = "subset"
                        else:
                            score = 100 + alias_len
                            match_type = "subset"

                    if score > 0:

                        # Bonus si coincide con sección actual en balance
                        if statement_norm == "balance" and current_section is not None:
                            if rule["section"] == current_section:
                                score += 25

                            # Penalización fuerte si contradice CP/LP
                            if (
                                current_section.endswith("_corriente")
                                and rule["section"].endswith("_no_corriente")
                            ):
                                score -= 200

                            if (
                                current_section.endswith("_no_corriente")
                                and rule["section"].endswith("_corriente")
                            ):
                                score -= 200

                        if score > best_score:
                            best_score = score
                            best_rule = rule
                            best_match_type = match_type
                            matched_alias = alias

        # -------------------------------------------------
        # 3. Section detected
        # -------------------------------------------------
        if best_rule:
            detected_section = best_rule["section"]
        else:
            if statement_norm == "balance" and is_ambiguous_concept(raw_concept):
                detected_section = current_section
            else:
                detected_section = None

        # -------------------------------------------------
        # 4. Validación simple activo / pasivo
        # -------------------------------------------------
        if detected_section is not None:
            t = raw_norm

            looks_like_asset = any(x in t for x in [
                "caja", "tesoreria", "tesorería", "efectivo",
                "clientes", "deudores", "existencias",
                "inventario", "inmovilizado", "intangibles",
                "goodwill", "activo", "assets", "receivables",
                "inventory", "inventories", "cash"
            ])

            looks_like_liability = any(x in t for x in [
                "deuda", "proveedores", "acreedores",
                "pasivo", "provisiones", "arrendamiento",
                "liabilities", "liability", "debt", "payables",
                "provisions"
            ])

            if looks_like_asset and detected_section.startswith("pasivo"):
                detected_section = best_rule["section"] if best_rule else None

            if looks_like_liability and detected_section.startswith("activo"):
                detected_section = best_rule["section"] if best_rule else None

        return {
            "best_rule": best_rule,
            "best_score": best_score,
            "best_match_type": best_match_type,
            "matched_alias": matched_alias,
            "detected_section": detected_section
        }

    # =====================================================
    # PRIMERA PASADA
    # Mapping independiente de cada línea
    # =====================================================

    for pos, row in df_work.iterrows():

        raw_concept = safe_get(row, "label", None)
        level = safe_get(row, "level", None)

        # -------------------------------------------------
        # NUEVO:
        # soporte para varios años / importes
        # -------------------------------------------------
        current_value = safe_get(row, "amount", None)
        current_year = safe_get(row, "amount_year", None)

        previous_value = safe_get(row, "amount_1", None)
        previous_year = safe_get(row, "amount_1_year", None)

        amounts_by_period = _build_amounts_by_period_from_row(row)

        raw_norm = normalize_text(raw_concept)

        # -------------------------------------------------
        # 1. Detectar header de sección
        # -------------------------------------------------
        detected_header = detect_section_from_header(raw_concept, level=level)

        # -------------------------------------------------
        # 2. Buscar mapping igualmente
        # Aunque sea header, queremos canonical si existe.
        # Ejemplo:
        # - Total equity -> equity_total
        # - Current assets -> current_assets_total
        # - Non-current liabilities -> non_current_liabilities_total
        # -------------------------------------------------
        mapping_info = find_best_mapping(
            raw_concept=raw_concept,
            level=level,
            current_section=current_section
        )

        best_rule = mapping_info["best_rule"]
        best_match_type = mapping_info["best_match_type"]
        matched_alias = mapping_info["matched_alias"]
        detected_section_from_mapping = mapping_info["detected_section"]

        original_canonical = best_rule["canonical"] if best_rule else None
        original_aggregate = best_rule["aggregate"] if best_rule else None
        original_section_expected = best_rule["section"] if best_rule else None

        # Si es header, la sección detectada por header manda como contexto.
        if detected_header:
            current_section = detected_header
            final_section_detected = detected_header
            is_header = True
            match_type = "section_header" if best_rule is None else f"section_header+{best_match_type}"
        else:
            final_section_detected = detected_section_from_mapping
            is_header = False
            match_type = best_match_type

        result_row = {
            "raw_concept": raw_concept,
            "normalized_concept": raw_norm,
            "statement": statement_type,
            "level": level,

            "section_detected": final_section_detected,
            "is_header": is_header,

            "canonical": original_canonical,
            "aggregate": original_aggregate,
            "section_expected": original_section_expected,

            "original_canonical": original_canonical,
            "original_aggregate": original_aggregate,
            "original_section_expected": original_section_expected,

            "parent_level": None,
            "parent_raw_concept": None,
            "parent_canonical": None,
            "parent_aggregate": None,
            "parent_section": None,

            # True solo si heredamos canonical/aggregate del padre.
            "mapping_overridden_by_parent": False,

            # True si mantenemos canonical del hijo pero alineamos section al padre.
            "section_aligned_to_parent": False,

            "match_type": match_type,
            "matched_alias": matched_alias,

            # -------------------------------------------------
            # Compatibilidad con bloques antiguos
            # -------------------------------------------------
            "current": current_value,
            "previous": previous_value,

            # -------------------------------------------------
            # NUEVO:
            # años asociados a current / previous
            # -------------------------------------------------
            "current_year": current_year,
            "previous_year": previous_year,

            # -------------------------------------------------
            # NUEVO:
            # mantener nombres originales del dataframe estructurado
            # -------------------------------------------------
            "amount": current_value,
            "amount_year": current_year,
            "amount_1": previous_value,
            "amount_1_year": previous_year,

            # -------------------------------------------------
            # NUEVO:
            # todos los periodos disponibles
            # -------------------------------------------------
            "amounts_by_period": amounts_by_period,
        }

        # -------------------------------------------------
        # NUEVO:
        # Añadir dinámicamente amount_2, amount_2_year, etc.
        # para que no se pierdan en el output plano.
        # -------------------------------------------------
        for period_key, period_info in amounts_by_period.items():
            if period_key == "amount":
                continue

            result_row[period_key] = period_info.get("value")
            result_row[f"{period_key}_year"] = period_info.get("year")
            result_row[f"{period_key}_raw"] = period_info.get("raw")
            result_row[f"{period_key}_source_col"] = period_info.get("source_col")

        results.append(result_row)

    # =====================================================
    # SEGUNDA PASADA
    # Aplicar contexto padre-hijo usando parent_index real
    # =====================================================

    for pos, row in df_work.iterrows():

        result_row = results[pos]

        parent_info = None
        parent_idx_raw = safe_get(row, "parent_index", None)

        # -------------------------------------------------
        # 1. Buscar padre por parent_index
        # -------------------------------------------------
        if _is_valid_parent_index(parent_idx_raw, len(results)):
            parent_idx = _to_int_index(parent_idx_raw)
            parent_info = results[parent_idx]

        # -------------------------------------------------
        # 2. Fallback por parent_name si parent_index no sirve
        # -------------------------------------------------
        if parent_info is None:
            parent_name = safe_get(row, "parent_name", None)

            if parent_name is not None:
                try:
                    if not pd.isna(parent_name):
                        parent_name_norm = normalize_text(parent_name)

                        if parent_name_norm:
                            for candidate in results:
                                if normalize_text(candidate.get("raw_concept")) == parent_name_norm:
                                    parent_info = candidate
                                    break
                except Exception:
                    pass

        # -------------------------------------------------
        # 3. Aplicar coherencia padre-hijo
        # -------------------------------------------------

        if parent_info is not None:
            parent_canonical = parent_info.get("canonical")
            parent_aggregate = parent_info.get("aggregate")
            parent_section = (
                parent_info.get("section_expected")
                or parent_info.get("section_detected")
            )

            result_row["parent_level"] = parent_info.get("level")
            result_row["parent_raw_concept"] = parent_info.get("raw_concept")
            result_row["parent_canonical"] = parent_canonical
            result_row["parent_aggregate"] = parent_aggregate
            result_row["parent_section"] = parent_section

            child_canonical = result_row.get("canonical")
            child_section_expected = result_row.get("section_expected")

            # -------------------------------------------------
            # CASO A:
            # El hijo NO tiene canonical.
            # Entonces sí hereda canonical / aggregate / section del padre.
            # -------------------------------------------------
            if child_canonical is None and parent_canonical is not None:
                result_row["canonical"] = parent_canonical
                result_row["aggregate"] = parent_aggregate
                result_row["section_expected"] = parent_section
                result_row["section_detected"] = parent_section
                result_row["mapping_overridden_by_parent"] = True
                result_row["section_aligned_to_parent"] = True
                result_row["match_type"] = (
                    str(result_row.get("match_type"))
                    + "|inherited_from_parent_no_child_canonical"
                )

            # -------------------------------------------------
            # CASO B:
            # El hijo SÍ tiene canonical.
            # Mantiene canonical y aggregate propios.
            # Pero alinea section_detected al padre.
            #
            # Ejemplo:
            # Parent: Current assets
            # Child: Inventories
            #
            # canonical hijo = inventory
            # section_detected = activo_corriente
            # parent_canonical = current_assets_total
            # -------------------------------------------------
            elif child_canonical is not None and parent_section is not None:

                # Solo alineamos section_detected.
                # No pisamos canonical ni aggregate del hijo.
                result_row["section_detected"] = parent_section

                # Si el hijo no tiene section_expected, también la rellenamos.
                if child_section_expected is None:
                    result_row["section_expected"] = parent_section

                result_row["mapping_overridden_by_parent"] = False
                result_row["section_aligned_to_parent"] = True
                result_row["match_type"] = (
                    str(result_row.get("match_type"))
                    + "|section_aligned_to_parent"
                )

        results[pos] = result_row

    return results

# =====================================================
# CANONICALS PERMITIDOS POR SECCIÓN
# Para usar con IA cuando canonical venga None
# =====================================================

CANONICALS_BY_SECTION = {
    "activo_no_corriente": [
        "non_current_assets_total",
        "right_of_use_assets",
        "intangible_assets",
        "goodwill",
        "property_plant_equipment",
        "investment_properties",
        "financial_assets_lp",
        "investments_in_associates",
        "loans_receivable_lp",
        "derivative_assets_lp",
        "deferred_tax_assets",
        "other_non_current_assets",
    ],

    "activo_corriente": [
        "current_assets_total",
        "inventory",
        "trade_receivables",
        "other_receivables",
        "income_tax_receivable",
        "other_current_assets",
        "other_financial_assets_cp",
        "temporary_financial_investments",
        "derivative_assets_cp",
        "prepayments",
        "assets_held_for_sale",
        "cash_and_equivalents",
    ],

    "activo_total": [
        "total_assets",
    ],

    "patrimonio_neto": [
        "equity_total",
        "equity_parent",
        "equity_minorities",
        "share_capital",
        "share_premium",
        "reserves",
        "retained_earnings",
        "treasury_shares",
        "valuation_adjustments",
        "grants_donations_bequests",
        "profit_loss_for_period_equity",
        "other_equity_items",
    ],

    "pasivo_no_corriente": [
        "non_current_liabilities_total",
        "financial_debt_lp",
        "lease_liability_lp",
        "provisions_lp",
        "deferred_tax_liabilities",
        "other_liabilities_lp",
        "other_financial_liabilities_lp",
        "derivative_liabilities_lp",
        "pension_obligations_lp",
        "deferred_income_lp",
    ],

    "pasivo_corriente": [
        "current_liabilities_total",
        "financial_debt_cp",
        "lease_liability_cp",
        "provisions_cp",
        "income_tax_payable",
        "trade_payables",
        "other_financial_liabilities_cp",
        "other_current_liabilities",
        "accruals",
        "contract_liabilities",
        "derivative_liabilities_cp",
        "payroll_social_security_payable",
    ],

    "pasivo_total": [
        "total_liabilities",
    ],

    "pasivo_total_pn": [
        "total_liabilities_and_equity",
    ],
}


# =====================================================
# BLOQUE 1 - PREPARAR COLUMNAS FINALES
# No pisa lo que ya detectó tu motor
# =====================================================

def initialize_final_columns(rows):
    final_rows = []

    for row in rows:
        new_row = dict(row)

        canonical = row.get("canonical")
        section_expected = row.get("section_expected")
        section_detected = row.get("section_detected")
        aggregate = row.get("aggregate")

        match_type = row.get("match_type")
        parent_canonical = row.get("parent_canonical")
        mapping_overridden_by_parent = row.get("mapping_overridden_by_parent", False)
        section_aligned_to_parent = row.get("section_aligned_to_parent", False)

        # Si hay canonical, priorizamos la sección esperada por el mapping.
        # Si no hay canonical, usamos la sección detectada por contexto.
        new_row["final_section"] = section_expected if canonical else section_detected
        new_row["final_canonical"] = canonical
        new_row["final_aggregate"] = aggregate

        # -------------------------------------------------
        # Diagnóstico del origen del mapping
        # -------------------------------------------------

        new_row["was_overridden_by_parent"] = mapping_overridden_by_parent
        new_row["section_aligned_to_parent"] = section_aligned_to_parent

        # Canonical genérico heredado del padre.
        # Ejemplo malo:
        #   Inventories -> current_assets_total
        # Esto conviene que lo revise la IA.
        new_row["canonical_equals_parent"] = (
            canonical is not None
            and parent_canonical is not None
            and canonical == parent_canonical
        )

        # Match débil: subset o heredado del padre.
        # Esto no significa que esté mal, solo que puede merecer revisión IA.
        match_type_str = str(match_type or "").lower()

        new_row["weak_rule_match"] = (
            "subset" in match_type_str
            or "inherited_from_parent" in match_type_str
            or mapping_overridden_by_parent is True
        )

        # Solo consideramos resuelto por reglas si hay canonical.
        # Pero ahora distinguimos si es regla fuerte o regla débil.
        if canonical:
            new_row["resolved_by"] = "rule_weak" if new_row["weak_rule_match"] else "rule"
        else:
            new_row["resolved_by"] = "unresolved"

        # -------------------------------------------------
        # Campos IA
        # -------------------------------------------------

        new_row["ai_confidence"] = None
        new_row["ai_reason"] = None
        new_row["needs_review"] = False

        new_row["ai_audit_section"] = None
        new_row["ai_audit_canonical"] = None
        new_row["ai_audit_aggregate"] = None
        new_row["ai_audit_confidence"] = None
        new_row["ai_audit_reason"] = None
        new_row["ai_disagrees_with_rule"] = False

        # Campo nuevo útil para saber si la IA sustituyó canonical/aggregate
        new_row["ai_applied"] = False
        new_row["ai_previous_canonical"] = canonical
        new_row["ai_previous_aggregate"] = aggregate

        final_rows.append(new_row)

    return final_rows

# =====================================================
# BLOQUE 2 - FILAS QUE PASAN A IA
# =====================================================

GENERIC_PARENT_CANONICALS = {
    # Activo
    "non_current_assets_total",
    "current_assets_total",
    "total_assets",

    # Patrimonio neto
    "equity_total",
    "equity_parent",
    "equity_minorities",

    # Pasivo
    "non_current_liabilities_total",
    "current_liabilities_total",
    "total_liabilities",
    "total_liabilities_and_equity",
}


def _is_empty_value(x):
    if x is None:
        return True

    try:
        if pd.isna(x):
            return True
    except Exception:
        pass

    return str(x).strip() in ["", "None", "nan", "NaN", "<NA>"]


def row_needs_ai(row):
    """
    Decide si una fila debe pasar a IA.

    Nueva lógica:
    - No analiza headers.
    - No analiza totales PyG sin detalle si ya están bien resueltos.
    - Sí analiza filas sin canonical.
    - Sí analiza filas con canonical genérico heredado del padre.
    - Sí analiza mappings débiles si tienen sección detectada.
    """

    # No analizar headers
    if row.get("is_header"):
        return False

    final_section = row.get("final_section") or row.get("section_detected")
    final_canonical = row.get("final_canonical")
    parent_canonical = row.get("parent_canonical")

    # Si no hay sección, la IA no tiene buen contexto.
    # Mejor no forzar.
    if _is_empty_value(final_section):
        return False

    # 1) Si no hay canonical, sí queremos IA
    if _is_empty_value(final_canonical):
        return True

    # 2) Si heredó canonical del padre, revisarlo con IA
    if row.get("was_overridden_by_parent"):
        return True

    # 3) Si canonical == parent_canonical y ese canonical es genérico,
    # probablemente perdió detalle.
    if (
        not _is_empty_value(parent_canonical)
        and final_canonical == parent_canonical
        and final_canonical in GENERIC_PARENT_CANONICALS
    ):
        return True

    # 4) Si lo marcamos como regla débil, que la IA audite.
    # Ej: subset, inherited_from_parent, etc.
    if row.get("weak_rule_match") is True:
        return True

    # 5) Si resolved_by viene como rule_weak, también
    if row.get("resolved_by") == "rule_weak":
        return True

    return False

# =====================================================
# BLOQUE 3 - CONTEXTO LOCAL PARA LA IA
# Le damos fila anterior, actual, siguiente, nivel, padre
# y canonicals permitidos por sección
# =====================================================

def build_ai_context(rows, idx):
    prev_row = rows[idx - 1] if idx > 0 else None
    curr_row = rows[idx]
    next_row = rows[idx + 1] if idx < len(rows) - 1 else None

    final_section = (
        curr_row.get("final_section")
        or curr_row.get("section_detected")
        or curr_row.get("parent_section")
    )

    statement = curr_row.get("statement")

    # Opciones cerradas para que la IA NO invente canonicals.
    allowed_canonicals = []

    if statement == "balance":
        allowed_canonicals = CANONICALS_BY_SECTION.get(final_section, [])

    elif statement == "pyg":
        # Puedes crear CANONICALS_BY_PYG_SECTION si quieres más adelante.
        # De momento sacamos los canonicals de LINE_MAPPING_RULES según sección.
        allowed_canonicals = sorted(list({
            rule["canonical"]
            for rule in LINE_MAPPING_RULES
            if rule.get("statement") == "pyg"
            and (
                rule.get("section") == final_section
                or final_section is None
            )
        }))

    return {
        "statement": statement,

        "previous_row": prev_row.get("raw_concept") if prev_row else None,
        "current_row": curr_row.get("raw_concept"),
        "next_row": next_row.get("raw_concept") if next_row else None,

        "level": curr_row.get("level"),
        "line_role": curr_row.get("line_role"),
        "is_header": curr_row.get("is_header"),

        "parent_row": curr_row.get("parent_raw_concept"),
        "parent_canonical": curr_row.get("parent_canonical"),
        "parent_aggregate": curr_row.get("parent_aggregate"),
        "parent_section": curr_row.get("parent_section"),

        "current_detected_section": final_section,
        "current_detected_canonical": curr_row.get("final_canonical"),
        "current_detected_aggregate": curr_row.get("final_aggregate"),

        "match_type": curr_row.get("match_type"),
        "matched_alias": curr_row.get("matched_alias"),
        "resolved_by": curr_row.get("resolved_by"),
        "weak_rule_match": curr_row.get("weak_rule_match"),
        "was_overridden_by_parent": curr_row.get("was_overridden_by_parent"),
        "section_aligned_to_parent": curr_row.get("section_aligned_to_parent"),

        "allowed_canonicals": allowed_canonicals,

        "current_value": curr_row.get("current"),
        "previous_value": curr_row.get("previous")
    }

# =====================================================
# BLOQUE 4 - LLAMADA REAL A IA
# Completa o audita canonical usando opciones cerradas
# =====================================================

import json

def _get_allowed_mapping_options(context):
    """
    Devuelve opciones cerradas para la IA.

    Prioridad:
    1) allowed_canonicals del contexto, ya filtrados por sección.
    2) Si no existen, usa reglas del statement + sección.
    """

    statement = normalize_text(context.get("statement"))
    section = context.get("current_detected_section")

    allowed_canonicals = context.get("allowed_canonicals") or []

    # Si no vienen opciones cerradas desde build_ai_context, las calculamos.
    if not allowed_canonicals:
        allowed_canonicals = sorted(list({
            rule["canonical"]
            for rule in LINE_MAPPING_RULES
            if normalize_text(rule.get("statement")) == statement
            and (
                section is None
                or rule.get("section") == section
            )
        }))

    # Construir sección y aggregate permitidos SOLO para esos canonicals.
    allowed_sections = sorted(list({
        rule["section"]
        for rule in LINE_MAPPING_RULES
        if normalize_text(rule.get("statement")) == statement
        and rule.get("canonical") in allowed_canonicals
    }))

    allowed_aggregates = sorted(list({
        rule["aggregate"]
        for rule in LINE_MAPPING_RULES
        if normalize_text(rule.get("statement")) == statement
        and rule.get("canonical") in allowed_canonicals
    }))

    # Añadir overrides críticos compatibles
    for canonical, rule in CRITICAL_LABEL_OVERRIDES.items():
        if normalize_text(rule.get("statement")) != statement:
            continue

        if canonical in allowed_canonicals:
            if rule.get("section") not in allowed_sections:
                allowed_sections.append(rule.get("section"))
            if rule.get("aggregate") not in allowed_aggregates:
                allowed_aggregates.append(rule.get("aggregate"))

    # En balance, si hay sección detectada, no dejamos que la IA se vaya a otra sección
    # salvo null. La sección viene de reglas y la consideramos contexto fuerte.
    if statement == "balance" and section is not None:
        allowed_sections = [section]

    if None not in allowed_sections:
        allowed_sections.append(None)

    if None not in allowed_canonicals:
        allowed_canonicals.append(None)

    if None not in allowed_aggregates:
        allowed_aggregates.append(None)

    return allowed_sections, allowed_canonicals, allowed_aggregates


def _aggregate_for_canonical(canonical):
    """
    Busca aggregate correspondiente a un canonical.
    """
    if canonical is None:
        return None

    for rule in LINE_MAPPING_RULES:
        if rule.get("canonical") == canonical:
            return rule.get("aggregate")

    if canonical in CRITICAL_LABEL_OVERRIDES:
        return CRITICAL_LABEL_OVERRIDES[canonical].get("aggregate")

    return None


def _section_for_canonical(canonical):
    """
    Busca section correspondiente a un canonical.
    """
    if canonical is None:
        return None

    for rule in LINE_MAPPING_RULES:
        if rule.get("canonical") == canonical:
            return rule.get("section")

    if canonical in CRITICAL_LABEL_OVERRIDES:
        return CRITICAL_LABEL_OVERRIDES[canonical].get("section")

    return None


def call_ai_for_mapping(context):
    statement = normalize_text(context.get("statement"))

    allowed_sections_for_statement, allowed_canonicals, allowed_aggregates = _get_allowed_mapping_options(context)

    current_canonical = context.get("current_detected_canonical")
    current_section = context.get("current_detected_section")
    parent_canonical = context.get("parent_canonical")
    parent_section = context.get("parent_section")

    prompt = f"""
Eres un clasificador financiero para líneas de balance y cuenta de resultados.

Tu tarea:
- Elegir el canonical más adecuado para la fila actual.
- Debes elegir SOLO entre los canonicals permitidos.
- Si no estás seguro, devuelve null.
- No inventes categorías nuevas.
- No mezcles Balance con PyG.
- En Balance, la sección detectada viene de reglas y debe respetarse.
- El padre sirve como contexto, pero NO debes copiar el canonical del padre si la línea hija tiene un concepto más específico.
- Ejemplo: si parent = Current assets y current_row = Inventories, el canonical correcto es inventory, NO current_assets_total.
- Ejemplo: si parent = Current assets y current_row = Trade receivables, el canonical correcto es trade_receivables.
- Ejemplo: si parent = Non-current liabilities y current_row = Long-term debt, el canonical correcto es financial_debt_lp.
- Si la fila actual es una partida genérica tipo Other current assets, Other liabilities, etc., elige el canonical genérico correspondiente.
- Si current_detected_canonical ya existe pero parece demasiado genérico o heredado del padre, puedes corregirlo.
- Si current_detected_canonical ya existe y es claramente específico y coherente, mantenlo.
- En PyG, diferencia earnings/profit/income antes de impuestos, impuestos, EBITDA, EBIT y net income.

Contexto:
{json.dumps(context, ensure_ascii=False, indent=2)}

Secciones permitidas:
{json.dumps(allowed_sections_for_statement, ensure_ascii=False)}

Canonicals permitidos:
{json.dumps(allowed_canonicals, ensure_ascii=False)}

Aggregates permitidos:
{json.dumps(allowed_aggregates, ensure_ascii=False)}

Devuelve:
- suggested_section: sección final sugerida o null.
- suggested_canonical: canonical final sugerido o null.
- suggested_aggregate: aggregate correspondiente al canonical o null.
- confidence: número entre 0 y 1.
- reason: explicación breve.

Reglas estrictas:
1. suggested_canonical debe estar en Canonicals permitidos.
2. suggested_section debe estar en Secciones permitidas.
3. suggested_aggregate debe estar en Aggregates permitidos.
4. Si suggested_canonical no es null, suggested_aggregate debe ser el aggregate coherente con ese canonical.
5. Si hay sección detectada en Balance, no la cambies.
6. No devuelvas texto fuera del JSON.
"""

    response = get_openai_client().responses.create(
        model="gpt-5.4",
        input=[
            {
                "role": "system",
                "content": "Responde únicamente con un JSON válido siguiendo el esquema solicitado."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        text={
            "format": {
                "type": "json_schema",
                "name": "financial_line_mapping",
                "schema": {
                    "type": "object",
                    "properties": {
                        "suggested_section": {
                            "type": ["string", "null"]
                        },
                        "suggested_canonical": {
                            "type": ["string", "null"]
                        },
                        "suggested_aggregate": {
                            "type": ["string", "null"]
                        },
                        "confidence": {
                            "type": "number"
                        },
                        "reason": {
                            "type": "string"
                        }
                    },
                    "required": [
                        "suggested_section",
                        "suggested_canonical",
                        "suggested_aggregate",
                        "confidence",
                        "reason"
                    ],
                    "additionalProperties": False
                },
                "strict": True
            }
        }
    )

    result = json.loads(response.output_text)

    # =====================================================
    # Limpieza defensiva
    # =====================================================

    suggested_canonical = result.get("suggested_canonical")
    suggested_section = result.get("suggested_section")
    suggested_aggregate = result.get("suggested_aggregate")

    if suggested_canonical not in allowed_canonicals:
        suggested_canonical = None

    # Si canonical válido, forzamos section/aggregate desde tu taxonomía
    # para que la IA no meta incoherencias.
    if suggested_canonical is not None:
        canonical_section = _section_for_canonical(suggested_canonical)
        canonical_aggregate = _aggregate_for_canonical(suggested_canonical)

        # En balance manda la sección reglada si existe.
        if statement == "balance" and current_section is not None:
            suggested_section = current_section
        else:
            suggested_section = canonical_section

        suggested_aggregate = canonical_aggregate

    else:
        suggested_section = None
        suggested_aggregate = None

    if suggested_section not in allowed_sections_for_statement:
        suggested_section = None

    if suggested_aggregate not in allowed_aggregates:
        suggested_aggregate = None

    result["suggested_section"] = suggested_section
    result["suggested_canonical"] = suggested_canonical
    result["suggested_aggregate"] = suggested_aggregate

    try:
        confidence = float(result.get("confidence", 0))
        result["confidence"] = max(0, min(1, confidence))
    except Exception:
        result["confidence"] = 0

    if not isinstance(result.get("reason"), str):
        result["reason"] = ""

    # Si no ha devuelto canonical, confianza máxima permitida baja
    if result["suggested_canonical"] is None:
        result["confidence"] = min(result["confidence"], 0.5)

    return result

# =====================================================
# BLOQUE EXTRA - AUDITORÍA IA VS REGLAS
# La IA revisa también filas ya clasificadas por reglas,
# pero NO pisa el resultado de las reglas.
# =====================================================

def row_needs_ai_audit(row):
    """
    Decide si una fila clasificada por reglas debe ser auditada por IA.
    """

    if row.get("is_header"):
        return False

    if row.get("resolved_by") != "rule":
        return False

    if row.get("final_canonical") is None:
        return False

    raw = row.get("raw_concept")
    if raw is None or str(raw).strip() == "":
        return False

    return True


def apply_ai_audit(row, ai_result, min_confidence=0.75):
    """
    Guarda la opinión de la IA sin modificar final_*.
    Marca discrepancia si la IA tiene confianza suficiente
    y opina distinto a las reglas.
    """

    if not ai_result:
        return row

    ai_section = ai_result.get("suggested_section")
    ai_canonical = ai_result.get("suggested_canonical")
    ai_aggregate = ai_result.get("suggested_aggregate")
    ai_confidence = ai_result.get("confidence", 0)

    row["ai_audit_section"] = ai_section
    row["ai_audit_canonical"] = ai_canonical
    row["ai_audit_aggregate"] = ai_aggregate
    row["ai_audit_confidence"] = ai_confidence
    row["ai_audit_reason"] = ai_result.get("reason")

    if ai_confidence < min_confidence:
        return row

    disagreements = []

    if (
        ai_canonical is not None
        and row.get("final_canonical") is not None
        and ai_canonical != row.get("final_canonical")
    ):
        disagreements.append("canonical")

    if (
        ai_section is not None
        and row.get("final_section") is not None
        and ai_section != row.get("final_section")
    ):
        disagreements.append("section")

    if (
        ai_aggregate is not None
        and row.get("final_aggregate") is not None
        and ai_aggregate != row.get("final_aggregate")
    ):
        disagreements.append("aggregate")

    row["ai_disagrees_with_rule"] = len(disagreements) > 0
    row["ai_disagreement_fields"] = disagreements

    return row

# =====================================================
# BLOQUE EXTRA - CALIDAD GLOBAL IA VS REGLAS
# Detecta si la IA discrepa mucho del motor de reglas
# =====================================================

def evaluate_rule_ai_agreement(rows, disagreement_threshold=0.25):
    audited_rows = [
        r for r in rows
        if str(r.get("resolved_by", "")).startswith("rule")
        and not r.get("is_header")
        and r.get("final_canonical") is not None
        and r.get("ai_audit_confidence") is not None
    ]

    total_audited = len(audited_rows)

    if total_audited == 0:
        return {
            "status": "not_audited",
            "message": "No hay filas auditadas por IA para comparar contra reglas.",
            "total_audited": 0
        }

    disagreements = [
        r for r in audited_rows
        if r.get("ai_disagrees_with_rule") is True
    ]

    pct_disagreement = len(disagreements) / total_audited

    if pct_disagreement >= disagreement_threshold:
        status = "review_required"
        message = (
            "Cuidado: la IA discrepa de las reglas en muchas partidas. "
            "Puede que el balance/PyG no se esté leyendo bien, que los niveles estén mal detectados "
            "o que el mapping esté forzando categorías incorrectas."
        )
    else:
        status = "ok"
        message = "La IA y las reglas parecen razonablemente alineadas."

    return {
        "status": status,
        "message": message,
        "total_audited": total_audited,
        "disagreement_rows": len(disagreements),
        "pct_disagreement": pct_disagreement,
        "examples": [
            {
                "raw_concept": r.get("raw_concept"),
                "rule_canonical": r.get("final_canonical"),
                "ai_canonical": r.get("ai_audit_canonical"),
                "rule_section": r.get("final_section"),
                "ai_section": r.get("ai_audit_section"),
                "ai_confidence": r.get("ai_audit_confidence"),
                "reason": r.get("ai_audit_reason")
            }
            for r in disagreements[:10]
        ]
    }

# =====================================================
# BLOQUE 5 - APLICAR IA SOBRE HUECOS O MAPPING DÉBIL
# =====================================================

def _is_blank_value(x):
    """
    Detecta None, NaN, pd.NA, "", "None", "nan".
    """
    if x is None:
        return True

    try:
        if pd.isna(x):
            return True
    except Exception:
        pass

    s = str(x).strip().lower()

    return s in ("", "none", "nan", "null", "<na>")


def apply_ai_suggestion(row, ai_result, min_confidence=0.80):
    if not ai_result:
        row["needs_review"] = True
        row["ai_applied"] = False
        row["ai_reason"] = "IA no devolvió resultado."
        return row

    confidence = ai_result.get("confidence", 0)

    try:
        confidence = float(confidence)
    except Exception:
        confidence = 0

    row["ai_confidence"] = confidence
    row["ai_reason"] = ai_result.get("reason")
    row["ai_applied"] = False

    suggested_section = ai_result.get("suggested_section")
    suggested_canonical = ai_result.get("suggested_canonical")
    suggested_aggregate = ai_result.get("suggested_aggregate")

    # Guardamos siempre lo que sugirió la IA para auditar
    row["ai_suggested_section"] = suggested_section
    row["ai_suggested_canonical"] = suggested_canonical
    row["ai_suggested_aggregate"] = suggested_aggregate

    if confidence < min_confidence:
        row["needs_review"] = True
        return row

    if _is_blank_value(suggested_canonical):
        row["needs_review"] = True
        return row

    applied = False

    current_canonical = row.get("final_canonical")
    current_aggregate = row.get("final_aggregate")
    current_section = row.get("final_section")

    # -------------------------------------------------
    # CASO 1: Huecos reales
    # Ahora detecta None, NaN, pd.NA, "", "None", etc.
    # -------------------------------------------------
    if _is_blank_value(current_canonical):
        row["final_canonical"] = suggested_canonical
        row["final_aggregate"] = suggested_aggregate
        row["final_section"] = suggested_section or current_section
        row["resolved_by"] = "ai"
        applied = True

    # -------------------------------------------------
    # CASO 2: Mapping débil / heredado / genérico
    # -------------------------------------------------
    else:
        can_replace = (
            row.get("weak_rule_match") is True
            or row.get("was_overridden_by_parent") is True
            or row.get("mapping_overridden_by_parent") is True
            or row.get("canonical_equals_parent") is True
            or row.get("resolved_by") == "rule_weak"
            or row.get("resolved_by") == "parent_inherited"
        )

        if can_replace and suggested_canonical != current_canonical:
            row["ai_previous_canonical"] = current_canonical
            row["ai_previous_aggregate"] = current_aggregate

            row["final_canonical"] = suggested_canonical
            row["final_aggregate"] = suggested_aggregate
            row["final_section"] = current_section or suggested_section

            row["resolved_by"] = "ai_corrected_rule"
            applied = True

        elif suggested_canonical == current_canonical:
            row["resolved_by"] = row.get("resolved_by") or "rule"
            applied = False

    row["ai_applied"] = applied

    if applied:
        row["needs_review"] = False

    return row

# =====================================================
# BLOQUE 6 - APLICAR IA A HUECOS / MAPPINGS DÉBILES
# =====================================================

def fill_none_with_ai(rows, verbose=False):
    updated_rows = []

    total_sent_to_ai = 0
    total_ai_applied = 0
    total_needs_review = 0

    def has_amount(row):
        """
        La fila tiene algún importe útil.
        """
        for key in ["current", "previous", "amount", "amount_1", "amount_2", "amount_3"]:
            v = row.get(key)

            if not _is_blank_value(v):
                try:
                    float(v)
                    return True
                except Exception:
                    pass

        return False

    def force_ai_for_unmapped_row(row):
        """
        Regla fuerte:
        Si una línea tiene importe y no tiene final_canonical,
        debe ir a IA.
        """
        if not has_amount(row):
            return False

        if _is_blank_value(row.get("final_canonical")):
            return True

        return False

    for idx, row in enumerate(rows):
        row = dict(row)

        send_to_ai = False

        # Regla existente
        try:
            if row_needs_ai(row):
                send_to_ai = True
        except Exception:
            pass

        # Nueva regla fuerte para huecos reales
        if force_ai_for_unmapped_row(row):
            send_to_ai = True

        if send_to_ai:
            total_sent_to_ai += 1

            context = build_ai_context(rows, idx)
            ai_result = call_ai_for_mapping(context)
            row = apply_ai_suggestion(row, ai_result)

            if row.get("ai_applied") is True:
                total_ai_applied += 1

            if row.get("needs_review") is True:
                total_needs_review += 1

        updated_rows.append(row)

    if verbose:
        print("Resumen IA:")
        print(f"- Filas enviadas a IA: {total_sent_to_ai}")
        print(f"- Correcciones/aplicaciones IA: {total_ai_applied}")
        print(f"- Filas marcadas para revisión: {total_needs_review}")

    return updated_rows

# =====================================================
# BLOQUE 7 - AUDITAR IA VS REGLAS
# No modifica final_*. Solo compara opinión IA vs reglas.
# =====================================================

def audit_rules_with_ai(rows):
    updated_rows = []

    for idx, row in enumerate(rows):
        row = dict(row)

        if row_needs_ai_audit(row):
            context = build_ai_context(rows, idx)
            ai_result = call_ai_for_mapping(context)
            row = apply_ai_audit(row, ai_result)

        updated_rows.append(row)

    quality_report = evaluate_rule_ai_agreement(updated_rows)

    return updated_rows, quality_report

# =====================================================
# BLOQUE UNICO - VALIDACIONES POST IA
# Se ejecuta despues de fill_none_with_ai(...)
# No modifica datos, solo valida
# =====================================================

# -------------------------------
# 1. Comparacion con tolerancia
# -------------------------------

def approx_equal(a, b, abs_tol=1.0, rel_tol=0.01):
    if a is None or b is None:
        return None

    try:
        a = float(a)
        b = float(b)
    except Exception:
        return None

    diff = abs(a - b)
    return diff <= abs_tol or diff <= abs(b) * rel_tol


# -------------------------------
# 2. Helpers de valores
# -------------------------------

def _is_number(x):
    return isinstance(x, (int, float)) and not isinstance(x, bool)


def get_first_value_by_canonical(rows, canonical_name):
    for row in rows:
        if row.get("final_canonical") == canonical_name:
            value = row.get("current")
            if _is_number(value):
                return value
    return None


def get_values_by_canonical(rows, canonical_name):
    values = []

    for row in rows:
        if row.get("final_canonical") == canonical_name:
            value = row.get("current")
            if _is_number(value):
                values.append(value)

    return values


def sum_canonicals(rows, canonical_names):
    total = 0
    found_any = False

    for row in rows:
        if row.get("final_canonical") not in canonical_names:
            continue

        value = row.get("current")
        if not _is_number(value):
            continue

        total += value
        found_any = True

    return total if found_any else None


def sum_section_components(rows, section_name, exclude_canonicals=None):
    if exclude_canonicals is None:
        exclude_canonicals = set()

    total = 0
    found_any = False

    for row in rows:
        if row.get("final_section") != section_name:
            continue

        canonical = row.get("final_canonical")

        if canonical in exclude_canonicals:
            continue

        value = row.get("current")
        if not _is_number(value):
            continue

        total += value
        found_any = True

    return total if found_any else None


# -------------------------------
# 3. Helper validación genérica
# -------------------------------

def validate_section_total(
    rows,
    section_name,
    total_canonical,
    exclude_canonicals=None,
    missing_total_reason=None,
    missing_components_reason=None
):
    if exclude_canonicals is None:
        exclude_canonicals = {total_canonical}
    else:
        exclude_canonicals = set(exclude_canonicals) | {total_canonical}

    reported_total = get_first_value_by_canonical(rows, total_canonical)

    if reported_total is None:
        return {
            "status": None,
            "reason": missing_total_reason or f"No existe total para {section_name}"
        }

    expected_sum = sum_section_components(
        rows,
        section_name,
        exclude_canonicals=exclude_canonicals
    )

    if expected_sum is None:
        return {
            "status": None,
            "reason": missing_components_reason or f"No hay suficientes partidas para {section_name}"
        }

    return {
        "status": approx_equal(expected_sum, reported_total),
        "expected": expected_sum,
        "reported": reported_total,
        "difference": None if expected_sum is None else reported_total - expected_sum
    }


# -------------------------------
# 4. Validaciones de secciones
# -------------------------------

def validate_non_current_assets(rows):
    return validate_section_total(
        rows,
        section_name="activo_no_corriente",
        total_canonical="non_current_assets_total",
        missing_total_reason="No existe total activo no corriente",
        missing_components_reason="No hay suficientes partidas ANC"
    )


def validate_current_assets(rows):
    return validate_section_total(
        rows,
        section_name="activo_corriente",
        total_canonical="current_assets_total",
        missing_total_reason="No existe total activo corriente",
        missing_components_reason="No hay suficientes partidas AC"
    )


def validate_equity(rows):
    return validate_section_total(
        rows,
        section_name="patrimonio_neto",
        total_canonical="equity_total",
        missing_total_reason="No existe total patrimonio neto",
        missing_components_reason="No hay suficientes partidas de patrimonio neto"
    )


def validate_non_current_liabilities(rows):
    return validate_section_total(
        rows,
        section_name="pasivo_no_corriente",
        total_canonical="non_current_liabilities_total",
        missing_total_reason="No existe total pasivo no corriente",
        missing_components_reason="No hay suficientes partidas PNC"
    )


def validate_current_liabilities(rows):
    return validate_section_total(
        rows,
        section_name="pasivo_corriente",
        total_canonical="current_liabilities_total",
        missing_total_reason="No existe total pasivo corriente",
        missing_components_reason="No hay suficientes partidas PC"
    )


# -------------------------------
# 5. Validaciones agregadas
# -------------------------------

def validate_total_assets(rows):
    total_assets = get_first_value_by_canonical(rows, "total_assets")
    total_anc = get_first_value_by_canonical(rows, "non_current_assets_total")
    total_ac = get_first_value_by_canonical(rows, "current_assets_total")

    if total_assets is None:
        return {
            "status": None,
            "reason": "No existe activo total"
        }

    if total_anc is None or total_ac is None:
        return {
            "status": None,
            "reason": "Falta ANC o AC total"
        }

    expected = total_anc + total_ac

    return {
        "status": approx_equal(expected, total_assets),
        "expected": expected,
        "reported": total_assets,
        "difference": total_assets - expected
    }


def validate_total_liabilities(rows):
    total_liabilities = get_first_value_by_canonical(rows, "total_liabilities")
    total_pnc = get_first_value_by_canonical(rows, "non_current_liabilities_total")
    total_pc = get_first_value_by_canonical(rows, "current_liabilities_total")

    if total_liabilities is None:
        return {
            "status": None,
            "reason": "No existe pasivo total"
        }

    if total_pnc is None or total_pc is None:
        return {
            "status": None,
            "reason": "Falta PNC o PC total"
        }

    expected = total_pnc + total_pc

    return {
        "status": approx_equal(expected, total_liabilities),
        "expected": expected,
        "reported": total_liabilities,
        "difference": total_liabilities - expected
    }


def validate_total_liabilities_and_equity(rows):
    total_liab_equity = get_first_value_by_canonical(rows, "total_liabilities_and_equity")
    total_liabilities = get_first_value_by_canonical(rows, "total_liabilities")
    total_pnc = get_first_value_by_canonical(rows, "non_current_liabilities_total")
    total_pc = get_first_value_by_canonical(rows, "current_liabilities_total")
    total_equity = get_first_value_by_canonical(rows, "equity_total")

    if total_liab_equity is None:
        return {
            "status": None,
            "reason": "No existe total pasivo + patrimonio neto"
        }

    if total_liabilities is not None and total_equity is not None:
        expected = total_liabilities + total_equity

        return {
            "status": approx_equal(expected, total_liab_equity),
            "expected": expected,
            "reported": total_liab_equity,
            "difference": total_liab_equity - expected,
            "method": "total_liabilities + equity_total"
        }

    if total_pnc is not None and total_pc is not None and total_equity is not None:
        expected = total_pnc + total_pc + total_equity

        return {
            "status": approx_equal(expected, total_liab_equity),
            "expected": expected,
            "reported": total_liab_equity,
            "difference": total_liab_equity - expected,
            "method": "pnc + pc + equity_total"
        }

    return {
        "status": None,
        "reason": "Faltan datos para validar total pasivo + patrimonio neto"
    }


def validate_balance_equation(rows):
    total_assets = get_first_value_by_canonical(rows, "total_assets")
    total_liab_equity = get_first_value_by_canonical(rows, "total_liabilities_and_equity")

    if total_assets is None:
        return {
            "status": None,
            "reason": "No existe activo total"
        }

    if total_liab_equity is not None:
        return {
            "status": approx_equal(total_assets, total_liab_equity),
            "expected": total_assets,
            "reported": total_liab_equity,
            "difference": total_liab_equity - total_assets,
            "method": "total_assets vs total_liabilities_and_equity"
        }

    total_liabilities = get_first_value_by_canonical(rows, "total_liabilities")
    total_pc = get_first_value_by_canonical(rows, "current_liabilities_total")
    total_pnc = get_first_value_by_canonical(rows, "non_current_liabilities_total")
    total_equity = get_first_value_by_canonical(rows, "equity_total")

    if total_liabilities is not None and total_equity is not None:
        reconstructed = total_liabilities + total_equity

        return {
            "status": approx_equal(total_assets, reconstructed),
            "expected": total_assets,
            "reported": reconstructed,
            "difference": reconstructed - total_assets,
            "method": "total_liabilities + equity_total"
        }

    if total_pc is None or total_pnc is None or total_equity is None:
        return {
            "status": None,
            "reason": "Faltan datos para reconstruir balance"
        }

    reconstructed = total_pc + total_pnc + total_equity

    return {
        "status": approx_equal(total_assets, reconstructed),
        "expected": total_assets,
        "reported": reconstructed,
        "difference": reconstructed - total_assets,
        "method": "pc + pnc + equity_total"
    }


# -------------------------------
# 6. Función final única
# -------------------------------

def run_all_validations(rows):
    return {
        "non_current_assets": validate_non_current_assets(rows),
        "current_assets": validate_current_assets(rows),
        "total_assets": validate_total_assets(rows),

        "equity": validate_equity(rows),
        "non_current_liabilities": validate_non_current_liabilities(rows),
        "current_liabilities": validate_current_liabilities(rows),
        "total_liabilities": validate_total_liabilities(rows),
        "total_liabilities_and_equity": validate_total_liabilities_and_equity(rows),

        "balance_equation": validate_balance_equation(rows),
    }

# =====================================================
# BLOQUE UNICO - WARNINGS NO BLOQUEANTES
# No frena el analisis. Solo avisa de posibles problemas.
# =====================================================

def _is_empty_warning_value(x):
    if x is None:
        return True

    try:
        if pd.isna(x):
            return True
    except Exception:
        pass

    return str(x).strip() in ["", "None", "nan", "NaN", "<NA>"]


def collect_analysis_warnings(rows, validation_results, min_ai_confidence=0.80):
    warnings = []

    # ---------------------------------
    # 1. Warnings de mapping / IA
    # ---------------------------------
    for row in rows:
        raw_concept = row.get("raw_concept")
        final_canonical = row.get("final_canonical")
        final_section = row.get("final_section")
        resolved_by = row.get("resolved_by")
        ai_confidence = row.get("ai_confidence")

        # Caso 1: sigue sin clasificar
        if _is_empty_warning_value(final_canonical) or _is_empty_warning_value(final_section):
            warnings.append({
                "type": "mapping_warning",
                "severity": "medium",
                "concept": raw_concept,
                "message": "Línea no clasificada completamente",
                "final_section": final_section,
                "final_canonical": final_canonical,
                "resolved_by": resolved_by
            })
            continue

        # Caso 2: marcada explícitamente para revisión
        if row.get("needs_review") is True:
            warnings.append({
                "type": "needs_review",
                "severity": "medium",
                "concept": raw_concept,
                "message": "Línea marcada para revisión tras reglas/IA",
                "final_section": final_section,
                "final_canonical": final_canonical,
                "resolved_by": resolved_by,
                "ai_confidence": ai_confidence,
                "ai_reason": row.get("ai_reason")
            })

        # Caso 3: clasificada/corregida por IA con baja confianza
        if (
            resolved_by in ["ai", "ai_corrected_rule"]
            and ai_confidence is not None
        ):
            try:
                conf = float(ai_confidence)
            except Exception:
                conf = None

            if conf is not None and conf < min_ai_confidence:
                warnings.append({
                    "type": "low_confidence_mapping",
                    "severity": "medium",
                    "concept": raw_concept,
                    "message": f"Línea clasificada/corregida por IA con confianza baja ({conf:.2f})",
                    "final_section": final_section,
                    "final_canonical": final_canonical,
                    "resolved_by": resolved_by,
                    "ai_reason": row.get("ai_reason")
                })

        # Caso 4: regla débil que no fue corregida por IA
        if row.get("weak_rule_match") is True and row.get("ai_applied") is not True:
            warnings.append({
                "type": "weak_rule_mapping",
                "severity": "low",
                "concept": raw_concept,
                "message": "Mapping resuelto por regla débil y no corregido por IA",
                "final_section": final_section,
                "final_canonical": final_canonical,
                "match_type": row.get("match_type")
            })

        # Caso 5: posible pérdida de detalle por canonical igual al padre
        if row.get("canonical_equals_parent") is True and row.get("ai_applied") is not True:
            warnings.append({
                "type": "possible_parent_overgeneralization",
                "severity": "medium",
                "concept": raw_concept,
                "message": "El canonical final coincide con el del padre; puede haber pérdida de detalle",
                "parent_raw_concept": row.get("parent_raw_concept"),
                "parent_canonical": row.get("parent_canonical"),
                "final_canonical": final_canonical
            })

    # ---------------------------------
    # 2. Validaciones contables
    # ---------------------------------
    if validation_results is not None:
        for check_name, result in validation_results.items():
            if not isinstance(result, dict):
                continue

            status = result.get("status")

            # Si no cuadra
            if status is False:
                warnings.append({
                    "type": "validation_error",
                    "severity": "high",
                    "check": check_name,
                    "message": f"La validación '{check_name}' no cuadra",
                    "expected": result.get("expected"),
                    "reported": result.get("reported"),
                    "difference": result.get("difference"),
                    "method": result.get("method")
                })

            # Si no se pudo validar
            elif status is None:
                warnings.append({
                    "type": "validation_skipped",
                    "severity": "low",
                    "check": check_name,
                    "message": result.get("reason", f"No se pudo ejecutar la validación '{check_name}'")
                })

    return warnings


# =====================================================
# HELPERS BASE
# =====================================================

import math
import pandas as pd


def is_missing_value(x):
    """
    Detecta None / NaN / pd.NA.
    """
    if x is None:
        return True
    try:
        return pd.isna(x)
    except Exception:
        return False


def safe_div(a, b):
    if is_missing_value(a) or is_missing_value(b) or b == 0:
        return None
    return a / b


def calc_yoy(current, previous):
    if is_missing_value(current) or is_missing_value(previous) or previous == 0:
        return None
    return (current - previous) / abs(previous)


# =====================================================
# EXTRAER METRICAS ROBUSTO
#
# Reglas importantes:
# 1) Usa SOLO final_canonical.
# 2) No usa canonical/original_canonical para métricas finales.
# 3) No mezcla niveles.
# 4) Si suma, suma únicamente filas del mismo level.
# 5) Evita filas heredadas del padre: mapping_overridden_by_parent=True.
# 6) Si hay una fila total/header clara, usa esa antes que sumar hijos.
# 7) EBITA queda unificado dentro de EBITDA.
# 8) Other operating income queda separado como other_gains_losses.
# =====================================================

def extract_financial_metrics(final_rows_balance, final_rows_pyg, debug=False):

    # -------------------------------------------------
    # Helpers internos
    # -------------------------------------------------

    def norm(x):
        if is_missing_value(x):
            return None
        return str(x).strip().lower()

    def is_number(x):
        if is_missing_value(x):
            return False
        return isinstance(x, (int, float)) and not isinstance(x, bool)

    def clean_number(x):
        return x if is_number(x) else None

    def row_bool(row, key):
        """
        Convierte valores tipo True/False/'True'/'False'/1/0 a bool.
        """
        v = row.get(key)

        if isinstance(v, bool):
            return v

        if isinstance(v, str):
            return v.strip().lower() in ("true", "1", "yes", "si", "sí")

        if isinstance(v, (int, float)) and not isinstance(v, bool):
            return v == 1

        return False

    def get_row_final_canonical(row):
        """
        Canonical final de la fila.

        Regla:
        - Para calcular métricas finales usamos SOLO final_canonical.
        - Si no existe o está vacío, la fila no sirve para métricas.
        """
        if "final_canonical" not in row:
            return None

        value = row.get("final_canonical")

        if is_missing_value(value):
            return None

        value = str(value).strip()

        if value == "":
            return None

        return value

    def get_row_value(row, period):
        """
        Lee current / previous.
        """
        value = row.get(period)
        return clean_number(value)

    def row_has_any_value(row):
        return (
            get_row_value(row, "current") is not None
            or get_row_value(row, "previous") is not None
        )

    def get_level(row):
        """
        Devuelve level como int si existe.
        Si no existe, devuelve None.
        """
        lvl = row.get("level")

        if is_missing_value(lvl):
            return None

        try:
            return int(float(lvl))
        except Exception:
            return None

    def is_total_or_header_row(row):
        """
        Detecta filas que pueden representar un total/agregado claro.
        """
        line_role = str(row.get("line_role", "")).strip().lower()

        if line_role == "total":
            return True

        if row_bool(row, "is_header"):
            return True

        raw = str(row.get("raw_concept", "") or "").strip().lower()

        total_words = [
            "total",
            "subtotal",
            "resultado",
            "beneficio",
            "pérdida",
            "perdida",
            "margen",
            "ebit",
            "ebitda",
            "ebita",
            "gross profit",
            "operating profit",
            "operating income",
            "operating result",
            "profit before tax",
            "net income",
            "net profit",
            "total assets",
            "total liabilities",
            "total equity",
            "current assets",
            "non current assets",
            "non-current assets",
            "current liabilities",
            "non current liabilities",
            "non-current liabilities",
        ]

        return any(w in raw for w in total_words)

    def is_usable_metric_row(row):
      """
      Decide si una fila puede usarse para extraer métricas.

      Regla:
      - Si no tiene final_canonical, no sirve.
      - Si no tiene importes, no sirve.
      - Si fue heredada del padre, normalmente se descarta.
      - Pero si la IA la corrigió y tiene final_canonical propio, sí se permite.
      """

      if get_row_final_canonical(row) is None:
          return False

      if not row_has_any_value(row):
          return False

      # Si fue heredada del padre, normalmente no la usamos
      if row_bool(row, "mapping_overridden_by_parent"):

          resolved_by = str(row.get("resolved_by", "")).strip().lower()

          # Excepción: la IA corrigió la fila.
          # Ejemplo:
          # canonical inicial heredado = current_assets_total
          # final_canonical corregido = cash_and_equivalents
          if resolved_by in [
              "ai_corrected_rule",
              "ai",
              "ai_filled_gap"
          ]:
              return True

          return False

      return True

    def row_score(row):
        """
        Score para elegir una fila única cuando hay total/header claro.

        Preferimos:
        - total/header
        - level 1 antes que level 2/3
        - filas con current y previous
        """
        score = 0

        if is_total_or_header_row(row):
            score += 100

        lvl = get_level(row)

        if lvl == 1:
            score += 30
        elif lvl == 2:
            score += 20
        elif lvl == 3:
            score += 10

        if get_row_value(row, "current") is not None:
            score += 5

        if get_row_value(row, "previous") is not None:
            score += 5

        return score

    def sum_rows_same_level(rows):
        """
        Suma varias filas, pero TODAS deben ser del mismo level.
        """
        current_values = []
        previous_values = []

        for row in rows:
            current = get_row_value(row, "current")
            previous = get_row_value(row, "previous")

            if current is not None:
                current_values.append(current)

            if previous is not None:
                previous_values.append(previous)

        current_sum = sum(current_values) if current_values else None
        previous_sum = sum(previous_values) if previous_values else None

        return current_sum, previous_sum, rows

    def choose_best_same_level_group(candidates):
        """
        Cuando no hay total/header claro, permite sumar varias filas
        pero SOLO del mismo level.
        """
        groups = {}

        for row in candidates:
            lvl = get_level(row)

            if lvl is None:
                lvl = 999

            groups.setdefault(lvl, []).append(row)

        ranked_groups = []

        for lvl, rows in groups.items():
            value_count = 0

            for row in rows:
                if get_row_value(row, "current") is not None:
                    value_count += 1
                if get_row_value(row, "previous") is not None:
                    value_count += 1

            ranked_groups.append({
                "level": lvl,
                "rows": rows,
                "row_count": len(rows),
                "value_count": value_count,
            })

        ranked_groups = sorted(
            ranked_groups,
            key=lambda x: (
                x["level"],
                -x["value_count"],
                -x["row_count"]
            )
        )

        return ranked_groups[0]["rows"], ranked_groups[0]["level"]

    def get_metric(rows, canonical):
        """
        Busca una métrica por final_canonical.

        Lógica:
        1) Filtra filas usables con final_canonical == canonical.
        2) Si hay total/header claro, usa la mejor fila única.
        3) Si no hay total/header claro:
             - si solo hay una fila, usa esa fila
             - si hay varias, suma SOLO filas del mismo level
        """

        canonical_norm = norm(canonical)

        candidates = []

        for row in rows:
            if not is_usable_metric_row(row):
                continue

            row_canonical = norm(get_row_final_canonical(row))

            if row_canonical == canonical_norm:
                candidates.append(row)

        if not candidates:
            if debug:
                print(f"[NO MATCH] final_canonical={canonical}")

            return {
                "current": None,
                "previous": None,
                "yoy": None,
                "matched_canonical": None,
                "matched_raw_concept": None,
                "matched_level": None,
                "matched_line_role": None,
                "aggregation_method": None,
                "matched_rows_count": 0,
                "matched_rows": [],
            }

        # -------------------------------------------------
        # 1) Si hay total/header claro, usamos una única fila.
        # -------------------------------------------------
        total_candidates = [
            row for row in candidates
            if is_total_or_header_row(row)
        ]

        if total_candidates:
            best_row = sorted(
                total_candidates,
                key=lambda r: row_score(r),
                reverse=True
            )[0]

            current = get_row_value(best_row, "current")
            previous = get_row_value(best_row, "previous")

            if debug:
                print(
                    f"[MATCH SINGLE TOTAL] canonical={canonical} | "
                    f"raw={best_row.get('raw_concept')} | "
                    f"level={best_row.get('level')} | "
                    f"current={current} | previous={previous}"
                )

            return {
                "current": current,
                "previous": previous,
                "yoy": calc_yoy(current, previous),
                "matched_canonical": canonical,
                "matched_raw_concept": best_row.get("raw_concept"),
                "matched_level": get_level(best_row),
                "matched_line_role": best_row.get("line_role"),
                "aggregation_method": "single_total_or_header",
                "matched_rows_count": 1,
                "matched_rows": [
                    {
                        "raw_concept": best_row.get("raw_concept"),
                        "level": get_level(best_row),
                        "current": current,
                        "previous": previous,
                    }
                ],
            }

        # -------------------------------------------------
        # 2) Si solo hay una fila, usamos esa fila.
        # -------------------------------------------------
        if len(candidates) == 1:
            row = candidates[0]

            current = get_row_value(row, "current")
            previous = get_row_value(row, "previous")

            if debug:
                print(
                    f"[MATCH SINGLE] canonical={canonical} | "
                    f"raw={row.get('raw_concept')} | "
                    f"level={row.get('level')} | "
                    f"current={current} | previous={previous}"
                )

            return {
                "current": current,
                "previous": previous,
                "yoy": calc_yoy(current, previous),
                "matched_canonical": canonical,
                "matched_raw_concept": row.get("raw_concept"),
                "matched_level": get_level(row),
                "matched_line_role": row.get("line_role"),
                "aggregation_method": "single_row",
                "matched_rows_count": 1,
                "matched_rows": [
                    {
                        "raw_concept": row.get("raw_concept"),
                        "level": get_level(row),
                        "current": current,
                        "previous": previous,
                    }
                ],
            }

        # -------------------------------------------------
        # 3) Hay varias filas. Sumamos SOLO mismo level.
        # -------------------------------------------------
        rows_to_sum, selected_level = choose_best_same_level_group(candidates)

        current_sum, previous_sum, used_rows = sum_rows_same_level(rows_to_sum)

        if debug:
            print(
                f"[MATCH SUM SAME LEVEL] canonical={canonical} | "
                f"selected_level={selected_level} | "
                f"rows={len(used_rows)} | "
                f"current={current_sum} | previous={previous_sum}"
            )

            for r in used_rows:
                print(
                    f"   - raw={r.get('raw_concept')} | "
                    f"level={r.get('level')} | "
                    f"current={get_row_value(r, 'current')} | "
                    f"previous={get_row_value(r, 'previous')}"
                )

        return {
            "current": current_sum,
            "previous": previous_sum,
            "yoy": calc_yoy(current_sum, previous_sum),
            "matched_canonical": canonical,
            "matched_raw_concept": "SUM_SAME_LEVEL",
            "matched_level": selected_level if selected_level != 999 else None,
            "matched_line_role": "derived_sum_same_level",
            "aggregation_method": "sum_same_level",
            "matched_rows_count": len(used_rows),
            "matched_rows": [
                {
                    "raw_concept": r.get("raw_concept"),
                    "level": get_level(r),
                    "current": get_row_value(r, "current"),
                    "previous": get_row_value(r, "previous"),
                }
                for r in used_rows
            ],
        }

    def get_metric_any(rows, metric_name, canonical_candidates):
        """
        Busca una métrica probando varios final_canonical posibles.
        """
        for canonical in canonical_candidates:
            result = get_metric(rows, canonical)

            if result["current"] is not None or result["previous"] is not None:
                result["metric_name"] = metric_name
                return result

        return {
            "current": None,
            "previous": None,
            "yoy": None,
            "matched_canonical": None,
            "matched_raw_concept": None,
            "matched_level": None,
            "matched_line_role": None,
            "aggregation_method": None,
            "matched_rows_count": 0,
            "matched_rows": [],
            "metric_name": metric_name,
        }

    def sum_metric_components(metrics, component_names):
        """
        Suma métricas ya extraídas.

        Esto NO suma filas de distinto nivel.
        Aquí ya estamos sumando métricas finales limpias.
        """
        current_values = []
        previous_values = []

        for name in component_names:
            metric = metrics.get(name, {})

            current = metric.get("current")
            previous = metric.get("previous")

            if current is not None:
                current_values.append(current)

            if previous is not None:
                previous_values.append(previous)

        current_sum = sum(current_values) if current_values else None
        previous_sum = sum(previous_values) if previous_values else None

        return current_sum, previous_sum

    # =================================================
    # INICIALIZAR METRICS
    # =================================================

    metrics = {}

    # =================================================
    # PYG
    # =================================================

    metrics["revenue"] = get_metric_any(final_rows_pyg, "revenue", [
        "operating_revenue",
    ])

    metrics["financial_income"] = get_metric_any(final_rows_pyg, "financial_income", [
        "financial_income",
    ])

    metrics["cost_of_goods_sold"] = get_metric_any(final_rows_pyg, "cost_of_goods_sold", [
        "cost_of_goods_sold",
    ])

    metrics["gross_profit"] = get_metric_any(final_rows_pyg, "gross_profit", [
        "gross_profit",
    ])

    metrics["gross_margin_pct_reported"] = get_metric_any(final_rows_pyg, "gross_margin_pct_reported", [
        "gross_margin_pct",
    ])

    metrics["personnel_expenses"] = get_metric_any(final_rows_pyg, "personnel_expenses", [
        "personnel_expenses",
    ])

    metrics["operating_expenses"] = get_metric_any(final_rows_pyg, "operating_expenses", [
        "operating_expenses",
    ])

    metrics["other_gains_losses"] = get_metric_any(final_rows_pyg, "other_gains_losses", [
        "other_gains_losses",
    ])

    metrics["ebitda"] = get_metric_any(final_rows_pyg, "ebitda", [
        "ebitda",
    ])

    metrics["depreciation_amortization"] = get_metric_any(final_rows_pyg, "depreciation_amortization", [
        "depreciation_amortization",
    ])

    metrics["ebit"] = get_metric_any(final_rows_pyg, "ebit", [
        "ebit",
    ])

    metrics["interest_expense"] = get_metric_any(final_rows_pyg, "interest_expense", [
        "interest_expense",
    ])

    metrics["equity_method_result"] = get_metric_any(final_rows_pyg, "equity_method_result", [
        "equity_method_result",
    ])

    metrics["profit_before_tax"] = get_metric_any(final_rows_pyg, "profit_before_tax", [
        "profit_before_tax",
    ])

    metrics["income_tax"] = get_metric_any(final_rows_pyg, "income_tax", [
        "income_tax",
    ])

    metrics["net_income"] = get_metric_any(final_rows_pyg, "net_income", [
        "net_income",
        "net_income_parent",
    ])

    metrics["minority_income"] = get_metric_any(final_rows_pyg, "minority_income", [
        "minority_income",
    ])

    metrics["net_income_parent"] = get_metric_any(final_rows_pyg, "net_income_parent", [
        "net_income_parent",
    ])

    metrics["eps"] = get_metric_any(final_rows_pyg, "eps", [
        "eps",
    ])

    # =================================================
    # BALANCE - ACTIVO
    # =================================================

    metrics["non_current_assets_total"] = get_metric_any(final_rows_balance, "non_current_assets_total", [
        "non_current_assets_total",
    ])

    metrics["right_of_use_assets"] = get_metric_any(final_rows_balance, "right_of_use_assets", [
        "right_of_use_assets",
    ])

    metrics["intangible_assets"] = get_metric_any(final_rows_balance, "intangible_assets", [
        "intangible_assets",
    ])

    metrics["goodwill"] = get_metric_any(final_rows_balance, "goodwill", [
        "goodwill",
    ])

    metrics["property_plant_equipment"] = get_metric_any(final_rows_balance, "property_plant_equipment", [
        "property_plant_equipment",
    ])

    metrics["investment_properties"] = get_metric_any(final_rows_balance, "investment_properties", [
        "investment_properties",
    ])

    metrics["financial_assets_lp"] = get_metric_any(final_rows_balance, "financial_assets_lp", [
        "financial_assets_lp",
    ])

    metrics["other_non_current_assets"] = get_metric_any(final_rows_balance, "other_non_current_assets", [
        "other_non_current_assets",
    ])

    metrics["deferred_tax_assets"] = get_metric_any(final_rows_balance, "deferred_tax_assets", [
        "deferred_tax_assets",
    ])

    metrics["investments_in_associates"] = get_metric_any(final_rows_balance, "investments_in_associates", [
        "investments_in_associates",
    ])

    metrics["loans_receivable_lp"] = get_metric_any(final_rows_balance, "loans_receivable_lp", [
        "loans_receivable_lp",
    ])

    metrics["derivative_assets_lp"] = get_metric_any(final_rows_balance, "derivative_assets_lp", [
        "derivative_assets_lp",
    ])

    metrics["current_assets_total"] = get_metric_any(final_rows_balance, "current_assets_total", [
        "current_assets_total",
    ])

    metrics["inventory"] = get_metric_any(final_rows_balance, "inventory", [
        "inventory",
    ])

    metrics["trade_receivables"] = get_metric_any(final_rows_balance, "trade_receivables", [
        "trade_receivables",
    ])

    metrics["other_receivables"] = get_metric_any(final_rows_balance, "other_receivables", [
        "other_receivables",
    ])

    metrics["income_tax_receivable"] = get_metric_any(final_rows_balance, "income_tax_receivable", [
        "income_tax_receivable",
    ])

    metrics["other_current_assets"] = get_metric_any(final_rows_balance, "other_current_assets", [
        "other_current_assets",
    ])

    metrics["other_financial_assets_cp"] = get_metric_any(final_rows_balance, "other_financial_assets_cp", [
        "other_financial_assets_cp",
    ])

    metrics["temporary_financial_investments"] = get_metric_any(final_rows_balance, "temporary_financial_investments", [
        "temporary_financial_investments",
    ])

    metrics["cash_and_equivalents"] = get_metric_any(final_rows_balance, "cash_and_equivalents", [
        "cash_and_equivalents",
    ])

    metrics["prepayments"] = get_metric_any(final_rows_balance, "prepayments", [
        "prepayments",
    ])

    metrics["derivative_assets_cp"] = get_metric_any(final_rows_balance, "derivative_assets_cp", [
        "derivative_assets_cp",
    ])

    metrics["assets_held_for_sale"] = get_metric_any(final_rows_balance, "assets_held_for_sale", [
        "assets_held_for_sale",
    ])

    metrics["total_assets"] = get_metric_any(final_rows_balance, "total_assets", [
        "total_assets",
    ])

    # =================================================
    # BALANCE - PATRIMONIO NETO
    # =================================================

    metrics["equity_total"] = get_metric_any(final_rows_balance, "equity_total", [
        "equity_total",
    ])

    metrics["equity_parent"] = get_metric_any(final_rows_balance, "equity_parent", [
        "equity_parent",
    ])

    metrics["equity_minorities"] = get_metric_any(final_rows_balance, "equity_minorities", [
        "equity_minorities",
    ])

    metrics["share_capital"] = get_metric_any(final_rows_balance, "share_capital", [
        "share_capital",
    ])

    metrics["share_premium"] = get_metric_any(final_rows_balance, "share_premium", [
        "share_premium",
    ])

    metrics["reserves"] = get_metric_any(final_rows_balance, "reserves", [
        "reserves",
    ])

    metrics["retained_earnings"] = get_metric_any(final_rows_balance, "retained_earnings", [
        "retained_earnings",
    ])

    metrics["treasury_shares"] = get_metric_any(final_rows_balance, "treasury_shares", [
        "treasury_shares",
    ])

    metrics["valuation_adjustments"] = get_metric_any(final_rows_balance, "valuation_adjustments", [
        "valuation_adjustments",
    ])

    metrics["grants_donations_bequests"] = get_metric_any(final_rows_balance, "grants_donations_bequests", [
        "grants_donations_bequests",
    ])

    metrics["profit_loss_for_period_equity"] = get_metric_any(final_rows_balance, "profit_loss_for_period_equity", [
        "profit_loss_for_period_equity",
    ])

    metrics["other_equity_items"] = get_metric_any(final_rows_balance, "other_equity_items", [
        "other_equity_items",
    ])

    # =================================================
    # BALANCE - PASIVO
    # =================================================

    metrics["non_current_liabilities_total"] = get_metric_any(final_rows_balance, "non_current_liabilities_total", [
        "non_current_liabilities_total",
    ])

    metrics["provisions_lp"] = get_metric_any(final_rows_balance, "provisions_lp", [
        "provisions_lp",
    ])

    metrics["other_liabilities_lp"] = get_metric_any(final_rows_balance, "other_liabilities_lp", [
        "other_liabilities_lp",
    ])

    metrics["financial_debt_lp"] = get_metric_any(final_rows_balance, "financial_debt_lp", [
        "financial_debt_lp",
    ])

    metrics["lease_liability_lp"] = get_metric_any(final_rows_balance, "lease_liability_lp", [
        "lease_liability_lp",
    ])

    metrics["deferred_tax_liabilities"] = get_metric_any(final_rows_balance, "deferred_tax_liabilities", [
        "deferred_tax_liabilities",
    ])

    metrics["other_financial_liabilities_lp"] = get_metric_any(final_rows_balance, "other_financial_liabilities_lp", [
        "other_financial_liabilities_lp",
    ])

    metrics["derivative_liabilities_lp"] = get_metric_any(final_rows_balance, "derivative_liabilities_lp", [
        "derivative_liabilities_lp",
    ])

    metrics["pension_obligations_lp"] = get_metric_any(final_rows_balance, "pension_obligations_lp", [
        "pension_obligations_lp",
    ])

    metrics["deferred_income_lp"] = get_metric_any(final_rows_balance, "deferred_income_lp", [
        "deferred_income_lp",
    ])

    metrics["current_liabilities_total"] = get_metric_any(final_rows_balance, "current_liabilities_total", [
        "current_liabilities_total",
    ])

    metrics["financial_debt_cp"] = get_metric_any(final_rows_balance, "financial_debt_cp", [
        "financial_debt_cp",
    ])

    metrics["other_financial_liabilities_cp"] = get_metric_any(final_rows_balance, "other_financial_liabilities_cp", [
        "other_financial_liabilities_cp",
    ])

    metrics["provisions_cp"] = get_metric_any(final_rows_balance, "provisions_cp", [
        "provisions_cp",
    ])

    metrics["lease_liability_cp"] = get_metric_any(final_rows_balance, "lease_liability_cp", [
        "lease_liability_cp",
    ])

    metrics["income_tax_payable"] = get_metric_any(final_rows_balance, "income_tax_payable", [
        "income_tax_payable",
    ])

    metrics["trade_payables"] = get_metric_any(final_rows_balance, "trade_payables", [
        "trade_payables",
    ])

    metrics["other_current_liabilities"] = get_metric_any(final_rows_balance, "other_current_liabilities", [
        "other_current_liabilities",
    ])

    metrics["accruals"] = get_metric_any(final_rows_balance, "accruals", [
        "accruals",
    ])

    metrics["contract_liabilities"] = get_metric_any(final_rows_balance, "contract_liabilities", [
        "contract_liabilities",
    ])

    metrics["derivative_liabilities_cp"] = get_metric_any(final_rows_balance, "derivative_liabilities_cp", [
        "derivative_liabilities_cp",
    ])

    metrics["payroll_social_security_payable"] = get_metric_any(final_rows_balance, "payroll_social_security_payable", [
        "payroll_social_security_payable",
    ])

    metrics["total_liabilities"] = get_metric_any(final_rows_balance, "total_liabilities", [
        "total_liabilities",
    ])

    metrics["total_liabilities_and_equity"] = get_metric_any(final_rows_balance, "total_liabilities_and_equity", [
        "total_liabilities_and_equity",
    ])

    # =================================================
    # MÉTRICAS DERIVADAS
    # =================================================

    # -----------------------------
    # Total operativo auxiliar
    # revenue + other_gains_losses
    #
    # Importante:
    # - No sustituye a revenue.
    # - Sirve para reconstrucciones operativas.
    # -----------------------------
    metrics["operating_income_total"] = {}

    for period in ["current", "previous"]:
        revenue = metrics["revenue"].get(period)
        other_operating = metrics["other_gains_losses"].get(period)

        if revenue is not None and other_operating is not None:
            metrics["operating_income_total"][period] = revenue + other_operating
        elif revenue is not None:
            metrics["operating_income_total"][period] = revenue
        else:
            metrics["operating_income_total"][period] = None

    metrics["operating_income_total"]["yoy"] = calc_yoy(
        metrics["operating_income_total"].get("current"),
        metrics["operating_income_total"].get("previous")
    )
    metrics["operating_income_total"]["matched_canonical"] = "operating_revenue + other_gains_losses"
    metrics["operating_income_total"]["matched_raw_concept"] = "DERIVED"
    metrics["operating_income_total"]["matched_level"] = None
    metrics["operating_income_total"]["matched_line_role"] = "derived"
    metrics["operating_income_total"]["aggregation_method"] = "derived_calculation"
    metrics["operating_income_total"]["matched_rows_count"] = None
    metrics["operating_income_total"]["matched_rows"] = []
    metrics["operating_income_total"]["metric_name"] = "operating_income_total"

    # -----------------------------
    # Pasivo total reconstruido si no viene directo
    # -----------------------------
    if metrics["total_liabilities"]["current"] is None and metrics["total_liabilities"]["previous"] is None:
        current_sum, previous_sum = sum_metric_components(
            metrics,
            [
                "current_liabilities_total",
                "non_current_liabilities_total"
            ]
        )

        metrics["total_liabilities"] = {
            "current": current_sum,
            "previous": previous_sum,
            "yoy": calc_yoy(current_sum, previous_sum),
            "matched_canonical": "current_liabilities_total + non_current_liabilities_total",
            "matched_raw_concept": "DERIVED",
            "matched_level": None,
            "matched_line_role": "derived",
            "aggregation_method": "derived_sum_metrics",
            "matched_rows_count": None,
            "matched_rows": [],
            "metric_name": "total_liabilities",
        }

    # -----------------------------
    # Deuda financiera total
    # Incluye deuda financiera + arrendamientos.
    # -----------------------------
    financial_debt_current, financial_debt_previous = sum_metric_components(
        metrics,
        [
            "financial_debt_lp",
            "financial_debt_cp",
            "lease_liability_lp",
            "lease_liability_cp"
        ]
    )

    metrics["financial_debt_total"] = {
        "current": financial_debt_current,
        "previous": financial_debt_previous,
        "yoy": calc_yoy(financial_debt_current, financial_debt_previous),
        "matched_canonical": "financial_debt_lp + financial_debt_cp + lease_liability_lp + lease_liability_cp",
        "matched_raw_concept": "DERIVED",
        "matched_level": None,
        "matched_line_role": "derived",
        "aggregation_method": "derived_sum_metrics",
        "matched_rows_count": None,
        "matched_rows": [],
        "metric_name": "financial_debt_total",
    }

    # -----------------------------
    # Deuda financiera sin leases
    # -----------------------------
    bank_debt_current, bank_debt_previous = sum_metric_components(
        metrics,
        [
            "financial_debt_lp",
            "financial_debt_cp"
        ]
    )

    metrics["bank_financial_debt_total"] = {
        "current": bank_debt_current,
        "previous": bank_debt_previous,
        "yoy": calc_yoy(bank_debt_current, bank_debt_previous),
        "matched_canonical": "financial_debt_lp + financial_debt_cp",
        "matched_raw_concept": "DERIVED",
        "matched_level": None,
        "matched_line_role": "derived",
        "aggregation_method": "derived_sum_metrics",
        "matched_rows_count": None,
        "matched_rows": [],
        "metric_name": "bank_financial_debt_total",
    }

    # -----------------------------
    # Deuda neta
    # -----------------------------
    metrics["net_debt"] = {}

    for period in ["current", "previous"]:
        debt = metrics["financial_debt_total"].get(period)
        cash = metrics["cash_and_equivalents"].get(period)

        if debt is not None and cash is not None:
            metrics["net_debt"][period] = debt - cash
        else:
            metrics["net_debt"][period] = None

    metrics["net_debt"]["yoy"] = calc_yoy(
        metrics["net_debt"].get("current"),
        metrics["net_debt"].get("previous")
    )
    metrics["net_debt"]["matched_canonical"] = "financial_debt_total - cash_and_equivalents"
    metrics["net_debt"]["matched_raw_concept"] = "DERIVED"
    metrics["net_debt"]["matched_level"] = None
    metrics["net_debt"]["matched_line_role"] = "derived"
    metrics["net_debt"]["aggregation_method"] = "derived_calculation"
    metrics["net_debt"]["matched_rows_count"] = None
    metrics["net_debt"]["matched_rows"] = []
    metrics["net_debt"]["metric_name"] = "net_debt"

    # -----------------------------
    # Working capital
    # -----------------------------
    metrics["working_capital"] = {}

    for period in ["current", "previous"]:
        current_assets = metrics["current_assets_total"].get(period)
        current_liabilities = metrics["current_liabilities_total"].get(period)

        if current_assets is not None and current_liabilities is not None:
            metrics["working_capital"][period] = current_assets - current_liabilities
        else:
            metrics["working_capital"][period] = None

    metrics["working_capital"]["yoy"] = calc_yoy(
        metrics["working_capital"].get("current"),
        metrics["working_capital"].get("previous")
    )
    metrics["working_capital"]["matched_canonical"] = "current_assets_total - current_liabilities_total"
    metrics["working_capital"]["matched_raw_concept"] = "DERIVED"
    metrics["working_capital"]["matched_level"] = None
    metrics["working_capital"]["matched_line_role"] = "derived"
    metrics["working_capital"]["aggregation_method"] = "derived_calculation"
    metrics["working_capital"]["matched_rows_count"] = None
    metrics["working_capital"]["matched_rows"] = []
    metrics["working_capital"]["metric_name"] = "working_capital"

    # -----------------------------
    # Operating working capital aproximado
    # inventory + trade_receivables - trade_payables
    # -----------------------------
    metrics["operating_working_capital"] = {}

    for period in ["current", "previous"]:
        inventory = metrics["inventory"].get(period)
        receivables = metrics["trade_receivables"].get(period)
        payables = metrics["trade_payables"].get(period)

        if inventory is not None and receivables is not None and payables is not None:
            metrics["operating_working_capital"][period] = inventory + receivables - payables
        else:
            metrics["operating_working_capital"][period] = None

    metrics["operating_working_capital"]["yoy"] = calc_yoy(
        metrics["operating_working_capital"].get("current"),
        metrics["operating_working_capital"].get("previous")
    )
    metrics["operating_working_capital"]["matched_canonical"] = "inventory + trade_receivables - trade_payables"
    metrics["operating_working_capital"]["matched_raw_concept"] = "DERIVED"
    metrics["operating_working_capital"]["matched_level"] = None
    metrics["operating_working_capital"]["matched_line_role"] = "derived"
    metrics["operating_working_capital"]["aggregation_method"] = "derived_calculation"
    metrics["operating_working_capital"]["matched_rows_count"] = None
    metrics["operating_working_capital"]["matched_rows"] = []
    metrics["operating_working_capital"]["metric_name"] = "operating_working_capital"

    return metrics

# =====================================================
# CALCULAR RATIOS AVANZADOS + DEBUG ROBUSTO
# =====================================================

import math
import pandas as pd


def calculate_ratios(metrics, verbose=False):

    ratios = {}
    debug_rows = []

    # =================================================
    # HELPERS BASE
    # =================================================

    def is_missing(x):
        return x is None or (isinstance(x, float) and math.isnan(x))

    def fmt_value(x):
        if is_missing(x):
            return None
        return x

    def safe_div(a, b):
        if is_missing(a) or is_missing(b) or b == 0:
            return None
        return a / b

    def calc_yoy(current, previous):
        if is_missing(current) or is_missing(previous) or previous == 0:
            return None
        return (current - previous) / abs(previous)

    def trace(msg):
        if verbose:
            print(msg)

    def metric_value(metric_name, period):
        metric = metrics.get(metric_name)

        if not isinstance(metric, dict):
            return None

        value = metric.get(period)

        if is_missing(value):
            return None

        return value

    def get_value(period, *metric_names, label=None):
        """
        Devuelve el primer valor no vacío entre varios nombres de metrics.
        """
        for name in metric_names:
            val = metric_value(name, period)

            if not is_missing(val):
                if label:
                    trace(f"[FOUND] {label} | period='{period}' | source='{name}' | value={val}")
                return val

        if label:
            trace(f"[MISSING] {label} | period='{period}' | aliases={metric_names}")

        return None

    def subtract_required(a, b):
        """
        Resta solo si ambos existen.
        Evita inventar ceros.
        """
        if is_missing(a) or is_missing(b):
            return None
        return a - b

    def sum_required(*values):
        """
        Suma solo si todos existen.
        """
        if any(is_missing(v) for v in values):
            return None
        return sum(values)

    def register_debug(
        ratio_name,
        current,
        previous,
        formula_desc,
        dependencies,
        status,
        note=None
    ):
        debug_rows.append({
            "ratio": ratio_name,
            "current": fmt_value(current),
            "previous": fmt_value(previous),
            "yoy": calc_yoy(current, previous),
            "status": status,
            "formula": formula_desc,
            "dependencies": ", ".join(dependencies) if dependencies else None,
            "note": note
        })

    def build_metric(name, func, formula_desc="", dependencies=None, note_if_missing=None):
        """
        Calcula una métrica/ratio para current y previous.
        """
        try:
            trace("\n==============================")
            trace(f"[BUILD] Calculando: {name}")

            current = func("current")
            previous = func("previous")

            trace(f"[BUILD] Resultado {name} -> current={current}, previous={previous}")

        except Exception as e:
            current = None
            previous = None

            ratios[name] = {
                "current": None,
                "previous": None,
                "yoy": None
            }

            register_debug(
                ratio_name=name,
                current=None,
                previous=None,
                formula_desc=formula_desc,
                dependencies=dependencies or [],
                status="error",
                note=str(e)
            )

            trace(f"[ERROR] {name}: {e}")
            return

        ratios[name] = {
            "current": current,
            "previous": previous,
            "yoy": calc_yoy(current, previous)
        }

        status = "ok" if (current is not None or previous is not None) else "not_calculated"

        note = None
        if status == "not_calculated":
            note = note_if_missing or "No se ha podido calcular con las métricas disponibles."

        register_debug(
            ratio_name=name,
            current=current,
            previous=previous,
            formula_desc=formula_desc,
            dependencies=dependencies or [],
            status=status,
            note=note
        )

    # =================================================
    # SNAPSHOT INICIAL
    # =================================================

    trace("\n############################################")
    trace("### SNAPSHOT INICIAL")
    trace("############################################")
    trace(f"[INFO] Keys disponibles en metrics: {sorted(list(metrics.keys()))}")

    # =================================================
    # RESOLVERS DE MÉTRICAS BASE
    # =================================================

    def get_revenue(p):
        return get_value(
            p,
            "revenue",
            label="REVENUE"
        )

    def get_cogs(p):
        """
        COGS se devuelve en positivo para márgenes.
        """
        cogs = get_value(
            p,
            "cost_of_goods_sold",
            label="COGS"
        )

        if cogs is None:
            return None

        return abs(cogs)

    def get_gross_profit(p):
        """
        Gross profit:
        1) directo
        2) revenue - abs(cogs)
        """
        direct = get_value(
            p,
            "gross_profit",
            label="GROSS_PROFIT_DIRECT"
        )

        if direct is not None:
            return direct

        revenue = get_revenue(p)
        cogs = get_cogs(p)

        if revenue is not None and cogs is not None:
            return revenue - cogs

        return None

    def get_ebit(p):
        """
        EBIT:
        1) directo
        2) profit_before_tax + abs(interest_expense)
        """
        direct = get_value(
            p,
            "ebit",
            label="EBIT_DIRECT"
        )

        if direct is not None:
            return direct

        pbt = get_value(
            p,
            "profit_before_tax",
            label="PBT_FOR_EBIT"
        )

        interest = get_value(
            p,
            "interest_expense",
            label="INTEREST_FOR_EBIT"
        )

        if pbt is not None and interest is not None:
            return pbt + abs(interest)

        return None

    def get_depreciation_amortization(p):
        da = get_value(
            p,
            "depreciation_amortization",
            label="DA"
        )

        if da is None:
            return None

        return abs(da)

    def get_ebitda(p):
          """
          EBITDA:
          1) directo
          2) EBIT + D&A
          3) reconstruido desde PyG operativo:
            operating_income_total
            - abs(cost_of_goods_sold)
            - abs(personnel_expenses)
            - abs(operating_expenses)

          Nota:
          operating_income_total = revenue + other_gains_losses
          """
          direct = get_value(
              p,
              "ebitda",
              label="EBITDA_DIRECT"
          )

          if direct is not None:
              return direct

          ebit = get_ebit(p)
          da = get_depreciation_amortization(p)

          if ebit is not None and da is not None:
              return ebit + da

          operating_income_total = metric_value("operating_income_total", p)
          cogs = get_cogs(p)
          personnel = metric_value("personnel_expenses", p)
          opex = metric_value("operating_expenses", p)

          if (
              operating_income_total is not None
              and cogs is not None
              and personnel is not None
              and opex is not None
          ):
              return (
                  operating_income_total
                  - abs(cogs)
                  - abs(personnel)
                  - abs(opex)
              )

          return None

    def get_total_liabilities(p):
        """
        Pasivo total:
        1) total_liabilities directo o derivado
        2) current_liabilities_total + non_current_liabilities_total
        """
        direct = get_value(
            p,
            "total_liabilities",
            label="TOTAL_LIABILITIES"
        )

        if direct is not None:
            return direct

        current_liabilities = metric_value("current_liabilities_total", p)
        non_current_liabilities = metric_value("non_current_liabilities_total", p)

        return sum_required(current_liabilities, non_current_liabilities)

    def get_financial_debt_total(p):
        return get_value(
            p,
            "financial_debt_total",
            label="FINANCIAL_DEBT_TOTAL"
        )

    def get_net_debt(p):
        return get_value(
            p,
            "net_debt",
            label="NET_DEBT"
        )

    def get_working_capital(p):
        return get_value(
            p,
            "working_capital",
            label="WORKING_CAPITAL"
        )

    def get_operating_working_capital(p):
        return get_value(
            p,
            "operating_working_capital",
            label="OPERATING_WORKING_CAPITAL"
        )

    # =================================================
    # MÉTRICAS BASE RECONSTRUIDAS
    # =================================================

    build_metric(
        "gross_profit",
        lambda p: get_gross_profit(p),
        formula_desc="gross_profit directo o revenue - abs(cost_of_goods_sold)",
        dependencies=[
            "gross_profit",
            "revenue",
            "cost_of_goods_sold"
        ]
    )

    build_metric(
        "ebit",
        lambda p: get_ebit(p),
        formula_desc="ebit directo o profit_before_tax + abs(interest_expense)",
        dependencies=[
            "ebit",
            "profit_before_tax",
            "interest_expense"
        ]
    )

    build_metric(
        "ebitda",
        lambda p: get_ebitda(p),
        formula_desc="ebitda directo o ebit + abs(depreciation_amortization)",
        dependencies=[
            "ebitda",
            "ebit",
            "depreciation_amortization"
        ]
    )

    build_metric(
        "financial_debt_total",
        lambda p: get_financial_debt_total(p),
        formula_desc="financial_debt_lp + financial_debt_cp + lease_liability_lp + lease_liability_cp",
        dependencies=[
            "financial_debt_lp",
            "financial_debt_cp",
            "lease_liability_lp",
            "lease_liability_cp"
        ]
    )

    build_metric(
        "net_debt",
        lambda p: get_net_debt(p),
        formula_desc="financial_debt_total - cash_and_equivalents",
        dependencies=[
            "financial_debt_total",
            "cash_and_equivalents"
        ]
    )

    build_metric(
        "working_capital",
        lambda p: get_working_capital(p),
        formula_desc="current_assets_total - current_liabilities_total",
        dependencies=[
            "current_assets_total",
            "current_liabilities_total"
        ]
    )

    build_metric(
        "operating_working_capital",
        lambda p: get_operating_working_capital(p),
        formula_desc="inventory + trade_receivables - trade_payables",
        dependencies=[
            "inventory",
            "trade_receivables",
            "trade_payables"
        ]
    )

    build_metric(
        "operating_income_total",
        lambda p: metric_value("operating_income_total", p),
        formula_desc="revenue + other_gains_losses",
        dependencies=[
            "revenue",
            "other_gains_losses"
        ]
    )

    # =================================================
    # RENTABILIDAD / MÁRGENES
    # =================================================

    build_metric(
        "gross_margin",
        lambda p: safe_div(get_gross_profit(p), get_revenue(p)),
        formula_desc="gross_profit / revenue",
        dependencies=[
            "gross_profit",
            "revenue"
        ]
    )

    build_metric(
        "ebitda_margin",
        lambda p: safe_div(get_ebitda(p), get_revenue(p)),
        formula_desc="ebitda / revenue",
        dependencies=[
            "ebitda",
            "revenue"
        ]
    )

    build_metric(
        "ebit_margin",
        lambda p: safe_div(get_ebit(p), get_revenue(p)),
        formula_desc="ebit / revenue",
        dependencies=[
            "ebit",
            "revenue"
        ]
    )

    build_metric(
        "net_margin",
        lambda p: safe_div(metric_value("net_income", p), get_revenue(p)),
        formula_desc="net_income / revenue",
        dependencies=[
            "net_income",
            "revenue"
        ]
    )

    build_metric(
        "personnel_expenses_ratio",
        lambda p: safe_div(
            abs(metric_value("personnel_expenses", p)) if metric_value("personnel_expenses", p) is not None else None,
            get_revenue(p)
        ),
        formula_desc="abs(personnel_expenses) / revenue",
        dependencies=[
            "personnel_expenses",
            "revenue"
        ]
    )

    build_metric(
        "operating_expenses_ratio",
        lambda p: safe_div(
            abs(metric_value("operating_expenses", p)) if metric_value("operating_expenses", p) is not None else None,
            get_revenue(p)
        ),
        formula_desc="abs(operating_expenses) / revenue",
        dependencies=[
            "operating_expenses",
            "revenue"
        ]
    )

    build_metric(
        "cogs_ratio",
        lambda p: safe_div(get_cogs(p), get_revenue(p)),
        formula_desc="abs(cost_of_goods_sold) / revenue",
        dependencies=[
            "cost_of_goods_sold",
            "revenue"
        ]
    )

    build_metric(
        "roa",
        lambda p: safe_div(metric_value("net_income", p), metric_value("total_assets", p)),
        formula_desc="net_income / total_assets",
        dependencies=[
            "net_income",
            "total_assets"
        ]
    )

    build_metric(
        "roe",
        lambda p: safe_div(metric_value("net_income", p), metric_value("equity_total", p)),
        formula_desc="net_income / equity_total",
        dependencies=[
            "net_income",
            "equity_total"
        ]
    )

    # =================================================
    # LIQUIDEZ
    # =================================================

    build_metric(
        "current_ratio",
        lambda p: safe_div(
            metric_value("current_assets_total", p),
            metric_value("current_liabilities_total", p)
        ),
        formula_desc="current_assets_total / current_liabilities_total",
        dependencies=[
            "current_assets_total",
            "current_liabilities_total"
        ]
    )

    build_metric(
        "quick_ratio",
        lambda p: safe_div(
            subtract_required(
                metric_value("current_assets_total", p),
                metric_value("inventory", p)
            ),
            metric_value("current_liabilities_total", p)
        ),
        formula_desc="(current_assets_total - inventory) / current_liabilities_total",
        dependencies=[
            "current_assets_total",
            "inventory",
            "current_liabilities_total"
        ],
        note_if_missing="No se calcula si falta inventory, para no convertir quick ratio en current ratio por error."
    )

    build_metric(
        "cash_ratio",
        lambda p: safe_div(
            metric_value("cash_and_equivalents", p),
            metric_value("current_liabilities_total", p)
        ),
        formula_desc="cash_and_equivalents / current_liabilities_total",
        dependencies=[
            "cash_and_equivalents",
            "current_liabilities_total"
        ]
    )

    build_metric(
        "working_capital_ratio",
        lambda p: safe_div(
            get_working_capital(p),
            get_revenue(p)
        ),
        formula_desc="working_capital / revenue",
        dependencies=[
            "working_capital",
            "revenue"
        ]
    )

    # =================================================
    # ENDEUDAMIENTO / SOLVENCIA
    # =================================================

    build_metric(
        "debt_ratio",
        lambda p: safe_div(
            get_total_liabilities(p),
            metric_value("total_assets", p)
        ),
        formula_desc="total_liabilities / total_assets",
        dependencies=[
            "total_liabilities",
            "total_assets"
        ]
    )

    build_metric(
        "debt_to_equity",
        lambda p: safe_div(
            get_total_liabilities(p),
            metric_value("equity_total", p)
        ),
        formula_desc="total_liabilities / equity_total",
        dependencies=[
            "total_liabilities",
            "equity_total"
        ]
    )

    build_metric(
        "equity_ratio",
        lambda p: safe_div(
            metric_value("equity_total", p),
            metric_value("total_assets", p)
        ),
        formula_desc="equity_total / total_assets",
        dependencies=[
            "equity_total",
            "total_assets"
        ]
    )

    build_metric(
        "financial_debt_to_equity",
        lambda p: safe_div(
            get_financial_debt_total(p),
            metric_value("equity_total", p)
        ),
        formula_desc="financial_debt_total / equity_total",
        dependencies=[
            "financial_debt_total",
            "equity_total"
        ]
    )

    build_metric(
        "financial_debt_to_ebitda",
        lambda p: safe_div(
            get_financial_debt_total(p),
            get_ebitda(p)
        ),
        formula_desc="financial_debt_total / ebitda",
        dependencies=[
            "financial_debt_total",
            "ebitda"
        ]
    )

    build_metric(
        "net_debt_to_ebitda",
        lambda p: safe_div(
            get_net_debt(p),
            get_ebitda(p)
        ),
        formula_desc="net_debt / ebitda",
        dependencies=[
            "net_debt",
            "ebitda"
        ]
    )

    build_metric(
        "short_term_debt_ratio",
        lambda p: safe_div(
            metric_value("financial_debt_cp", p),
            get_financial_debt_total(p)
        ),
        formula_desc="financial_debt_cp / financial_debt_total",
        dependencies=[
            "financial_debt_cp",
            "financial_debt_total"
        ]
    )

    build_metric(
        "current_liabilities_to_total_liabilities",
        lambda p: safe_div(
            metric_value("current_liabilities_total", p),
            get_total_liabilities(p)
        ),
        formula_desc="current_liabilities_total / total_liabilities",
        dependencies=[
            "current_liabilities_total",
            "total_liabilities"
        ]
    )

    # =================================================
    # COBERTURA
    # =================================================

    build_metric(
        "interest_coverage_ebit",
        lambda p: safe_div(
            get_ebit(p),
            abs(metric_value("interest_expense", p)) if metric_value("interest_expense", p) is not None else None
        ),
        formula_desc="ebit / abs(interest_expense)",
        dependencies=[
            "ebit",
            "interest_expense"
        ]
    )

    build_metric(
        "interest_coverage_ebitda",
        lambda p: safe_div(
            get_ebitda(p),
            abs(metric_value("interest_expense", p)) if metric_value("interest_expense", p) is not None else None
        ),
        formula_desc="ebitda / abs(interest_expense)",
        dependencies=[
            "ebitda",
            "interest_expense"
        ]
    )

    build_metric(
        "pbt_to_interest",
        lambda p: safe_div(
            metric_value("profit_before_tax", p),
            abs(metric_value("interest_expense", p)) if metric_value("interest_expense", p) is not None else None
        ),
        formula_desc="profit_before_tax / abs(interest_expense)",
        dependencies=[
            "profit_before_tax",
            "interest_expense"
        ]
    )

    # =================================================
    # EFICIENCIA / ACTIVIDAD
    # =================================================

    build_metric(
        "asset_turnover",
        lambda p: safe_div(
            get_revenue(p),
            metric_value("total_assets", p)
        ),
        formula_desc="revenue / total_assets",
        dependencies=[
            "revenue",
            "total_assets"
        ]
    )

    build_metric(
        "receivables_ratio",
        lambda p: safe_div(
            metric_value("trade_receivables", p),
            get_revenue(p)
        ),
        formula_desc="trade_receivables / revenue",
        dependencies=[
            "trade_receivables",
            "revenue"
        ]
    )

    build_metric(
        "inventory_ratio",
        lambda p: safe_div(
            metric_value("inventory", p),
            get_revenue(p)
        ),
        formula_desc="inventory / revenue",
        dependencies=[
            "inventory",
            "revenue"
        ]
    )

    build_metric(
        "payables_ratio",
        lambda p: safe_div(
            metric_value("trade_payables", p),
            get_revenue(p)
        ),
        formula_desc="trade_payables / revenue",
        dependencies=[
            "trade_payables",
            "revenue"
        ]
    )

    build_metric(
        "operating_working_capital_ratio",
        lambda p: safe_div(
            get_operating_working_capital(p),
            get_revenue(p)
        ),
        formula_desc="operating_working_capital / revenue",
        dependencies=[
            "inventory",
            "trade_receivables",
            "trade_payables",
            "revenue"
        ]
    )

    # =================================================
    # ESTRUCTURA DE BALANCE
    # =================================================

    build_metric(
        "current_assets_weight",
        lambda p: safe_div(
            metric_value("current_assets_total", p),
            metric_value("total_assets", p)
        ),
        formula_desc="current_assets_total / total_assets",
        dependencies=[
            "current_assets_total",
            "total_assets"
        ]
    )

    build_metric(
        "non_current_assets_weight",
        lambda p: safe_div(
            metric_value("non_current_assets_total", p),
            metric_value("total_assets", p)
        ),
        formula_desc="non_current_assets_total / total_assets",
        dependencies=[
            "non_current_assets_total",
            "total_assets"
        ]
    )

    build_metric(
        "intangibles_weight",
        lambda p: safe_div(
            metric_value("intangible_assets", p),
            metric_value("total_assets", p)
        ),
        formula_desc="intangible_assets / total_assets",
        dependencies=[
            "intangible_assets",
            "total_assets"
        ]
    )

    build_metric(
        "goodwill_to_equity",
        lambda p: safe_div(
            metric_value("goodwill", p),
            metric_value("equity_total", p)
        ),
        formula_desc="goodwill / equity_total",
        dependencies=[
            "goodwill",
            "equity_total"
        ]
    )

    # =================================================
    # PRINT RESUMEN FINAL
    # =================================================

    trace("\n############################################")
    trace("### RESUMEN FINAL")
    trace("############################################")

    for p in ["current", "previous"]:
        revenue_final = get_revenue(p)
        ebitda_final = get_ebitda(p)
        ebit_final = get_ebit(p)
        net_income_final = metric_value("net_income", p)
        debt_final = get_financial_debt_total(p)
        net_debt_final = get_net_debt(p)

        trace(
            f"[SUMMARY] period='{p}' | "
            f"Revenue={revenue_final} | "
            f"EBITDA={ebitda_final} | "
            f"EBIT={ebit_final} | "
            f"Net Income={net_income_final} | "
            f"Financial Debt={debt_final} | "
            f"Net Debt={net_debt_final}"
        )

    # =================================================
    # TABLA DEBUG / RESUMEN
    # =================================================

    ratios_debug_table = pd.DataFrame(debug_rows)

    if verbose:
        trace("\n############################################")
        trace("### TABLA DEBUG RATIOS")
        trace("############################################")
        try:
            print(ratios_debug_table[["ratio", "current", "previous", "yoy", "status", "note"]])
        except Exception:
            print(ratios_debug_table)

    return ratios, ratios_debug_table

# =====================================================
# INSIGHTS AVANZADOS
# =====================================================

def build_relational_insights(metrics, ratios):

    insights = []

    # =================================================
    # HELPERS
    # =================================================

    def val(metric_name, period):
        return metrics.get(metric_name, {}).get(period)

    def ratio_val(ratio_name, period="current"):
        return ratios.get(ratio_name, {}).get(period)

    def yoy_metric(metric_name):
        return calc_yoy(val(metric_name, "current"), val(metric_name, "previous"))

    def add_insight(text):
        if text and text not in insights:
            insights.append(text)

    def has_strong_change(metric_name, threshold=0.25):
        yoy = yoy_metric(metric_name)
        return yoy is not None and abs(yoy) >= threshold

    def classify_change(yoy):
        if yoy is None:
            return None
        if yoy >= 0.50:
            return "muy fuerte aumento"
        elif yoy >= 0.25:
            return "aumento relevante"
        elif yoy <= -0.50:
            return "muy fuerte caída"
        elif yoy <= -0.25:
            return "caída relevante"
        return None

    # =================================================
    # CRECIMIENTO
    # =================================================

    revenue_yoy = yoy_metric("revenue")

    if revenue_yoy is not None:
        if revenue_yoy > 0.15:
            add_insight("Fuerte crecimiento de ingresos (>15%)")
        elif revenue_yoy > 0.05:
            add_insight("Crecimiento moderado de ingresos")
        elif revenue_yoy < -0.10:
            add_insight("Caída significativa de ingresos")
        else:
            add_insight("Ingresos estables")

    # =================================================
    # MÁRGENES
    # =================================================

    ebitda_margin = ratio_val("ebitda_margin")
    net_margin = ratio_val("net_margin")
    gross_margin = ratio_val("gross_margin")
    ebit_margin = ratio_val("ebit_margin")

    if gross_margin is not None:
        if gross_margin < 0:
            add_insight("Margen bruto negativo o inconsistente")
        elif gross_margin < 0.20:
            add_insight("Margen bruto reducido")
        elif gross_margin > 0.50:
            add_insight("Margen bruto elevado")

    if ebitda_margin is not None:
        if ebitda_margin < 0:
            add_insight("EBITDA negativo (negocio no rentable operativamente)")
        elif ebitda_margin < 0.10:
            add_insight("Rentabilidad operativa baja")
        elif ebitda_margin > 0.25:
            add_insight("Alta rentabilidad operativa")

    if ebit_margin is not None:
        if ebit_margin < 0:
            add_insight("EBIT negativo")
        elif ebit_margin < 0.08:
            add_insight("Margen EBIT ajustado")

    if net_margin is not None:
        if net_margin < 0:
            add_insight("La empresa presenta pérdidas")
        elif net_margin < 0.05:
            add_insight("Margen neto muy ajustado")
        elif net_margin > 0.15:
            add_insight("Elevada rentabilidad neta")

    # =================================================
    # CALIDAD DE BENEFICIO
    # =================================================

    if ebitda_margin is not None and net_margin is not None:
        if ebitda_margin > 0.15 and net_margin < 0.03:
            add_insight("Posible impacto elevado de amortizaciones, intereses o impuestos sobre el resultado final")

    # =================================================
    # LIQUIDEZ
    # =================================================

    current_ratio = ratio_val("current_ratio")
    quick_ratio = ratio_val("quick_ratio")
    cash_ratio = ratio_val("cash_ratio")

    if current_ratio is not None:
        if current_ratio < 1:
            add_insight("Riesgo de liquidez a corto plazo")
        elif current_ratio < 1.5:
            add_insight("Liquidez ajustada")
        else:
            add_insight("Buena liquidez")

    if quick_ratio is not None and quick_ratio < 1:
        add_insight("Dependencia de inventario para cubrir obligaciones")

    if cash_ratio is not None and cash_ratio < 0.2:
        add_insight("Cobertura de caja reducida frente a pasivos corrientes")

    # =================================================
    # ENDEUDAMIENTO
    # =================================================

    debt_to_equity = ratio_val("debt_to_equity")
    debt_ratio = ratio_val("debt_ratio")

    if debt_to_equity is not None:
        if debt_to_equity > 2:
            add_insight("Alto apalancamiento financiero")
        elif debt_to_equity > 1:
            add_insight("Apalancamiento moderado")

    if debt_ratio is not None and debt_ratio > 0.7:
        add_insight("Alta dependencia de financiación externa")

    # =================================================
    # EFICIENCIA
    # =================================================

    asset_turnover = ratio_val("asset_turnover")

    if asset_turnover is not None:
        if asset_turnover < 0.5:
            add_insight("Baja rotación de activos")
        elif asset_turnover > 1.5:
            add_insight("Alta eficiencia en el uso de activos")

    # =================================================
    # CAPITAL CIRCULANTE
    # =================================================

    wc = ratio_val("working_capital")
    receivables_ratio = ratio_val("receivables_ratio")
    payables_ratio = ratio_val("payables_ratio")
    inventory_ratio = ratio_val("inventory_ratio")

    if wc is not None:
        if wc < 0:
            add_insight("Capital circulante negativo (potencial tensión)")
        else:
            add_insight("Capital circulante positivo")

    if receivables_ratio is not None and payables_ratio is not None:
        if receivables_ratio > payables_ratio:
            add_insight("Mayor peso de clientes que proveedores (consumo de caja)")
        else:
            add_insight("Financiación operativa vía proveedores")

    if inventory_ratio is not None and inventory_ratio > 0.25:
        add_insight("Peso relevante de inventario sobre ventas")

    # =================================================
    # RENTABILIDAD SOBRE CAPITAL
    # =================================================

    roe = ratio_val("roe")
    roa = ratio_val("roa")

    if roe is not None:
        if roe < 0.05:
            add_insight("Baja rentabilidad para accionistas (ROE bajo)")
        elif roe > 0.20:
            add_insight("Alta rentabilidad para accionistas (ROE elevado)")

    if roa is not None and roa < 0.03:
        add_insight("Rentabilidad reducida sobre activos")

    # =================================================
    # INSIGHTS EXPLICATIVOS POR VARIACIONES FUERTES
    # =================================================

    # -------------------------------------------------
    # 1. CAJA
    # -------------------------------------------------

    cash_yoy = yoy_metric("cash_and_equivalents")
    if cash_yoy is not None and abs(cash_yoy) >= 0.30:
        change_label = classify_change(cash_yoy)

        explanations = []

        inv_group_yoy = yoy_metric("investments_group_long_term")
        inv_fin_yoy = yoy_metric("financial_investments_long_term")
        debt_cp_yoy = yoy_metric("current_liabilities_total")
        debt_lp_yoy = yoy_metric("non_current_liabilities_total")
        receivables_yoy = yoy_metric("trade_receivables")
        inventory_yoy = yoy_metric("inventory")
        equity_yoy = yoy_metric("equity_total")
        net_income_yoy = yoy_metric("net_income")

        if inv_group_yoy is not None and inv_group_yoy > 0.20:
            explanations.append("aumento de inversiones a largo plazo o en empresas del grupo")
        if inv_fin_yoy is not None and inv_fin_yoy > 0.20:
            explanations.append("incremento de inversiones financieras")
        if debt_cp_yoy is not None and debt_cp_yoy < -0.20:
            explanations.append("reducción de pasivos corrientes, compatible con amortización o pagos")
        if debt_lp_yoy is not None and debt_lp_yoy < -0.20:
            explanations.append("reducción de deuda a largo plazo")
        if receivables_yoy is not None and receivables_yoy > 0.20:
            explanations.append("aumento de clientes/deudores, que podría absorber caja")
        if inventory_yoy is not None and inventory_yoy > 0.20:
            explanations.append("acumulación de inventario")
        if net_income_yoy is not None and net_income_yoy < -0.30:
            explanations.append("deterioro del resultado neto")
        if equity_yoy is not None and equity_yoy < -0.10:
            explanations.append("movimientos patrimoniales que convendría revisar")

        if cash_yoy < 0:
            if explanations:
                add_insight(
                    f"Se observa una {change_label} de la caja, que podría estar relacionada con: "
                    + "; ".join(explanations) + "."
                )
            else:
                add_insight(
                    "Se observa una caída relevante de la caja, pero con la información disponible en balance y PyG "
                    "no puede determinarse con suficiente fiabilidad la causa exacta; conviene validarlo con el cliente."
                )
        else:
            if explanations:
                add_insight(
                    f"Se observa un {change_label} de la caja, posiblemente apoyado por: "
                    + "; ".join(explanations) + "."
                )
            else:
                add_insight("Fuerte aumento de caja respecto al ejercicio anterior.")

    # -------------------------------------------------
    # 2. DEUDA
    # -------------------------------------------------

    total_debt_current = (val("current_liabilities_total", "current") or 0) + (val("non_current_liabilities_total", "current") or 0)
    total_debt_previous = (val("current_liabilities_total", "previous") or 0) + (val("non_current_liabilities_total", "previous") or 0)
    total_debt_yoy = calc_yoy(total_debt_current, total_debt_previous)

    if total_debt_yoy is not None and abs(total_debt_yoy) >= 0.25:
        explanations = []

        cash_yoy = yoy_metric("cash_and_equivalents")
        total_assets_yoy = yoy_metric("total_assets")
        revenue_yoy = yoy_metric("revenue")
        net_income_yoy = yoy_metric("net_income")

        if total_debt_yoy > 0:
            if cash_yoy is not None and cash_yoy > 0.15:
                explanations.append("refuerzo de liquidez")
            if total_assets_yoy is not None and total_assets_yoy > 0.15:
                explanations.append("crecimiento del balance o financiación de inversiones")
            if revenue_yoy is not None and revenue_yoy < -0.10:
                explanations.append("posibles necesidades de financiación ante debilidad operativa")
            if net_income_yoy is not None and net_income_yoy < -0.20:
                explanations.append("presión sobre la generación interna de recursos")

            if explanations:
                add_insight(
                    "La deuda aumenta de forma relevante, posiblemente en relación con "
                    + "; ".join(explanations) + "."
                )
            else:
                add_insight(
                    "La deuda aumenta de forma relevante y convendría revisar si responde a financiación de inversión, "
                    "necesidades operativas o tensiones de caja."
                )
        else:
            if cash_yoy is not None and cash_yoy < -0.20:
                explanations.append("salidas de caja asociadas a amortización de deuda")
            if explanations:
                add_insight(
                    "La deuda se reduce de forma relevante, lo que podría estar relacionado con "
                    + "; ".join(explanations) + "."
                )
            else:
                add_insight("Reducción relevante del endeudamiento respecto al ejercicio anterior.")

    # -------------------------------------------------
    # 3. CLIENTES / DEUDORES
    # -------------------------------------------------

    receivables_yoy = yoy_metric("trade_receivables")
    if receivables_yoy is not None and abs(receivables_yoy) >= 0.25:
        if receivables_yoy > 0:
            explanations = []

            if revenue_yoy is not None and revenue_yoy > 0.10:
                explanations.append("crecimiento de ventas")
            if cash_yoy is not None and cash_yoy < -0.20:
                explanations.append("mayor consumo de caja por circulante")

            if explanations:
                add_insight(
                    "Los clientes/deudores aumentan de forma relevante, posiblemente por "
                    + "; ".join(explanations) + "."
                )
            else:
                add_insight(
                    "Los clientes/deudores aumentan de forma relevante, lo que podría reflejar mayores plazos de cobro "
                    "o presión sobre el circulante."
                )
        else:
            add_insight("Reducción relevante de clientes/deudores, con posible liberación de circulante.")

    # -------------------------------------------------
    # 4. INVENTARIO
    # -------------------------------------------------

    inventory_yoy = yoy_metric("inventory")
    if inventory_yoy is not None and abs(inventory_yoy) >= 0.25:
        if inventory_yoy > 0:
            explanations = []
            if revenue_yoy is not None and revenue_yoy < 0.05:
                explanations.append("crecimiento de inventario superior al de ventas")
            if cash_yoy is not None and cash_yoy < -0.20:
                explanations.append("posible consumo de caja")

            if explanations:
                add_insight(
                    "El inventario aumenta de forma relevante, posiblemente por "
                    + "; ".join(explanations) + "."
                )
            else:
                add_insight("Aumento relevante del inventario, que conviene revisar por su impacto en liquidez y rotación.")
        else:
            add_insight("Reducción relevante del inventario respecto al ejercicio anterior.")

    # -------------------------------------------------
    # 5. PATRIMONIO NETO
    # -------------------------------------------------

    equity_yoy = yoy_metric("equity_total")
    if equity_yoy is not None and abs(equity_yoy) >= 0.20:
        net_income_current = val("net_income", "current")
        if equity_yoy > 0:
            if net_income_current is not None and net_income_current > 0:
                add_insight("El patrimonio neto aumenta de forma relevante, previsiblemente apoyado por resultados positivos o refuerzo patrimonial.")
            else:
                add_insight("El patrimonio neto aumenta de forma relevante; conviene revisar si se debe a aportaciones, ajustes patrimoniales o resultados.")
        else:
            add_insight("El patrimonio neto cae de forma relevante, lo que puede responder a pérdidas, distribución al accionista o ajustes patrimoniales.")

    # -------------------------------------------------
    # 6. INGRESOS VS RENTABILIDAD
    # -------------------------------------------------

    ebitda_yoy = calc_yoy(
        metrics.get("ebitda", {}).get("current"),
        metrics.get("ebitda", {}).get("previous")
    )

    if revenue_yoy is not None and ebitda_yoy is not None:
        if revenue_yoy > 0.10 and ebitda_yoy < 0:
            add_insight("Los ingresos crecen, pero el EBITDA empeora, lo que sugiere presión en costes o deterioro de márgenes.")
        elif revenue_yoy < 0 and ebitda_yoy > 0.10:
            add_insight("Aunque los ingresos retroceden, el EBITDA mejora, lo que sugiere ajuste de costes o mejora del mix de negocio.")

    # -------------------------------------------------
    # 7. MÁRGENES Y COSTES
    # -------------------------------------------------

    gross_margin_prev = ratio_val("gross_margin", "previous")
    if gross_margin is not None and gross_margin_prev is not None:
        gross_margin_delta = gross_margin - gross_margin_prev
        if gross_margin_delta < -0.05:
            add_insight("Deterioro relevante del margen bruto, compatible con presión en precios, mix o coste de ventas/aprovisionamientos.")
        elif gross_margin_delta > 0.05:
            add_insight("Mejora relevante del margen bruto.")

    ebitda_margin_prev = ratio_val("ebitda_margin", "previous")
    if ebitda_margin is not None and ebitda_margin_prev is not None:
        ebitda_margin_delta = ebitda_margin - ebitda_margin_prev
        if ebitda_margin_delta < -0.05:
            add_insight("Deterioro relevante del margen EBITDA, lo que apunta a mayor presión operativa.")
        elif ebitda_margin_delta > 0.05:
            add_insight("Mejora relevante del margen EBITDA.")

    return insights

import os
import json

# INFORME IA - VERSION EJECUTIVA + RIESGOS
# =====================================================

def generate_ai_financial_report(metrics, ratios, insights, warnings_balance=None, warnings_pyg=None):
    if warnings_balance is None:
        warnings_balance = []
    if warnings_pyg is None:
        warnings_pyg = []

    input_payload = {
        "metrics": metrics,
        "ratios": ratios,
        "insights": insights,
        "warnings_balance": warnings_balance,
        "warnings_pyg": warnings_pyg
    }

    prompt = f"""
Eres un analista financiero senior y debes redactar un informe ejecutivo en español.

Objetivo:
Redactar un análisis financiero claro, profesional y útil para negocio a partir de métricas, ratios, variaciones interanuales, insights deterministas y observaciones técnicas.

ESTILO:
- Tono ejecutivo, claro y profesional.
- Orientado a dirección financiera o gerencia.
- No sonar técnico de programación.
- No listar JSON ni nombres internos de variables.
- No inventar datos ni conclusiones no soportadas.

ESTRUCTURA OBLIGATORIA DEL INFORME:

1. RESUMEN EJECUTIVO
- Explica de forma general la situación financiera y operativa de la empresa.
- Resume la evolución del negocio entre ambos años.
- Destaca si la lectura general es positiva, neutral o preocupante.

2. PRINCIPALES RIESGOS IDENTIFICADOS
Agrupa los riesgos identificados en categorías cuando aplique, por ejemplo:
- Riesgos operativos
- Riesgos financieros
- Riesgos de liquidez
- Riesgos de crédito / circulante
- Riesgos de rentabilidad
Describe solo los riesgos que realmente se desprendan de los ratios, métricas e insights proporcionados.

3. FORTALEZAS O ASPECTOS POSITIVOS
- Señala fortalezas relevantes detectadas en los datos.
- Por ejemplo: buena liquidez, crecimiento de ingresos, rentabilidad sólida, bajo endeudamiento, etc.

4. INSIGHTS CLAVE
- Resume los insights más importantes de forma ejecutiva.
- Prioriza los que ayuden a entender el negocio y sus riesgos.

5. OBSERVACIONES TÉCNICAS
- Incluye aquí los warnings técnicos.
- Si alguna validación no se ha podido realizar, dilo de forma neutra.
- Si hay limitaciones de datos, explícalas brevemente.
- Esta sección no debe bloquear ni invalidar el análisis, solo contextualizarlo.

INSTRUCCIONES IMPORTANTES:
- Usa los ratios, métricas y variaciones para justificar el análisis.
- Si hay una variación interanual relevante, coméntala.
- Si faltan datos para algún punto, indícalo con naturalidad.
- No menciones campos como canonical, aggregate, payload, final_section, etc.
- No hagas tablas.
- Redacta en párrafos y bullets cortos si ayuda a la claridad.
- Sé más ejecutivo que académico.

DATOS DE ENTRADA:
{json.dumps(input_payload, ensure_ascii=False, indent=2)}
"""

    response = get_openai_client().responses.create(
        model="gpt-5.4",
        input=[
            {
                "role": "system",
                "content": "Eres un analista financiero senior especializado en redacción de informes ejecutivos en español."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return response.output_text

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# =====================================================
# GUARDAR INFORME EN DOCX - VERSION CON SUBTITULOS DE RIESGO VISIBLES
# =====================================================

def save_full_financial_report_to_docx(
    ai_report,
    insights,
    warnings_balance=None,
    warnings_pyg=None,
    output_path="informe_financiero_completo.docx"
):
    if warnings_balance is None:
        warnings_balance = []
    if warnings_pyg is None:
        warnings_pyg = []

    doc = Document()

    # -------------------------
    # Título principal
    # -------------------------
    title = doc.add_heading("Informe financiero automático", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -------------------------
    # Helpers de detección
    # -------------------------
    main_sections = {
        "resumen ejecutivo",
        "principales riesgos identificados",
        "fortalezas o aspectos positivos",
        "insights clave",
        "observaciones técnicas"
    }

    def is_main_section(line):
        clean = line.strip().lower().rstrip(":")
        return clean in main_sections

    def is_risk_subsection(line):
        clean = line.strip().lower().rstrip(":")
        return clean.startswith("riesgo") or clean.startswith("riesgos")

    def is_bullet(line):
        stripped = line.strip()
        return (
            stripped.startswith("- ")
            or stripped.startswith("• ")
            or stripped.startswith("* ")
        )

    def clean_bullet(line):
        stripped = line.strip()
        if len(stripped) >= 2 and stripped[:2] in ["- ", "• ", "* "]:
            return stripped[2:].strip()
        return stripped

    def add_normal_paragraph(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)

    def add_bullet(text, level=1):
        style_name = "List Bullet" if level == 1 else "List Bullet 2"
        p = doc.add_paragraph(style=style_name)
        run = p.add_run(text)
        run.font.size = Pt(11)

    def add_risk_subheading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)

    # -------------------------
    # Procesado del informe IA
    # -------------------------
    doc.add_heading("Informe ejecutivo", level=1)

    lines = [line.strip() for line in ai_report.splitlines() if line.strip()]

    in_risk_section = False
    current_risk_subsection = None

    for line in lines:
        # 1. Sección principal
        if is_main_section(line):
            doc.add_heading(line.rstrip(":"), level=1)

            if line.strip().lower().rstrip(":") == "principales riesgos identificados":
                in_risk_section = True
            else:
                in_risk_section = False

            current_risk_subsection = None
            continue

        # 2. Subtítulo de riesgo
        if in_risk_section and is_risk_subsection(line) and not is_bullet(line):
            add_risk_subheading(line.rstrip(":"))
            current_risk_subsection = line
            continue

        # 3. Bullet
        if is_bullet(line):
            bullet_text = clean_bullet(line)

            if in_risk_section and current_risk_subsection is not None:
                add_bullet(bullet_text, level=2)
            else:
                add_bullet(bullet_text, level=1)
            continue

        # 4. Texto normal
        add_normal_paragraph(line)

    # -------------------------
    # Insights utilizados
    # -------------------------
    doc.add_heading("Insights utilizados", level=1)
    if insights:
        for insight in insights:
            add_bullet(str(insight), level=1)
    else:
        add_normal_paragraph("No se generaron insights.")

    # -------------------------
    # Observaciones técnicas balance
    # -------------------------
    doc.add_heading("Observaciones técnicas balance", level=1)
    if warnings_balance:
        for w in warnings_balance:
            msg = w.get("message", "")
            add_bullet(msg, level=1)
    else:
        add_normal_paragraph("No se han detectado incidencias técnicas relevantes en balance.")

    # -------------------------
    # Observaciones técnicas PyG
    # -------------------------
    doc.add_heading("Observaciones técnicas PyG", level=1)
    if warnings_pyg:
        for w in warnings_pyg:
            msg = w.get("message", "")
            add_bullet(msg, level=1)
    else:
        add_normal_paragraph("No se han detectado incidencias técnicas relevantes en PyG.")

    doc.save(output_path)
    return output_path


# =====================================================
# FUNCIÓN PRINCIPAL PARA STREAMLIT
# =====================================================

def run_credit_risk_analysis(file_path):
    """
    Función principal llamada desde app.py.
    """
    warnings = []

    try:
        structured_pyg_df = assign_levels_from_excel_openpyxl(
            file_path=file_path,
            sheet_name="PyG",
            verbose=False,
            latest_position="left"
        )

        structured_balance_df = assign_levels_from_excel_openpyxl(
            file_path=file_path,
            sheet_name="Balance",
            verbose=False,
            latest_position="left"
        )

        mapped_rows_balance = map_lines_with_context(
            structured_balance_df.copy(),
            "balance"
        )
        final_rows_balance = initialize_final_columns(mapped_rows_balance)
        final_rows_balance = fill_none_with_ai(final_rows_balance, verbose=False)
        validation_results_balance = run_all_validations(final_rows_balance)

        mapped_rows_pyg = map_lines_with_context(
            structured_pyg_df.copy(),
            "pyg"
        )
        final_rows_pyg = initialize_final_columns(mapped_rows_pyg)
        final_rows_pyg = fill_none_with_ai(final_rows_pyg, verbose=False)

        warnings_balance = collect_analysis_warnings(
            final_rows_balance,
            validation_results_balance
        )
        warnings_pyg = collect_analysis_warnings(
            final_rows_pyg,
            None
        )
        warnings = warnings_balance + warnings_pyg

        metrics = extract_financial_metrics(
            final_rows_balance,
            final_rows_pyg,
            debug=False
        )

        ratios, ratios_debug_table = calculate_ratios(
            metrics,
            verbose=False
        )

        insights = build_relational_insights(metrics, ratios)

        report_text = generate_ai_financial_report(
            metrics=metrics,
            ratios=ratios,
            insights=insights,
            warnings_balance=warnings_balance,
            warnings_pyg=warnings_pyg
        )

        output_dir = tempfile.mkdtemp()
        word_path = os.path.join(output_dir, "informe_financiero_completo.docx")

        save_full_financial_report_to_docx(
            ai_report=report_text,
            insights=insights,
            warnings_balance=warnings_balance,
            warnings_pyg=warnings_pyg,
            output_path=word_path
        )

        return {
            "warnings": warnings,
            "report_text": report_text,
            "ratios": ratios,
            "ratios_debug_table": ratios_debug_table,
            "output_files": {
                "Informe Word": word_path
            },
            "debug": {
                "structured_balance_shape": structured_balance_df.shape,
                "structured_pyg_shape": structured_pyg_df.shape,
                "final_balance_rows": len(final_rows_balance),
                "final_pyg_rows": len(final_rows_pyg),
                "validation_results_balance": validation_results_balance,
                "insights": insights,
            }
        }

    except Exception as e:
        return {
            "warnings": [
                {
                    "type": "pipeline_error",
                    "severity": "high",
                    "message": str(e)
                }
            ],
            "report_text": f"Ha ocurrido un error durante la generación del informe: {str(e)}",
            "ratios": None,
            "ratios_debug_table": None,
            "output_files": {},
            "debug": {}
        }
